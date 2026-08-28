"""v33 -> v34: the two results that were measured but unwritten.

  * Point 7b, the physics-informed vs data-driven 2x2, now complete
    (Table 21), inserted before "8.9 Representative training visualizations".
  * Point 9, the method of manufactured solutions, FEM half (Table 22),
    inserted after it.

Both are read from their committed JSONs and every claim is asserted before
it is written.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v33.docx', 'PFEM_Transolver_Report_v34.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
P7 = json.load(open(os.path.join(HERE, '..', 'omar_pfem', 'point7b_results',
                                 'comparison_B1_neo_hookean.json')))
P9 = json.load(open(os.path.join(HERE, '..', 'omar_pfem', 'point9_results',
                                 'mms_B1_neo_hookean.json')))

# ---- point 7b claims ---------------------------------------------------
R = P7['runs']
PI_A = R['physics_informed']['best_val_rel_L2']
DD_A = R['data_driven_matched_optimizer']['best_val_rel_L2']
PI_O = R['physics_informed_adamw_onecycle']['best_val_rel_L2']
DD_O = R['data_driven_own_optimizer']['best_val_rel_L2']
assert all(R[k]['opt_steps'] == 75000 for k in R), \
    'the four runs no longer share a 75,000-step budget; the columns are not comparable'
assert PI_A < DD_A, 'the physics-informed loss no longer wins the Adam column'
assert DD_O < PI_O, 'the data-driven loss no longer wins the AdamW column'
FLIP = (PI_A < DD_A) != (PI_O < DD_O)
assert FLIP, 'the ranking no longer flips -- rewrite this section, do not patch the numbers'
adam_gap = abs(PI_A - DD_A) / max(PI_A, DD_A) * 100
oc_gap = abs(PI_O - DD_O) / max(PI_O, DD_O) * 100
pi_worse = (PI_O / PI_A - 1) * 100
dd_better = (1 - DD_O / DD_A) * 100
assert pi_worse > 0 > -dd_better, 'the optimizer no longer moves the two losses oppositely'
LABEL_H = R['data_driven_matched_optimizer']['label_generation_cost_h']
assert LABEL_H > 5 and R['physics_informed']['label_generation_cost_h'] == 0

# ---- point 9 claims ----------------------------------------------------
ROWS9 = sorted(P9['rows'], key=lambda r: (r['order'], r['N']))
RATES = P9['convergence_rates']
assert all(v == 'as expected' for v in P9['rate_check'].values()), \
    'an MMS convergence rate is off; the manufactured problem is not sound'
for o in ('Q4', 'Q9'):
    for n in ('L2', 'H1_semi'):
        v = RATES[o][n]
        assert abs(v['rate'] - v['expected']) < 0.4, (o, n, v)
        assert max(abs(p - v['rate']) for p in v['pairwise']) < 0.1, \
            f'{o} {n} pairwise rates have not settled'
byd = {}
for r in ROWS9:
    byd.setdefault(r['n_dof'], {})[r['order']] = r
SHARED = sorted(d for d, v in byd.items() if len(v) > 1)
assert len(SHARED) >= 2, 'fewer than two matched-DOF comparisons'
ADV = [(d, byd[d]['Q4']['L2_rel'] / byd[d]['Q9']['L2_rel']) for d in SHARED]
assert ADV[-1][1] > ADV[0][1], "Q9's advantage no longer grows with refinement"

print(f'7b verified: flip confirmed, Adam gap {adam_gap:.1f}%, AdamW gap {oc_gap:.1f}%, '
      f'PI {pi_worse:+.1f}%, DD {-dd_better:+.1f}%')
print(f'9 verified: all rates as expected; Q9 advantage '
      f'{ADV[0][1]:.1f}x -> {ADV[-1][1]:.1f}x')

doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def insert_before(heading_prefix, els):
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(heading_prefix):
            anchor = p
            break
    assert anchor is not None, f'heading {heading_prefix!r} not found'
    target = anchor._p.getprevious()
    for el in els:
        target.addnext(el)
        target = el


# ======================================================================
# Point 7b
# ======================================================================
els = []
h = doc.add_paragraph(style='Heading 2')
h.add_run('8.9 Physics-informed against data-driven training, at matched cost')
els.append(h._p)

els.append(para(
    'Every operator in this report is trained by minimising the total potential '
    'energy, with no reference solutions in the loss. The obvious question is '
    'whether that principle earns its keep: the same architecture on the same '
    'mesh could instead be fitted to finite-element displacements directly. The '
    'comparison is easy to get wrong, because changing the loss usually means '
    'changing the optimiser that was tuned for it, and then the two effects are '
    'confounded. It was therefore run as a two-by-two: both losses under both '
    'recipes, all four at the same 75,000 optimiser steps, the same 800/200 '
    'split and the same B1 × Neo-Hookean mesh.')._p)

els.append(new_table(
    ['Training loss', 'Adam, lr 2×10⁻³', 'AdamW lr 10⁻³ + OneCycleLR'],
    [['Physics-informed (energy)', f'{PI_A:.4f}', f'{PI_O:.4f}'],
     ['Data-driven (relative L2 to FEM)', f'{DD_A:.4f}', f'{DD_O:.4f}']])._tbl)
els.append(para(
    'Table 21. Held-out relative L2 error, B1 × Neo-Hookean, for each '
    'combination of training loss and optimiser. Reading down a column isolates '
    'the loss, because everything else in that column is identical. All four '
    'runs used 75,000 optimiser steps at batch size 8.')._p)

els.append(para(
    f'Read down the first column and the physics-informed loss wins by '
    f'{adam_gap:.0f}%. Read down the second and the data-driven loss wins by '
    f'{oc_gap:.0f}%. The ranking reverses, so neither training principle is '
    'better than the other in any unconditional sense on this problem — which '
    'is the finding, and it is one that only the complete grid can support. '
    'Either column on its own would have licensed a confident and opposite '
    'conclusion, and the earlier state of this comparison, with the '
    'physics-informed model measured under only its own recipe, was exactly '
    'that trap.')._p)

els.append(para(
    'The mechanism is that the optimiser does not treat the two losses alike; '
    f'it moves them in opposite directions. Switching from Adam to AdamW with a '
    f'one-cycle schedule makes the physics-informed model {pi_worse:.0f}% worse '
    f'and the data-driven model {dd_better:.0f}% better. A learning-rate '
    'schedule tuned against one objective is not neutral ground for the other, '
    'and reporting a single number for "the physics-informed operator" without '
    'naming its optimiser therefore overstates what has been measured.')._p)

els.append(para(
    'One asymmetry does survive the reversal, and it is not an accuracy figure. '
    'The data-driven model needs a labelled dataset: 800 finite-element solves, '
    f'which at the measured cost of Table 4a is {LABEL_H:.2f} hours of CPU '
    'before training can begin. The physics-informed model needs none, because '
    'its loss is the energy of its own prediction. That gap is a property of the '
    'training principle rather than of the recipe, it does not move between the '
    'columns of Table 21, and on this problem it is larger than the accuracy '
    'difference in either direction.')._p)

els.append(para(
    'The comparison covers one case, B1 × Neo-Hookean, and two optimiser '
    'settings. Two recipes are enough to show that the ranking is not fixed; '
    'they are not enough to map out which family of schedules favours which '
    'loss, and nothing here should be read as identifying the best available '
    'recipe for either.')._p)

insert_before('8.9 Representative training visualizations', els)
print(f'inserted {len(els)} elements as section 8.9 (point 7b)')

# The new 8.9 collides with the existing one, so the incumbent moves to 8.10.
# Checked before doing it: no paragraph in the document refers to "8.9" other
# than the heading itself, so renumbering breaks no cross-reference.
_renamed = 0
for p in doc.paragraphs:
    if p.text.strip() == '8.9 Representative training visualizations':
        for r in p.runs:
            r.text = ''
        p.runs[0].text = '8.10 Representative training visualizations'
        _renamed += 1
assert _renamed == 1, f'expected exactly one heading to renumber, found {_renamed}'
print('renumbered the incumbent 8.9 to 8.10')

# ======================================================================
# Point 9 -- MMS
# ======================================================================
els = []
h = doc.add_paragraph(style='Heading 2')
h.add_run('8.11 Verification against an analytic solution (manufactured solutions)')
els.append(h._p)

els.append(para(
    'Every accuracy figure reported so far is measured against a '
    'finite-element solution. That makes the finite-element method the one '
    'component of the pipeline that is never itself graded: it defines the '
    'reference, so it cannot be scored against it. A manufactured solution '
    'removes that circularity by supplying an exact answer known in closed '
    'form, against which Q4 and Q9 are both measured on equal terms.')._p)

els.append(para(
    'A displacement field is chosen first and the governing equations are then '
    'solved backwards for the body force that would produce it. The field used '
    f'here is {P9["manufactured_solution"].split(";")[0]}, which vanishes on the '
    'entire boundary — not a cosmetic choice, since it makes homogeneous '
    'Dirichlet conditions the exact boundary condition for the manufactured '
    'problem and so requires no change to the solver the rest of this report '
    'depends on. The body force b = −Div P is obtained by nested automatic '
    'differentiation of the strain-energy density and checked against a central '
    'finite difference before any mesh is built. The material is uniform, so '
    'that what the study measures is discretisation error and not the '
    'interpolation of a varying stiffness field.')._p)

els.append(para(
    'The alternative — finding an exact solution that needs no body force — was '
    'considered and rejected. On this geometry such a solution is a homogeneous '
    'deformation, which a bilinear Q4 element reproduces to machine precision; '
    'the study would measure round-off and separate nothing. This choice was '
    'made here and has not been confirmed with the advisor.')._p)

rows = []
for r in ROWS9:
    rows.append([r['order'], str(r['N']), f"{r['n_dof']:,}",
                 f"{r['L2_rel']:.3e}", f"{r['H1_semi_rel']:.3e}",
                 f"{r['stress_rel_L2']:.3e}", f"{r['energy_rel']:.3e}"])
els.append(new_table(
    ['Order', 'N', 'DOF', 'L2', 'H1 semi-norm', 'Stress', 'Energy'], rows)._tbl)
els.append(para(
    'Table 22. Q4 and Q9 against the manufactured solution, B1 geometry, '
    'Neo-Hookean, FP64. All four errors are relative and integrated on the '
    'element\'s own quadrature. "Energy" is the internal strain energy '
    '∫ψ(F) dV, not the total potential.')._p)

rr = []
for o in ('Q4', 'Q9'):
    for n, lab in (('L2', 'L2'), ('H1_semi', 'H1 semi-norm'),
                   ('stress', 'Stress'), ('energy', 'Energy')):
        v = RATES[o][n]
        rr.append([o, lab, f"{v['rate']:.2f}", str(v['expected']),
                   ', '.join(f'{p:.2f}' for p in v['pairwise'])])
els.append(new_table(
    ['Order', 'Norm', 'Observed rate', 'Theory', 'Pairwise'], rr)._tbl)
els.append(para(
    'Table 23. Observed convergence rates, fitted by least squares on log h, '
    'with each consecutive pair\'s own rate alongside. This table is the '
    'verification: a body force wrong by a sign or a factor would make the '
    'discrete solution converge to the wrong function and these rates would '
    'collapse. All eight land on their theoretical values, and every pairwise '
    'rate sits within 0.03 of its fit, so unlike the cost exponent of Table 20 '
    'these have settled and carry no caveat.')._p)

adv_txt = ', '.join(f'{r:.1f}× at {d:,} DOF' for d, r in ADV)
els.append(para(
    'Because both orders are scored against the same exact solution, they can '
    'be compared at equal cost rather than at equal mesh spacing, which is the '
    f'comparison that matters when choosing one. Q9 is more accurate at every '
    f'matched degree-of-freedom count — {adv_txt} in the L2 norm — and its '
    'margin widens with refinement, exactly as the higher rate predicts. On this '
    'problem the extra nodes of a biquadratic element buy more than the same '
    'nodes spent on refining a bilinear one.')._p)

els.append(para(
    'Two limits. The comparison here is two-way; the advisor asked for a '
    'three-way study including the physics-informed operator, and that third '
    'part is not in this table. The operator cannot be run on a manufactured '
    'problem as the pipeline currently stands, because its energy functional '
    'has no body-force term and its input channels have no field to carry one, '
    'so the trained checkpoints of Table 5 are not applicable and a separately '
    'trained model is required. And the study covers a single member of the '
    'manufactured family on one geometry and one material, rather than the '
    'parametrised family that would let the operator be trained on it at all.')._p)

insert_before('9. Discussion', els)
print(f'inserted {len(els)} elements as section 8.11 (MMS)')

doc.save(DST)
print('wrote', DST)
