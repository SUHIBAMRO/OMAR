"""Adds the OOD attribution (report Table 19) to the parallel summary."""
import copy, json, os
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HERE='/home/user/OMAR/Practical_Examples/report_builders'
R=json.load(open(os.path.join(HERE,'..','omar_pfem','point6_results',
                              'ood_progressive_B1_neo_hookean.json')))
BASE=next(r['mean_rel_L2'] for r in R['rows'] if r['factor']=='baseline')
KS=[0.5,1.0,1.5,2.0,2.5,3.0]
def col(f):
    d={r['shift_sigma']:r for r in R['rows'] if r['factor']==f}
    return [d[k] for k in KS]
L,M,B=col('loading'),col('material'),col('both')
ld=[r['degradation_vs_baseline'] for r in L]; md=[r['degradation_vs_baseline'] for r in M]
bd=[r['degradation_vs_baseline'] for r in B]
assert max(ld)<1.10 and md[-1]>5.0

doc=Document('PFEM_Summary_Completed_Work.pre_v5.docx')
ref=next(doc.element.body.iter(qn('w:tbl')))
def new_table(header,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(header)); t.style=doc.tables[0].style
    pr=ref.find(qn('w:tblPr'))
    if pr is not None:
        old=t._tbl.find(qn('w:tblPr'))
        if old is not None: t._tbl.remove(old)
        t._tbl.insert(0,copy.deepcopy(pr))
    for j,h in enumerate(header):
        c=t.cell(0,j); c.text=''; c.paragraphs[0].add_run(h).bold=True
    for i,row in enumerate(rows,1):
        for j,v in enumerate(row): t.cell(i,j).text=v
    return t
def para(t):
    p=doc.add_paragraph(); p.add_run(t); return p

rows=[['0.0 (ID)',f'{BASE:.4f}','—','—','—','—','—']]
for i,k in enumerate(KS):
    rows.append([f'{k}',f"{L[i]['mean_rel_L2']:.4f}",f'{ld[i]:.2f}×',
                 f"{M[i]['mean_rel_L2']:.4f}",f'{md[i]:.2f}×',
                 f"{B[i]['mean_rel_L2']:.4f}",f'{bd[i]:.2f}×'])

hs=None
for p in doc.paragraphs:
    if p.text.strip().startswith('6. Out-of-distribution'): hs=p.style; break
assert hs is not None

els=[para('Shift in units of the training distribution’s own std, one factor at '
          'a time, 10 held-out samples per point, B1 × Neo-Hookean, N=21. '
          'Same per-component metric as Table 11.')._p,
     new_table(['k (σ)','Loading','×','Material','×','Both','×'],rows)._tbl,
     para(f'Table 19. Loading causes no degradation anywhere ({min(ld):.2f}–{max(ld):.2f}× '
          f'over the whole sweep); material causes all of it ({md[0]:.2f}–{md[-1]:.2f}×). '
          f'Table 11’s 4.11× for this case sits inside the combined column '
          f'({bd[3]:.2f}× at k=2.0, {bd[4]:.2f}× at k=2.5). Smooth, no threshold. '
          f'The two shifts partially cancel: “both” is {(1-bd[-1]/md[-1])*100:.0f}% below '
          '“material” at k=3.')._p]

anchor=None
for p in doc.paragraphs:
    if p.text.strip().startswith('6. Out-of-distribution'): anchor=p._p; break
cur=last=anchor
while True:
    nxt=cur.getnext()
    if nxt is None: break
    if nxt.tag==qn('w:p'):
        pp=Paragraph(nxt,doc)
        if pp.style is not None and pp.style==hs and pp.text.strip(): break
    last=nxt; cur=nxt
target=last
for el in els: target.addnext(el); target=el
doc.save('PFEM_Summary_Completed_Work.docx')
print(f'inserted {len(els)} elements into summary section 6')
