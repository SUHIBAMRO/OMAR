"""v46 -> v47. Fixes three duplicate table numbers found by
check_report_tables.py: Table 3, 4 and 5 each captioned two entirely
different tables. The four mesh-convergence captions (B1xMR, B1xAB,
B2xMR, B2xAB -- currently 3, 4, 5, 6) are renumbered to 1a, 1b, 2a, 2b,
alongside Table 1 (B1xNH) and Table 2 (B2xNH) they belong with. Nothing
else is touched: the "real" Table 3 (hyperparameters), 4 (training
protocol, 2 inline references), 5 (best validation error, 11 inline
references) and 6/"6 (revised)" (batch-size sweep) keep their numbers
unchanged, since renumbering THEM would mean chasing down and rewording
13+ references elsewhere versus 0 for the mesh-convergence quartet.

The one collateral fix: paragraph "Across all six ... (Tables 1-6)"
named the mesh-convergence set as a numeric range, which no longer holds
once the four are 1a/1b/2a/2b instead of 3-6. Reworded to name the new
labels directly.

Verified afterward by re-running check_report_tables.py on the output --
zero duplicates must remain.
"""
from docx import Document

SRC, DST = 'PFEM_Transolver_Report_v46.docx', 'PFEM_Transolver_Report_v47.docx'

doc = Document(SRC)
ORIGINAL = list(doc.paragraphs)


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; expected exactly 1')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


# The four mesh-convergence captions -- distinguished from their
# duplicate-number siblings by their unique full text (each is the only
# paragraph with this exact wording, unlike bare "Table 3."/"Table 4."/
# "Table 5." which each match twice).
old_3 = find_para('Table 3. B1 × Mooney-Rivlin mesh convergence')
assert old_3.text.strip().startswith('Table 3. B1 × Mooney-Rivlin')
retext('Table 3. B1 × Mooney-Rivlin mesh convergence',
       old_3.text.replace('Table 3.', 'Table 1a.', 1))

old_4 = find_para('Table 4. B1 × Arruda-Boyce mesh convergence')
retext('Table 4. B1 × Arruda-Boyce mesh convergence',
       old_4.text.replace('Table 4.', 'Table 1b.', 1))

old_5 = find_para('Table 5. B2 × Mooney-Rivlin mesh convergence')
retext('Table 5. B2 × Mooney-Rivlin mesh convergence',
       old_5.text.replace('Table 5.', 'Table 2a.', 1))

old_6 = find_para('Table 6. B2 × Arruda-Boyce mesh convergence')
retext('Table 6. B2 × Arruda-Boyce mesh convergence',
       old_6.text.replace('Table 6.', 'Table 2b.', 1))

# The compound reference to the old numeric range.
old_131 = find_para('Across all six (geometry, material) combinations')
assert 'Tables 1–6' in old_131.text or 'Tables 1-6' in old_131.text
retext('Across all six (geometry, material) combinations',
       old_131.text.replace('Tables 1–6', 'Tables 1, 1a, 1b, 2, 2a, 2b')
                    .replace('Tables 1-6', 'Tables 1, 1a, 1b, 2, 2a, 2b'))

doc.save(DST)
print(f'wrote {DST}')
print('Renamed: Table 3(B1xMR)->1a, Table 4(B1xAB)->1b, '
      'Table 5(B2xMR)->2a, Table 6(B2xAB)->2b')
print('Reworded the "Tables 1-6" compound reference at paragraph', ORIGINAL.index(old_131))
