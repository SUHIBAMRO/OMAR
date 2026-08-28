"""v31 -> v32: where the out-of-distribution degradation actually comes from.

Timon's round-6 point 1: separate material from loading, and sweep the shift
progressively rather than reporting one ID/OOD pair.

Goes into section 8.6, immediately after Table 11, because it explains that
table's own numbers. Table 19, continuing 15-18.

The claims are asserted against the result JSON before they are written --
the same mechanism that caught four false statements in earlier revisions.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v31.docx', 'PFEM_Transolver_Report_v32.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
RES = json.load(open(os.path.join(
    HERE, '..', 'omar_pfem', 'point6_results', 'ood_progressive_B1_neo_hookean.json')))

BASE = next(r['mean_rel_L2'] for r in RES['rows'] if r['factor'] == 'baseline')
KS = [0.5, 1.0, 1.5, 2.0, 2.5, 3.0]


def col(factor):
    d = {r['shift_sigma']: r for r in RES['rows'] if r['factor'] == factor}
    return [d[k] for k in KS]


LOAD, MAT, BOTH = col('loading'), col('material'), col('both')

# ---- the claims, verified before they are written ----------------------
load_deg = [r['degradation_vs_baseline'] for r in LOAD]
mat_deg = [r['degradation_vs_baseline'] for r in MAT]
both_deg = [r['degradation_vs_baseline'] for r in BOTH]
assert max(load_deg) < 1.10, f'loading DOES degrade: up to {max(load_deg):.2f}x'
assert mat_deg == sorted(mat_deg), 'material degradation is not monotone'
assert mat_deg[-1] > 5.0, f'material tops out at only {mat_deg[-1]:.2f}x'
assert all(b <= m * 1.02 for b, m in zip(both_deg[2:], mat_deg[2:])), \
    'both is not at or below material alone at k>=1.5'
incr = [mat_deg[i + 1] - mat_deg[i] for i in range(len(mat_deg) - 1)]
assert all(incr[i] <= incr[i + 1] + 1e-9 for i in range(len(incr) - 1)), \
    'material increments are not monotonically increasing'
print(f'claims verified: loading {min(load_deg):.2f}-{max(load_deg):.2f}x, '
      f'material {mat_deg[0]:.2f}-{mat_deg[-1]:.2f}x, '
      f'both at k=3.0 is {(1 - both_deg[-1] / mat_deg[-1]) * 100:.0f}% below material')

doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


rows = [['0.0 (in-distribution)', f'{BASE:.4f}', '—', '—', '—', '—', '—']]
for i, k in enumerate(KS):
    rows.append([f'{k}', f'{LOAD[i]["mean_rel_L2"]:.4f}', f'{load_deg[i]:.2f}×',
                 f'{MAT[i]["mean_rel_L2"]:.4f}', f'{mat_deg[i]:.2f}×',
                 f'{BOTH[i]["mean_rel_L2"]:.4f}', f'{both_deg[i]:.2f}×'])

els = []
h = doc.add_paragraph(style='Heading 3')
h.add_run('Which factor causes the degradation, and how it accumulates')
els.append(h._p)

els.append(para(
    'Table 11 shifts material stiffness and loading together and reports a '
    'single number per case. That measurement cannot say which of the two '
    'factors is responsible, nor whether the damage appears gradually or at a '
    'threshold, and both questions matter for whether anything can be done '
    'about it. The two were therefore separated and swept. Each shift is '
    'expressed in units of the training distribution\'s own standard deviation, '
    'so k = 2 means the mean of that field has moved two standard deviations '
    'from where the network was trained; the sweep runs from k = 0 to k = 3 in '
    'half-sigma steps, with ten held-out samples per point, on the study\'s own '
    'N=21 mesh. Poisson\'s ratio is deliberately excluded: the parametric field '
    'clips it to (0.2, 0.4), so a shifted mean saturates against the clip and '
    'any curve drawn through it would be measuring the clip rather than the '
    'physics. Errors use the per-component definition of Tables 5 and 11, so '
    'the degradation columns read directly against Table 11\'s factors.')._p)

els.append(new_table(
    ['Shift k (σ)', 'Loading', '×', 'Material', '×', 'Both', '×'], rows)._tbl)
els.append(para(
    'Table 19. Progressive out-of-distribution shift, B1 × Neo-Hookean, one '
    'factor at a time. Each entry is the mean relative L2 error over ten '
    'held-out samples, with its ratio to the in-distribution baseline. '
    '"Loading" shifts the traction mean, "material" the Young\'s modulus mean, '
    '"both" shifts them together — which is what Table 11 measured.')._p)

els.append(para(
    'The separation is unambiguous. Shifting the loading produces no '
    f'measurable degradation anywhere in the sweep: the error stays between '
    f'{min(load_deg):.2f}× and {max(load_deg):.2f}× of its in-distribution '
    'value across the full three-sigma range, which is sample noise at ten '
    'samples rather than a trend. Shifting the material carries all of it, '
    f'from {mat_deg[0]:.2f}× at half a sigma to {mat_deg[-1]:.2f}× at three. '
    'The 4.11× reported for this case in Table 11, measured with both factors '
    'shifted at roughly two to two and a half sigma, falls exactly inside the '
    f'combined column here ({both_deg[3]:.2f}× at k=2.0, {both_deg[4]:.2f}× at '
    'k=2.5), so the two studies agree and this one identifies the cause.')._p)

els.append(para(
    'The damage accumulates smoothly. Degradation increments per half sigma '
    f'are {", ".join(f"{x:.2f}" for x in incr)} — monotonically increasing, '
    'with no threshold and no cliff. There is no shift small enough to be safe; '
    'there is only a slope, which is a less comfortable finding than a '
    'threshold would have been because it means the operator has no stated '
    'domain of validity, only a rate at which it degrades outside the one it '
    'was trained on.')._p)

els.append(para(
    'The two shifts do not compound. From k = 1.5 upward the combined column '
    'is consistently below the material column, by '
    f'{(1 - both_deg[-1] / mat_deg[-1]) * 100:.0f}% at k = 3. A plausible '
    'mechanism, offered as a hypothesis because nothing in this experiment '
    'isolates it, is that the two move the solution\'s scale in opposite '
    'directions: raising the modulus stiffens the body and shrinks the '
    'displacement field, while raising the traction magnitude grows it, so a '
    'simultaneous shift partially cancels in the quantity being scored.')._p)

els.append(para(
    'Why the two factors behave so differently is worth stating, because it '
    'points at what could be done. Displacement depends on load magnitude '
    'close to proportionally — scale the traction and, away from strong '
    'geometric nonlinearity, the field scales with it. A network that has '
    'learned that proportionality extrapolates it without difficulty, which is '
    'what the flat loading row shows. Stiffness enters inversely: displacement '
    'scales roughly as 1/E. An inverse dependence learned across a narrow band '
    'of E does not extrapolate, because the part of the curve being asked for '
    'was never seen. That asymmetry — linear in load, inverse in stiffness — '
    'is the simplest explanation consistent with these measurements.')._p)

els.append(para(
    'It also identifies where a mitigation would have to act. Since the '
    'degradation enters entirely through E, and E is one of the network\'s four '
    'input channels, the cheapest candidate is to normalize that channel by '
    'the training distribution\'s own mean and standard deviation, so the '
    'network sees a standardized quantity rather than a raw stiffness; that '
    'costs one training run and no new finite-element data. Widening the '
    'training range of E is the alternative, but it requires regenerating the '
    'dataset and does not remove the extrapolation problem, only moves its '
    'boundary. A third option — having the network predict a scaled quantity '
    'such as u·E, so the inverse dependence is removed from what must be '
    'learned — is a design change rather than a setting. Only the first has '
    'been costed here; none has been tested, and this section reports the '
    'diagnosis rather than a fix.')._p)

els.append(para(
    'This diagnosis covers B1 × Neo-Hookean. Whether the same attribution '
    'holds for the other five cases has not been measured, and the B2 rows of '
    'Table 11 in particular degrade by a different factor.')._p)

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('8.7 Resolution invariance'):
        anchor = p
        break
assert anchor is not None, 'section 8.7 heading not found'
target = anchor._p.getprevious()
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements at the end of section 8.6')

doc.save(DST)
print('wrote', DST)
