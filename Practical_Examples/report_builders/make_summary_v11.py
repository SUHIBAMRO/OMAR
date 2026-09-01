"""Mirrors report v38 into the summary.

Two edits, matching make_v38.py's two substantive ones and reading the same
JSONs so the documents cannot disagree:

  * "That did not fix them ... the cause is under investigation" is replaced
    by the cause, which was our own selection metric, and by what B2 x
    Neo-Hookean does once it is corrected -- including the part that is not
    good news: the old flatness was insensitivity, and the corrected model's
    spread across the mesh is 3.8x where B1's worst is 2.1x;
  * "one scored member of the family" is replaced by the family sweep, which
    keeps the finding and sharpens it.

Run from the directory holding PFEM_Summary_Completed_Work.docx; it copies
the current file to .pre_v11.docx first.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

FIX = json.load(open(os.path.join(PF, 'point7a_results',
                                  'B2_zeroshot_fixedselection.json')))
FAM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_family_fem_B1_neo_hookean.json')))
B1 = {m: json.load(open(os.path.join(PF, 'point7a_results',
                                     f'zeroshot_B1_{m}.json')))
      for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')}

ROWS = FIX['rows']
NS = [r['N'] for r in ROWS]
NEW_PC = [r['mean_rel_L2_vs_fine_reference'] for r in ROWS]
NEW_BC = [r['mean_combined_rel_L2_vs_fine_reference'] for r in ROWS]
OLD_PC = [r['superseded_mean_rel_L2_vs_fine_reference'] for r in ROWS]
TR = FIX['training']

for c in B1.values():
    assert [r['N'] for r in c['rows']] == NS, 'a B1 case is on other meshes'
B1_ALL = [r['mean_rel_L2_vs_fine_reference']
          for c in B1.values() for r in c['rows']]
B1_SPREAD = max(max(r['mean_rel_L2_vs_fine_reference'] for r in c['rows'])
                / min(r['mean_rel_L2_vs_fine_reference'] for r in c['rows'])
                for c in B1.values())
B2_SPREAD = max(NEW_PC) / min(NEW_PC)
OLD_SPREAD_PCT = (max(OLD_PC) / min(OLD_PC) - 1) * 100
MID = [25, 29, 37]
B2_MID = [p for N, p in zip(NS, NEW_PC) if N in MID]
B1_MID = [r['mean_rel_L2_vs_fine_reference']
          for c in B1.values() for r in c['rows'] if r['N'] in MID]
assert B2_SPREAD > B1_SPREAD, (B2_SPREAD, B1_SPREAD)
assert OLD_SPREAD_PCT < 0.2, OLD_SPREAD_PCT
for a, b in zip(NEW_PC, OLD_PC):
    assert a < b

FROWS = {r['N']: r for r in FAM['rows']}
FNS = sorted(FROWS)
RATE = FAM['observed_rates_N9_to_N33_two_point']
PERI = FAM['per_interval_rates']
RATIO = {N: FROWS[N]['operator_over_Q4_on_the_family']['L2_rel'] for N in FNS}
SINGLE_RATIO = {N: (FROWS[N]['operator_single_member_as_published']['L2_rel']
                    / FROWS[N]['Q4']['L2_rel']) for N in FNS}
OP_GROWTH = (FROWS[33]['operator_family_mean']['L2_rel']
             / FROWS[9]['operator_family_mean']['L2_rel'])
Q4_FALL = FROWS[9]['Q4']['L2_rel'] / FROWS[33]['Q4']['L2_rel']
assert RATE['operator']['L2_rel'] < 0
assert abs(RATE['Q4']['L2_rel'] - 1.98) < 0.05, RATE['Q4']['L2_rel']
for iv in ('N9_to_N17', 'N17_to_N33'):
    assert PERI['operator'][iv]['L2_rel'] < 0, iv
assert abs(SINGLE_RATIO[17] - 2.42) < 0.02, SINGLE_RATIO[17]

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v11.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v11.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def replace(prefix, els):
    v = find(prefix)
    t = v._p
    for el in els:
        t.addnext(el)
        t = el
    v._p.getparent().remove(v._p)


# ---- 1. B2: the cause was ours, and the case works -------------------
replace('The three B2 cases are NOT here.', [
    para(
        'B2 was excluded from Table 12 twice, and the second exclusion was our '
        'own mistake. First its sample caches were found to carry an applied '
        'load overstated by a mesh-dependent factor — about 13× at N=21 '
        'against 21× at N=33 — which for a study measuring transfer across '
        'meshes invalidates them outright. The caches were repaired and '
        'checked (one fixed pressure field now assembles to 0.007% across the '
        'two meshes) and all three cases retrained under the B1 protocol with '
        'the per-sample force normalisation section 9.1 requires. They still '
        'reported best validation errors of 0.9986, 0.9752 and 1.0267 — what '
        'predicting zero scores — and that was written up as an unexplained '
        'failure on the curved geometry.')._p,
    para(
        'The cause was the criterion used to stop training, not the models. '
        'Validation used the per-component average, which divides each '
        'displacement component by its own magnitude. On B2 the per-sample '
        'ratio of the two component magnitudes averages 1.90 while the ratio '
        'of the averaged components is 0.90 — a skewed distribution, so the '
        'mean of the per-sample ratios reports its tail — and the resulting '
        'quantity rises over intervals where the field is in fact getting '
        'closer to the reference. Every B2 run exhausted its patience at its '
        'first or second validation event, epochs 25, 25 and 225, so every '
        'downstream B2 diagnosis — the load, the ramp, the batch size, the '
        'functional, joint training — was measuring a model trained for a few '
        'dozen epochs. B1 is unaffected: its two metrics differ by a stable '
        '1.36–1.71× and all three B1 cases agree on which checkpoint is '
        'better, so every B1 figure stands as measured.')._p,
    para(
        'Re-running B2 × Neo-Hookean under exactly the Table 12 protocol, '
        'changing only the selection criterion to the combined norm ‖e‖/‖u‖ '
        'that Table 18 already uses, moves validation from 0.9986 to '
        f'{TR["per_component_val_error_at_that_checkpoint"]:.4f} on the same '
        'metric that had condemned it, and to '
        f'{TR["best_both_components_val_error"]:.4f} on the combined norm. Its '
        f'best checkpoint is at epoch {TR["best_epoch"]:,} against a previous '
        'best at epoch 25. Zero-shot at Table 12\'s seven unseen resolutions, '
        'per-component:')._p,
    new_table(['N', 'before', 'after', 'combined norm, after'],
              [[str(N), f'{o:.4f}', f'{n:.4f}', f'{b:.4f}']
               for N, o, n, b in zip(NS, OLD_PC, NEW_PC, NEW_BC)])._tbl,
    para(
        'Read honestly, that is a working case and not a match for B1. In the '
        f'middle of the range the two are comparable — {min(B2_MID):.4f} to '
        f'{max(B2_MID):.4f} at N = 25, 29 and 37 against B1\'s '
        f'{min(B1_MID):.4f} to {max(B1_MID):.4f} — but B2 degrades at both '
        f'ends, {NEW_PC[0]:.4f} at N = 13 and {NEW_PC[-1]:.4f} at N = 49, for '
        f'a spread of {B2_SPREAD:.1f}× across the sweep where the worst B1 '
        f'case spreads {B1_SPREAD:.1f}×. Training was at N = 21 and 33 and B2 '
        'falls away from those meshes in both directions. Its resolution '
        'invariance is real and weaker than B1\'s.')._p,
    para(
        'One correction to how the old numbers were described. The superseded '
        f'model\'s error was flat across the mesh — {min(OLD_PC):.4f} to '
        f'{max(OLD_PC):.4f}, {OLD_SPREAD_PCT:.3f}% across a fourfold '
        'refinement, where the B1 columns move by 79%–111%. That flatness '
        'reads like the strongest resolution invariance in the report and is '
        'the opposite: a model whose output barely responds to its input gives '
        'nearly the same field, and so nearly the same error, on every mesh. '
        'Insensitivity and invariance are indistinguishable in that column. '
        'Flatness alone is not evidence of the property being tested.')._p,
    para(
        'Only Neo-Hookean has been re-run. Mooney-Rivlin and Arruda-Boyce '
        'carry the identical defect and no B2 zero-shot number is quoted for '
        'them until they are re-run — about 3 h 46 m of A100 apiece at the '
        'measured rate. Why the corrected model degrades at the ends of the '
        'range was not measured and is not asserted.')._p,
])

# ---- 2. the MMS comparison is now on a family ------------------------
replace('What is left is optimisation error, not discretisation error', [
    para(
        'What is left is optimisation error, not discretisation error: best '
        'held-out L2 went 1.429e-02 → 8.826e-03 over the second half of '
        'training, a further 38%, still falling slowly.')._p,
    para(
        'The three-way comparison is now on a family rather than one member. '
        'Every ratio in Tables 22–24b is computed on the single manufactured '
        'member α = 0.05, β = 0.7, while the operator is scored by its mean '
        'over 16 held-out members. Q4 and Q9 were therefore solved on those '
        'same 16 members, drawn with the operator run\'s own call, at all '
        'three meshes; the single member is not among them. Q4\'s observed '
        f'rate comes out at {RATE["Q4"]["L2_rel"]:.2f} in L2 and '
        f'{RATE["Q4"]["H1_semi_rel"]:.2f} in H1, reproducing the 1.98 and 1.00 '
        'Table 23 measures on eight resolutions, which is what makes it a '
        'control:')._p,
    new_table(['N', 'Q4 (family)', 'Q9 (family)', 'operator (family)',
               'operator/Q4, family', 'operator/Q4, single member'],
              [[str(N),
                f'{FROWS[N]["Q4"]["L2_rel"]:.4e}',
                f'{FROWS[N]["Q9"]["L2_rel"]:.4e}',
                f'{FROWS[N]["operator_family_mean"]["L2_rel"]:.4e}',
                f'{RATIO[N]:.2f}×',
                f'{SINGLE_RATIO[N]:.2f}×'] for N in FNS])._tbl,
    para(
        'The finding survives and is sharper on the family: the operator\'s '
        f'error grows {OP_GROWTH:.2f}× from N = 9 to N = 33 while Q4\'s falls '
        f'{Q4_FALL:.1f}×, for observed rates of Q4 {RATE["Q4"]["L2_rel"]:.2f}, '
        f'Q9 {RATE["Q9"]["L2_rel"]:.2f} and operator '
        f'{RATE["operator"]["L2_rel"]:+.2f}. The operator\'s rate is negative '
        f'on each half separately, {PERI["operator"]["N9_to_N17"]["L2_rel"]:+.2f} '
        f'and {PERI["operator"]["N17_to_N33"]["L2_rel"]:+.2f}, so what is '
        'established is the sign everywhere rather than one exponent. Table '
        f'24\'s N = 17 comparison was representative — {RATIO[17]:.2f}× on the '
        f'family against {SINGLE_RATIO[17]:.2f}× on the member — while at N = 9 '
        'the single member was materially easier and any statement there '
        f'should quote the family\'s {RATIO[9]:.2f}×, not '
        f'{SINGLE_RATIO[9]:.2f}×. A ratio below one at N = 9 is not a defect: '
        'the ceiling constrains Π, not L2.')._p,
    para(
        'Limits: the operator\'s GPU training cost is not commensurable with '
        'the CPU Newton solves and is not forced onto a common axis; Q4\'s '
        'spread across the family is negligible (0.2–0.3% in L2) but the '
        'operator\'s cannot be reported, because the operator runs stored only '
        'their family mean; and the whole section still rests on one geometry '
        'and one material.')._p,
])

doc.save(SRC)
print(f'summary v11: B2 works ('
      f'{TR["per_component_val_error_at_that_checkpoint"]:.4f} val, '
      f'{min(NEW_PC):.4f}-{max(NEW_PC):.4f} zero-shot, spread '
      f'{B2_SPREAD:.2f}x against B1 {B1_SPREAD:.2f}x); MMS on the family '
      f'(operator/Q4 ' + ', '.join(f'{RATIO[N]:.2f}x' for N in FNS) + ')')
