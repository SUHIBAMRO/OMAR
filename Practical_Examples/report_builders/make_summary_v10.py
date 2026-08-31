"""Mirrors report v37 into the summary.

Four edits, matching make_v37.py's four and reading the same JSONs so the
two documents cannot disagree:

  * the Table 20 CG caveat gains its resolution -- CG was allowed to
    converge at N = 501 and 701, and Table 20 UNDERSTATES by +28% and +18%;
  * the section-9 ceiling is restated: it constrains Pi, it holds
    empirically in the derivative norms, and it does not hold in L2;
  * "one mesh, so the operator has no convergence rate of its own" is
    replaced by the measured rate, which is negative;
  * "the cases are being regenerated" is replaced by what happened when
    they were.

Run from the directory holding PFEM_Summary_Completed_Work.docx; it copies
the current file to .pre_v10.docx first.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')
P8 = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_scaling_B1_neo_hookean.json')))
CG = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_cg_converged_B1_neo_hookean.json')))
MMS = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_operator_rate_B1_neo_hookean.json')))
BAD = json.load(open(os.path.join(PF, 'point7a_results',
                                  'INVALID_B2_zeroshot.json')))
RETRAIN = json.load(open(os.path.join(PF, 'point7a_results',
                                      'B2_zeroshot_retrain_status.json')))

TRUNC = {r['N']: r for r in P8['rows']}
CONV = {r['N']: r for r in CG['rows']}
DELTA = {N: (CONV[N]['solve_s'] / TRUNC[N]['solve_s_in_source'] - 1) * 100
         for N in CONV}
assert all(r['cg_failures'] == 0 for r in CG['rows'])
assert all(v > 0 for v in DELTA.values()), DELTA
EXTRA = CG['EXTRAPOLATED_NOT_MEASURED']

MROWS = sorted(MMS['rows'], key=lambda r: r['N'])
RATE = MMS['fitted_rates_in_h']
RAT = MMS['ratios_operator_over_Q4']
L2R = [RAT['L2'][str(r['N'])] for r in MROWS]
assert RATE['operator_L2'] < 0 < RATE['operator_H1_semi'], RATE
assert L2R[0] < 1.0 < L2R[2], L2R
assert all(RAT['H1_semi'][str(r['N'])] > 1.0 for r in MROWS), RAT['H1_semi']

BEST = RETRAIN['result']['best_by_case']
BS = RETRAIN['candidates_tested'][1]['result']
MESH_OK = BAD['repair']['mesh_independence_check_after_repair']
assert all(abs(v['combined_val_error'] - 1.0) < 0.05 for v in BEST.values())

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v10.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v10.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def replace(prefix, els):
    v = find(prefix)
    t = v._p
    for el in els:
        t.addnext(el)
        t = el
    v._p.getparent().remove(v._p)


def after(prefix, els):
    t = find(prefix)._p
    for el in els:
        t.addnext(el)
        t = el


# ---- 1. the CG caveat is now resolved --------------------------------
after('What survives cleanly, and is the better result', [
    para(
        f'That caveat has since been closed. N = 501 and 701 were re-run with '
        f'the CG cap raised from 2,000 to 8,000 and nothing else changed. CG '
        f'converged at both — zero capped solves — and the O(N) law, fitted on '
        f'meshes no larger than N = 301, predicted '
        f'{CONV[501]["predicted_cg_per_newton"]:,.0f} and '
        f'{CONV[701]["predicted_cg_per_newton"]:,.0f} iterations per Newton '
        f'solve against the {CONV[501]["cg_per_newton"]:,.1f} and '
        f'{CONV[701]["cg_per_newton"]:,.1f} measured, 0.4% at both. The '
        f'prediction was printed before the runs were launched. Truncation '
        f'does cost Newton steps as expected — at N = 701 the count fell from '
        f'{TRUNC[701]["stats"]["newton_iters_total"]} to '
        f'{CONV[701]["newton_iters_total"]} — but not enough to pay for the '
        f'missing iterations: Table 20 UNDERSTATES the converged cost, by '
        f'{DELTA[501]:.0f}% at N = 501 ({TRUNC[501]["solve_s_in_source"]:,.0f} '
        f's → {CONV[501]["solve_s"]:,.0f} s) and {DELTA[701]:.0f}% at N = 701 '
        f'({TRUNC[701]["solve_s_in_source"]:,.0f} s → '
        f'{CONV[701]["solve_s"]:,.0f} s). The per-CG-iteration cost agrees '
        f'between the truncated and converged runs to within 1.3%, which is '
        f'two independent measurements of the quantity the O(DOF) claim rests '
        f'on. N = 1001 and 1401 were not re-run; the same model puts them at '
        f'{EXTRA["N1001"]["change_pct"]:+.0f}% and '
        f'{EXTRA["N1401"]["change_pct"]:+.0f}%, and those two are predictions.'
    )._p])

# ---- 2. the ceiling ---------------------------------------------------
replace('Read the table with the ceiling in hand', [
    para(
        f'Read the table with the ceiling in hand, and read the ceiling '
        f'correctly. The operator minimises the same functional over the same '
        f'Q4 space that the Q4 solver solves, so the Q4 solution is the '
        f'minimiser and nothing the operator produces attains a lower Π. That '
        f'is a statement about Π, and Π is none of the four errors in the '
        f'table. It does not transfer to L2: a field that fails to minimise Π '
        f'can still sit closer to the exact solution in L2, because Q4\'s '
        f'discretisation error is a systematic bias the network\'s '
        f'optimisation error may partly cancel. An earlier revision said a '
        f'ratio below one would be a defect; that is right for Π and wrong for '
        f'L2, and the three-mesh run below gives {L2R[0]:.2f}× in L2 at N = 9 '
        f'with nothing defective. What the ceiling does protect, empirically, '
        f'is the derivative norms: operator/Q4 is above one in H1 and in '
        f'stress at all three meshes. The functional itself was verified '
        f'against the solver before training — Π(u_FEM) = -7.999050 is a true '
        f'minimum, the excess grows quadratically with ratio 4.000, and a work '
        f'term mis-scaled by eight moves that minimum to 0.125.')._p])

replace('operator / Q4 = 2.42× in L2, so the ceiling holds.', [
    para(
        'operator / Q4 = 2.42× in L2 at this mesh. The finding is that the '
        'four norms disagree: 1.03× in H1 and 1.03× in stress — effectively at '
        'the Q4 optimum — against 2.42× in L2 and 3.11× in energy. That '
        'inverts the usual ordering. The loss is built from the deformation '
        'gradient, so strain and stress are what it constrains hardest and the '
        'displacement is pinned only through them; the same inversion appears '
        'in an independent N = 9 CPU run (1.35× against 4.71×). For a '
        'physics-informed operator an L2 displacement error overstates how '
        'wrong the mechanics are.')._p])

# ---- 3. the operator now has a rate, and it is negative --------------
els = [para(
    f'The operator now has a convergence rate, and it is negative. Two more '
    f'operators were trained, at N = 9 and N = 33, under exactly the N = 17 '
    f'protocol, so the three points differ in the mesh alone.')._p]
els.append(new_table(
    ['N', 'DOF', 'Operator L2', 'Q4 L2', 'op/Q4 L2', 'op/Q4 H1', 'op/Q4 stress'],
    [[str(r['N']), f"{r['n_dof']:,}", f"{r['operator']['L2']:.3e}",
      f"{r['Q4']['L2']:.3e}", f"{RAT['L2'][str(r['N'])]:.2f}×",
      f"{RAT['H1_semi'][str(r['N'])]:.2f}×",
      f"{RAT['stress'][str(r['N'])]:.2f}×"] for r in MROWS])._tbl)
els.append(para(
    f'Table 24a/24b. The operator across three meshes, same manufactured '
    f'solution and same scoring routine as Tables 22–24.')._p)
els.append(para(
    f'Fitted in h: operator L2 {RATE["operator_L2"]:+.2f} against Q4 '
    f'{RATE["Q4_L2"]:.2f}; operator H1 {RATE["operator_H1_semi"]:+.2f} against '
    f'Q4 {RATE["Q4_H1_semi"]:.2f}. Q4\'s two are the control and they land on '
    f'Table 23\'s measured 1.98 and 1.00, so the operator\'s can be quoted '
    f'beside them. The L2 error therefore gets WORSE with refinement. The '
    f'reason is that the operator\'s error is optimisation error, not '
    f'discretisation error: refining reduces what limits Q4 and leaves the '
    f'network where it was, while enlarging the problem it has to optimise. '
    f'The crossover is inside the three points — at N = 9 the operator is the '
    f'more accurate of the two in L2 ({L2R[0]:.2f}×), by N = 33 Q4 is '
    f'{L2R[2]:.1f} times better. Two qualifications: three points over a '
    f'narrow span, and the budget was held at 2,000 epochs while the problem '
    f'grew from {MROWS[0]["n_dof"]:,} to {MROWS[2]["n_dof"]:,} DOF, so this is '
    f'the rate at a fixed budget, not a property of the method at '
    f'convergence.')._p)
# The block belongs after the paragraph that reads Table 24, not between
# the table and its reading, so it goes in ahead of the closing paragraph
# rather than through after() -- that paragraph is itself replaced below and
# could not be used as an anchor afterwards.
replace('What is left is optimisation error, not discretisation error', els + [
    para(
        'What is left is optimisation error, not discretisation error: best '
        'held-out L2 went 1.429e-02 → 8.826e-03 over the second half of '
        'training, a further 38%, still falling slowly. Limits: the operator\'s '
        'GPU training cost is not commensurable with the CPU Newton solves and '
        'is not forced onto a common axis; and the whole section rests on one '
        'geometry, one material and one scored member of the family.')._p])

# ---- 4. what happened when the B2 cases were regenerated -------------
replace('The three B2 cases are NOT here.', [
    para(
        f'The three B2 cases are NOT here. Their sample caches carried an '
        f'applied load overstated by a mesh-dependent factor — about 13× at '
        f'N=21 against 21× at N=33 — which for a study that measures transfer '
        f'across meshes invalidates them outright. The caches were repaired '
        f'and checked (one fixed pressure field now assembles to '
        f'{MESH_OK["spread_pct"]:.3f}% across the two meshes) and all three '
        f'cases retrained under the B1 protocol with the per-sample force '
        f'normalisation section 9.1 requires. That did not fix them: best '
        f'combined validation error '
        + ', '.join(f'{BEST[m]["combined_val_error"]:.4f}' for m in
                    ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
        + f' against 0.0658–0.0827 for the three B1 cases. One is what '
        f'predicting zero scores on that metric. A batch-size arm at matched '
        f'optimiser steps moved it from {BS["batch_8"]:.4f} to '
        f'{BS["batch_1"]:.4f}, so it is not that either. The cause is under '
        f'investigation and no B2 zero-shot number is admissible until it is '
        f'found.')._p])

doc.save(SRC)
print('summary v10: CG resolved, ceiling corrected, operator rate, B2 status')
