"""v48 -> v49. Omar quoted a stale line back at us: "Still open: the same
high-DOF study for B2, and the five remaining resolution-invariance
cases." That sentence lives in the summary (fixed by make_summary_v23.py)
but it sent me back to check the report for the same kind of leftover
text, and it is there too, in a place the earlier table-by-table audit
did not look: the point-mapping overview table at the very top of the
report (doc.tables[0], "# / point / Headline result"), row 7.

That row still described the OLD, single-resolution protocol (B1 x
Neo-Hookean only, trained once at N=21, evaluated at five unseen
resolutions) and said "Extending the same zero-shot protocol to the
other 5 cases is in progress." Section 8.7's actual body has long since
moved past this: all six geometry x material cases are done, using the
current two-resolution joint-training protocol (N=21 and 33), evaluated
at seven unseen resolutions, not five. This table row was simply never
updated when that section was.

Rewritten to match Section 8.7's own summary sentences exactly (para
"Across all three materials and all seven unseen resolutions the error
stays between 5.0% and 10.6%" and para "Every B1/B2 x material
combination now has a valid zero-shot result... spread across the seven
meshes -- 4.91x for Mooney-Rivlin, 3.78x for Neo-Hookean, 3.62x for
Arruda-Boyce -- exceeds B1's worst case of 2.11x").
"""
from docx import Document

SRC, DST = 'PFEM_Transolver_Report_v48.docx', 'PFEM_Transolver_Report_v49.docx'

doc = Document(SRC)

t = doc.tables[0]
row = t.rows[7]
assert row.cells[0].text == '7', 'not row 7'
old = row.cells[2].text
assert old.startswith('A single trained model (B1 × Neo-Hookean'), (
    f'unexpected stale text: {old!r}')
assert 'in progress' in old

new = (
    'Every trained model (all six geometry × material combinations) '
    'evaluated zero-shot, with no retraining, on seven unseen resolutions '
    '(N=13, 17, 25, 29, 37, 41, 49) against a common N=101 fine-mesh '
    'reference. All three B1 materials stay within 5.0–10.6% error across '
    'every resolution tested. B2 shows the same real resolution '
    'invariance, weaker than B1: each material’s spread across the '
    'seven meshes (2.11× for B1’s worst case, up to 4.91× for B2 × '
    'Mooney-Rivlin) exceeds anything seen on B1. All six cases now have a '
    'valid result; none remain. (§8.7)'
)
row.cells[2].text = new
assert row.cells[2].text == new

doc.save(DST)
print('wrote', DST)
