"""Mirrors report v47 into the summary: renumbers the four
mesh-convergence captions (B1xMR, B1xAB, B2xMR, B2xAB) from 3/4/5/6 to
1a/1b/2a/2b, matching the report's fix and keeping both documents'
numbering identical for the same underlying data. Fixes the duplicate
Table 5 and Table 6 found by check_report_tables.py.

The summary never had separate "hyperparameters" or "training protocol"
tables, so 3/4 were not technically ambiguous here the way they are in
the report -- renumbered anyway so the two documents agree on what
"Table 1a" etc. means.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v21.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import shutil

from docx import Document

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v21.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v21.docx')
ORIGINAL = list(doc.paragraphs)


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


retext('Table 3. B1 x Mooney-Rivlin mesh convergence',
       'Table 1a. B1 x Mooney-Rivlin mesh convergence (N=6-51)')
retext('Table 4. B1 x Arruda-Boyce mesh convergence',
       'Table 1b. B1 x Arruda-Boyce mesh convergence (N=6-51)')
retext('Table 5. B2 x Mooney-Rivlin mesh convergence',
       'Table 2a. B2 x Mooney-Rivlin mesh convergence (N=6-51)')
retext('Table 6. B2 x Arruda-Boyce mesh convergence',
       'Table 2b. B2 x Arruda-Boyce mesh convergence (N=6-51)')

doc.save(SRC)
print('summary v21: mesh-convergence captions renumbered to 1a/1b/2a/2b')
