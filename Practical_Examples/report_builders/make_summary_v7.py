"""Adds point 7b (report Table 21) and the MMS (Tables 22-23) to the summary.

Reads the same JSONs as make_v34.py and re-checks the same claims, so the two
documents cannot disagree.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
P7 = json.load(open(os.path.join(HERE, '..', 'omar_pfem', 'point7b_results',
                                 'comparison_B1_neo_hookean.json')))
P9 = json.load(open(os.path.join(HERE, '..', 'omar_pfem', 'point9_results',
                                 'mms_B1_neo_hookean.json')))

R = P7['runs']
PI_A = R['physics_informed']['best_val_rel_L2']
DD_A = R['data_driven_matched_optimizer']['best_val_rel_L2']
PI_O = R['physics_informed_adamw_onecycle']['best_val_rel_L2']
DD_O = R['data_driven_own_optimizer']['best_val_rel_L2']
assert all(R[k]['opt_steps'] == 75000 for k in R)
assert (PI_A < DD_A) != (PI_O < DD_O), 'the ranking no longer flips'
adam_gap = abs(PI_A - DD_A) / max(PI_A, DD_A) * 100
oc_gap = abs(PI_O - DD_O) / max(PI_O, DD_O) * 100
pi_worse = (PI_O / PI_A - 1) * 100
dd_better = (1 - DD_O / DD_A) * 100
LABEL_H = R['data_driven_matched_optimizer']['label_generation_cost_h']

ROWS9 = sorted(P9['rows'], key=lambda r: (r['order'], r['N']))
RATES = P9['convergence_rates']
assert all(v == 'as expected' for v in P9['rate_check'].values())
byd = {}
for r in ROWS9:
    byd.setdefault(r['n_dof'], {})[r['order']] = r
SHARED = sorted(d for d, v in byd.items() if len(v) > 1)
ADV = [(d, byd[d]['Q4']['L2_rel'] / byd[d]['Q9']['L2_rel']) for d in SHARED]
assert ADV[-1][1] > ADV[0][1]

doc = Document('PFEM_Summary_Completed_Work.pre_v7.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def heading(t):
    p = doc.add_paragraph(style=HS)
    p.add_run(t)
    return p


HS = None
for p in doc.paragraphs:
    if p.text.strip().startswith('9. Accuracy/cost Pareto'):
        HS = p.style
        break
assert HS is not None, 'summary heading style not found'

els = []
els.append(heading('10. Physics-informed vs data-driven, the complete 2×2')._p)
els.append(para(
    'Both losses under both optimisers, all four at 75,000 optimiser steps, '
    'B1 × Neo-Hookean, 800/200 split. Held-out relative L2.')._p)
els.append(new_table(
    ['Training loss', 'Adam lr 2e-3', 'AdamW lr 1e-3 + OneCycle'],
    [['Physics-informed (energy)', f'{PI_A:.4f}', f'{PI_O:.4f}'],
     ['Data-driven (rel. L2 to FEM)', f'{DD_A:.4f}', f'{DD_O:.4f}']])._tbl)
els.append(para(
    f'Table 21. THE RANKING FLIPS. Physics-informed wins the Adam column by '
    f'{adam_gap:.0f}%; data-driven wins the AdamW column by {oc_gap:.0f}%. '
    'Neither principle wins outright on this problem, and only the complete '
    'grid shows it — either column alone supports a confident, opposite '
    'conclusion.')._p)
els.append(para(
    f'Why: the optimiser moves the two losses in OPPOSITE directions. '
    f'AdamW+OneCycle makes the physics-informed model {pi_worse:.0f}% worse and '
    f'the data-driven model {dd_better:.0f}% better. A schedule tuned on one '
    'loss is not neutral ground for the other.')._p)
els.append(para(
    f'What survives the flip: the data-driven model needs 800 FEM solves = '
    f'{LABEL_H:.2f} h of CPU for labels before training starts; the '
    'physics-informed one needs none. That is a property of the principle, not '
    'of the recipe, and it does not move between columns.')._p)

els.append(heading('11. Verification against an analytic solution (MMS)')._p)
els.append(para(
    'Everywhere else in this report FEM defines the reference, so FEM is never '
    'itself graded. A manufactured solution supplies an exact answer in closed '
    'form and scores Q4 and Q9 on equal terms. u* vanishes on the whole '
    'boundary, so homogeneous Dirichlet is exact and the shared solver needed '
    'no change; the body force b = −Div P is derived by nested autodiff and '
    'checked against a central difference. Uniform material, FP64.')._p)
els.append(new_table(
    ['Order', 'N', 'DOF', 'L2', 'H1 semi', 'Stress', 'Energy'],
    [[r['order'], str(r['N']), f"{r['n_dof']:,}", f"{r['L2_rel']:.3e}",
      f"{r['H1_semi_rel']:.3e}", f"{r['stress_rel_L2']:.3e}",
      f"{r['energy_rel']:.3e}"] for r in ROWS9])._tbl)
els.append(para('Table 22. Q4 and Q9 against the manufactured solution.')._p)
els.append(new_table(
    ['Order', 'Norm', 'Observed', 'Theory', 'Pairwise'],
    [[o, lab, f"{RATES[o][n]['rate']:.2f}", str(RATES[o][n]['expected']),
      ', '.join(f'{p:.2f}' for p in RATES[o][n]['pairwise'])]
     for o in ('Q4', 'Q9')
     for n, lab in (('L2', 'L2'), ('H1_semi', 'H1 semi'),
                    ('stress', 'Stress'), ('energy', 'Energy'))])._tbl)
els.append(para(
    'Table 23. This is the verification, not decoration: a body force wrong by '
    'a sign or a factor would collapse these rates. All eight land on theory, '
    'and every pairwise rate is within 0.03 of its fit — unlike Table 20\'s '
    'exponent, these have settled.')._p)
els.append(para(
    'Q9 wins at equal cost and its margin grows with refinement: '
    + ', '.join(f'{r:.1f}× at {d:,} DOF' for d, r in ADV) + ' in L2.')._p)
els.append(para(
    'NOT here: the operator third of the three-way Timon asked for. It cannot '
    'run on a manufactured problem as the pipeline stands — no body-force term '
    'in Π, no body-force input channel — so the Table 5 checkpoints do not '
    'apply and a separate model must be trained. Also: the body-force-vs-'
    'homogeneous fork was decided HERE, not confirmed by Timon; a body-force-'
    'free exact solution on this domain is a homogeneous deformation, which Q4 '
    'reproduces to machine precision and which would separate nothing.')._p)

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('9. Accuracy/cost Pareto'):
        anchor = p._p
        break
cur = last_el = anchor
while True:
    nxt = cur.getnext()
    if nxt is None:
        break
    if nxt.tag == qn('w:p'):
        pp = Paragraph(nxt, doc)
        if pp.style is not None and pp.style == HS and pp.text.strip():
            break
    last_el = nxt
    cur = nxt
target = last_el
for el in els:
    target.addnext(el)
    target = el

doc.save('PFEM_Summary_Completed_Work.docx')
print(f'inserted {len(els)} elements as summary sections 10 and 11')
