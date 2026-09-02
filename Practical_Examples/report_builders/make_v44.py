"""v43 -> v44. Fixes a direct internal contradiction found while auditing
the whole document: Section 4.4 (paragraph "Scope of this section") says
the B2 Q4-vs-Q9 study is "deliberately confined to the B1 geometry... not
planned", but Section 10's opening paragraph still called that same study
"open for B2" and "one genuinely unfinished measurement" -- the two
sections flatly disagreed about whether this item exists.

This paragraph predates the whole make_vNN.py builder chain (grep across
every make_v27.py..v43.py finds no retext() call that ever touched it) --
it is a leftover from the original base document that no automated pass
ever addressed. No JSON to assert against here; the correction is reading
Section 4.4's own already-correct sentence back and making Section 10
agree with it, plus recounting the bulleted "remaining items" list below
it, which now holds exactly one open item (OOD attribution), not two.
"""
from docx import Document

SRC, DST = 'PFEM_Transolver_Report_v43.docx', 'PFEM_Transolver_Report_v44.docx'

doc = Document(SRC)
ORIGINAL = list(doc.paragraphs)


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


# Section 4.4's own scope note (unchanged, read back to confirm what
# Section 10 must agree with):
scope_note = find_para('Scope of this section')
assert 'not planned' in scope_note.text and 'deliberately confined' in \
    scope_note.text, 'Section 4.4 no longer says what this fix assumes it says'

retext(
    'The h-refinement mesh-convergence study against fixed analytical '
    'fields (Section 4.3) and the GPU-native FEM timing benchmark',
    'The h-refinement mesh-convergence study against fixed analytical '
    'fields (Section 4.3) and the GPU-native FEM timing benchmark '
    '(Section 8.5) are both complete for all six geometry–material '
    'combinations, including B2 × Neo-Hookean, with convergence confirmed '
    'through (Table 2), demonstrating the same rapid-convergence pattern '
    'observed in the other cases. This is distinct from the deeper, '
    'DOF-referenced Q4-vs-Q9 convergence study of Section 4.4, which — as '
    'that section\'s own scope note states — is deliberately confined to '
    'the B1 geometry and not planned for B2; it is a scope decision, not '
    'an unfinished measurement. The remaining item below is a single '
    'scientifically motivated extension, not a missing core measurement.')

doc.save(DST)
print(f'wrote {DST} -- Section 10 now agrees with Section 4.4 on the '
      f'B2 Q4-vs-Q9 study')
