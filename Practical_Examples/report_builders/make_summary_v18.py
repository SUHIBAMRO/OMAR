"""Mirrors report v45 into the summary: Table 18d, B2 x Mooney-Rivlin's
Pareto and the replicated anchoring finding. Same JSON and assertions as
make_v45.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v18.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2MR = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_mooney_rivlin.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
rows = {r['N']: r for r in B2MR['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = {int(k): v for k, v in B2MR['shape']['training_resolution_anchoring']['anchor_ratios'].items()}
nh_anchor = {int(k): v for k, v in B2MR['shape']['training_resolution_anchoring'][
    'neo_hookean_anchor_ratios_for_comparison'].items()}
assert anchor[21] > 2.5 and anchor[33] > 2.5

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v18.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v18.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']
table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

insert_after('Unlike every B1 curve', [
    para(
        f'Table 18d. B2 × Mooney-Rivlin (second of three B2 materials to '
        f'finish). Speed-up {min(speedup.values()):,.0f}×–'
        f'{max(speedup.values()):,.0f}×.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'The anchoring effect replicates: local minima at N=21 and N=33 '
        f'again, {anchor[21]:.2f}× and {anchor[33]:.2f}× lower than each '
        f'point\'s neighbours — smaller than Neo-Hookean\'s '
        f'{nh_anchor[21]:.2f}×/{nh_anchor[33]:.2f}× but the same shape, at '
        f'the same two resolutions. Two of three B2 materials now show it; '
        f'Arruda-Boyce, the last one, decides whether this generalises. '
        f'Arruda-Boyce\'s B2 sweep was still running as this was '
        f'written.')._p,
])

retext('Unlike every B1 curve',
       'Unlike every B1 curve, this one is not smooth in N: sharp local '
       'minima at N=21 and N=33 — the two resolutions this checkpoint was '
       'jointly trained on — 4.30× and 3.65× lower than each point\'s '
       'neighbours. Table 18\'s B1 × Neo-Hookean curve, from a checkpoint '
       'trained at N=21 only, shows no such dip (1.01×, flat) — candidate '
       'explanation (joint two-resolution training leaves visible anchor '
       'points), not established, since geometry and protocol differ at '
       'once between the two checkpoints compared.')

doc.save(SRC)
print(f'summary v18: Table 18d added -- anchor ratios {anchor[21]:.2f}x/'
      f'{anchor[33]:.2f}x')
