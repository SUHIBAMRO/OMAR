"""v51 -> v52. Omar found VINO's own README sitting unedited at the repo
root and asked why. Checked the full (previously shallow) git history:
an early prototype (`Practical_Examples/omar/`, July 3-6) vendored VINO's
code and briefly used its exact closed-form energy-integration method,
but three days later that was abandoned for a separate, "isolated from
omar/" rewrite (`Practical_Examples/omar_pfem/`, from July 9), which is
the ONLY codebase behind every result in this report. It uses ordinary
Gauss-quadrature energy assembly (already correctly described in
Section 5.2), not VINO's method.

The one real, current connection: `materials_torch.py`'s own docstring
says the Mooney-Rivlin and Arruda-Boyce strain-energy densities "match
the energy densities already validated in omar/losses.py (the VINO
project)". That is a genuine cross-check worth citing -- not a claim
that the training methodology comes from VINO, which it does not.

Adds one sentence after Section 2.4 (right where the report already
discusses how the Mooney-Rivlin/Arruda-Boyce derivatives were computed)
noting this cross-check, and adds VINO as reference [4].
"""
from docx import Document
import copy

SRC, DST = 'PFEM_Transolver_Report_v51.docx', 'PFEM_Transolver_Report_v52.docx'

doc = Document(SRC)
ORIGINAL = list(doc.paragraphs)


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; expected exactly 1')
    return hits[0]


def insert_after(prefix, text):
    from docx.text.paragraph import Paragraph
    p = find_para(prefix)
    new_p = copy.deepcopy(p._p)
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    for r in list(np.runs):
        r._r.getparent().remove(r._r)
    np.add_run(text)
    return np


anchor = (
    'is used exclusively by the reference finite-element Newton–Raphson '
    'solver. For the Neo-Hookean model'
)
insert_after(
    anchor,
    'The Mooney-Rivlin and Arruda-Boyce strain-energy densities above '
    'were cross-checked against the implementation released with VINO '
    '[4], a variational physics-informed neural operator from the same '
    'research group, and match it. An early prototype of this project '
    'briefly used VINO’s own closed-form energy-integration method before '
    'being superseded by the Gauss-quadrature assembly described in '
    'Section 5.2, which produced every result in this report; the '
    'material formulas are the one part of VINO’s implementation this '
    'work still relies on, for correctness, not for methodology.'
)

insert_after(
    '[3] E, W., and Yu, B. “The Deep Ritz Method',
    '[4] Eshaghi, M.S., Anitescu, C., Thombre, M., Wang, Y., Zhuang, X., '
    'and Rabczuk, T. “Variational Physics-informed Neural Operator '
    '(VINO) for Learning Partial Differential Equations.” Computer '
    'Methods in Applied Mechanics and Engineering, 437, 117785, 2025.'
)

doc.save(DST)
print('wrote', DST)
