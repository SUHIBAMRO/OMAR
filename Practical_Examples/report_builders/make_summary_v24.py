"""v23 -> v24. Mirrors make_v51.py's table-formatting fix into the
summary. Same root cause, different numbers: 21 of 48 tables have a
styled header (light-gray fill #E8EAED, bold, 8pt) and the other 27 have
no header shading or explicit font size at all.

Standardizes on the REPORT's header style (dark fill #1F2937, bold,
white, 9pt) rather than the summary's own lighter one, so the two
documents match each other as well as being internally consistent --
Omar has both open together, and a reader flipping between them
shouldn't see two different table languages.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC, DST = 'PFEM_Summary_Completed_Work.docx', 'PFEM_Summary_Completed_Work.docx'
BACKUP = 'PFEM_Summary_Completed_Work.pre_v24.docx'
HEADER_FILL = '1F2937'
FONT_SIZE = Pt(9)

doc = Document(SRC)
doc.save(BACKUP)


def set_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


changed = 0
for t in doc.tables:
    header = t.rows[0]
    for cell in header.cells:
        set_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                changed += 1
                r.font.bold = True
                r.font.size = FONT_SIZE
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in t.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = FONT_SIZE

assert changed > 0
print(f'{changed} header runs restyled')

doc.save(DST)
print('overwrote', DST, '(backup at', BACKUP, ')')
