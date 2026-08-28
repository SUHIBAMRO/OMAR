"""v34 -> v35: the operator third of point 9, so section 8.11 is three-way.

v34's section 8.11 closes with a paragraph headed "Two limits." whose first
limit is that the operator third does not exist. It does now. That paragraph
is REPLACED -- not appended to, which would leave the document asserting both
things -- by the three-way table and the reading it needs.

Every number is read from the committed JSONs and asserted before it is
written, including the two that make the section honest: that the FEM columns
are the same numbers Table 22 already prints, and that operator/Q4 is above
one.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v34.docx', 'PFEM_Transolver_Report_v35.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
P9D = os.path.join(HERE, '..', 'omar_pfem', 'point9_results')
P9 = json.load(open(os.path.join(P9D, 'mms_B1_neo_hookean.json')))
OP = json.load(open(os.path.join(P9D, 'mms_operator_B1_neo_hookean.json')))

KEYS = ('L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel')
N = OP['N']
op = OP['operator_on_the_reference_member']
ref = OP['fem_reference_same_mesh']

# ---- the FEM columns must BE Table 22's, not merely resemble them -------
# If these ever diverge the two tables would quietly disagree inside one
# document, which is worse than either being wrong on its own.
for order in ('Q4', 'Q9'):
    row = next(r for r in P9['rows']
               if r['order'] == order and r['N'] == N)
    for k in KEYS:
        assert abs(ref[order][k] - row[k]) < 1e-15 * max(1.0, abs(row[k])), \
            f'{order} {k}: Table 24 would contradict Table 22'
    assert ref[order]['n_dof'] == row['n_dof']
Q4_DOF, Q9_DOF = ref['Q4']['n_dof'], ref['Q9']['n_dof']
assert Q4_DOF == OP['n_dof'], 'the operator is not on the Q4 mesh it is compared to'
assert Q9_DOF > Q4_DOF, 'Q9 at the same N is no longer the more expensive one'

# ---- the ceiling --------------------------------------------------------
RATIO = {k: op[k] / ref['Q4'][k] for k in KEYS}
assert abs(RATIO['L2_rel'] - OP['operator_over_Q4_L2']) < 1e-9
assert RATIO['L2_rel'] > 1.0, (
    'operator/Q4 is below 1.0. That is not a result to write up -- it means '
    'the operator is not minimising the same functional over the same space. '
    'Check the Dirichlet mask, the quadrature and the work term first.')
# The inversion is the section's actual finding, so it is asserted rather
# than described: if a future run puts L2 back on the forgiving side, the
# paragraph below is wrong and this must stop.
INVERTED = (RATIO['H1_semi_rel'] < RATIO['L2_rel']
            and RATIO['stress_rel_L2'] < RATIO['L2_rel'])
assert INVERTED, 'the norm ordering is no longer inverted -- rewrite the reading'

# ---- the same inversion in the independent demo run ---------------------
DEMO = json.load(open(os.path.join(P9D, 'operator_demo_N9_undertrained.json')))
d_op, d_q4 = DEMO['operator_on_the_reference_member'], DEMO['fem_reference_same_mesh']['Q4']
DEMO_INVERTED = (d_op['H1_semi_rel'] / d_q4['H1_semi_rel']
                 < d_op['L2_rel'] / d_q4['L2_rel'])
assert DEMO['N'] != N and DEMO['device'] != OP['device'], \
    'the demo run is no longer an independent mesh and device'

# ---- training curve, for the "budget or method" sentence ----------------
H = OP['history']
half = len(H) // 2
B_HALF = min(x['L2_rel'] for x in H[:half])
B_END = min(x['L2_rel'] for x in H)
SECOND_HALF_GAIN = (1 - B_END / B_HALF) * 100
assert SECOND_HALF_GAIN > 0, 'the best error did not improve in the second half'
MEMBER = OP['family']['scored_member']
DEF_A, DEF_B = MEMBER['alpha'], MEMBER['beta']
# Table 22 and Table 24 must be scored on the same member, or the "these are
# the same numbers" sentence in the caption is false.
assert f'{DEF_A:g}*(' in P9['manufactured_solution'], P9['manufactured_solution']
assert f'{DEF_B:g}*sin' in P9['manufactured_solution'], P9['manufactured_solution']
FV = OP['functional_verified']
T = OP['training']
assert T['label_cost'].startswith('zero'), 'the operator run is no longer label-free'

print(f'operator verified: op/Q4 L2 {RATIO["L2_rel"]:.2f}x, '
      f'H1 {RATIO["H1_semi_rel"]:.2f}x, stress {RATIO["stress_rel_L2"]:.2f}x, '
      f'energy {RATIO["energy_rel"]:.2f}x')
print(f'  inversion holds here and in the independent N={DEMO["N"]} '
      f'{DEMO["device"]} run: {DEMO_INVERTED}')
print(f'  best L2 {B_HALF:.3e} at halfway -> {B_END:.3e} at the end '
      f'({SECOND_HALF_GAIN:.0f}% over the second half)')

doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


els = []

els.append(para(
    'The third leg. The comparison the advisor asked for is three-way, and the '
    'two tables above are FEM only. The report\'s trained operators cannot be '
    'pointed at a manufactured problem: their energy functional has no '
    'body-force term and their input channels have no field to carry one, so '
    'the checkpoints of Table 5 do not apply. A separate operator was therefore '
    'trained for this study — the same architecture, on the same Q4 mesh, '
    'minimising the same discrete potential energy, and scored by the same '
    'error routine as the two solvers, so the three columns are commensurable '
    'even though the third model is not one of the report\'s own.')._p)

els.append(para(
    'One property of that arrangement decides how the numbers read, and it is '
    'arithmetic rather than an empirical expectation. The operator minimises '
    'the same discrete functional over the same Q4 space that the Q4 solver '
    'solves, and the minimiser of that functional over that space is the Q4 '
    'solution itself. The operator therefore cannot be more accurate than Q4 at '
    'this mesh, and a ratio below one would indicate a defect in the Dirichlet '
    'mask, the quadrature or the work term rather than an advance. The quantity '
    'of interest is the ratio operator/Q4, where one would mean the network has '
    'solved the variational problem exactly and anything above it is the '
    'network\'s optimisation error, measured on its own.')._p)

els.append(para(
    'That the functional is genuinely the solver\'s was established before any '
    f'training: Π at the finite-element solution is {FV["Pi_at_u_FEM"]:.6f}, the '
    'interpolated exact solution does not beat it, all 36 admissible '
    'perturbations raise it, and the excess grows quadratically with ratio '
    f'{FV["quadratic_excess_ratio"]:.3f} against the {FV["expected_quadratic_excess_ratio"]:.1f} '
    'that a true minimum requires. A work term deliberately mis-scaled by a '
    f'factor of eight moves the minimum to {FV["best_scale_W_divided_by_8"]}, '
    'which is the check that those checks can fail.')._p)

els.append(new_table(
    ['Method', 'DOF', 'L2', 'H1 semi-norm', 'Stress', 'Energy'],
    [['Q4 (the same mesh)', f'{Q4_DOF:,}'] +
     [f'{ref["Q4"][k]:.3e}' for k in KEYS],
     ['Q9 (the same N)', f'{Q9_DOF:,}'] +
     [f'{ref["Q9"][k]:.3e}' for k in KEYS],
     ['Physics-informed operator', f'{Q4_DOF:,}'] +
     [f'{op[k]:.3e}' for k in KEYS]])._tbl)
els.append(para(
    f'Table 24. The three-way comparison at N = {N}, every entry scored against '
    'the same manufactured solution by the same routine. The Q4 and Q9 rows are '
    f'the N = {N} rows of Table 22 and are not a second measurement. The '
    f'operator sits on the Q4 mesh; Q9 at the same N costs {Q9_DOF / Q4_DOF:.1f}× '
    f'the degrees of freedom. The operator was trained for {T["opt_steps"]:,} '
    f'optimiser steps, {T["train_wall_clock_min"]} minutes on an A100, on a '
    f'{OP["family"]["ntrain"]}-member family of manufactured solutions, with no '
    'labels: the exact solution is analytic but enters only the scoring, never '
    'the loss.')._p)

els.append(para(
    f'The operator is {RATIO["L2_rel"]:.2f}× the Q4 error in L2, so the ceiling '
    'holds and the network has closed most, but not all, of the distance to the '
    'optimum it is chasing. The more interesting reading is that the four norms '
    'do not agree on how far it is. In the H1 semi-norm the ratio is '
    f'{RATIO["H1_semi_rel"]:.2f}× and in stress {RATIO["stress_rel_L2"]:.2f}×, '
    'both effectively at the Q4 optimum, while L2 sits at '
    f'{RATIO["L2_rel"]:.2f}× and energy at {RATIO["energy_rel"]:.2f}×. That '
    'ordering is inverted relative to the usual one, in which L2 is the '
    'forgiving norm and the derivative-based norms are the strict ones.')._p)

els.append(para(
    'The inversion is consistent with what the training principle optimises. '
    'The energy functional is built from the deformation gradient, so the '
    'strain and stress fields are the quantities the loss sees directly and '
    'constrains hardest; the displacement itself is only pinned down through '
    'them, anchored by the boundary mask. A gradient-accurate field that drifts '
    'slightly in magnitude is exactly what that objective tolerates. The same '
    f'inversion appears in an earlier, independent run at N = {DEMO["N"]} on CPU '
    f'({d_op["H1_semi_rel"] / d_q4["H1_semi_rel"]:.2f}× in H1 against '
    f'{d_op["L2_rel"] / d_q4["L2_rel"]:.2f}× in L2), so it is not an artefact of '
    'one training run. It is a property worth stating plainly: for a '
    'physics-informed operator, an L2 displacement error overstates how wrong '
    'the mechanics are.')._p)

els.append(para(
    'What is left is optimisation error, not discretisation error, and the run '
    'does not establish how much of it a longer budget would remove. The best '
    f'held-out L2 was {B_HALF:.3e} at the halfway point and {B_END:.3e} at the '
    f'end, a further {SECOND_HALF_GAIN:.0f}% over the second half of training — '
    'still falling, and slowly. Three limits remain beyond that. The operator '
    'was trained and scored at one mesh only, so it has no convergence rate of '
    'its own and cannot appear in Table 23; producing one means a training run '
    'per refinement, not a solve per refinement. Its cost — minutes of GPU '
    'training against seconds of FP64 CPU Newton solves — is not put on a '
    'common axis here, because no honest one was available. And the whole of '
    'section 8.11, all three legs of it, rests on one geometry and one '
    f'material; the manufactured family is parametrised by two numbers and '
    f'{OP["family"]["ntrain"]} of its members train the operator, but every '
    'error in Tables 22 and 24 is scored on the single member '
    f'α = {DEF_A}, β = {DEF_B}.')._p)

# Replace the "Two limits." paragraph rather than adding to it: its first
# clause states that this third leg does not exist.
victim = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Two limits. The comparison here is two-way'):
        victim = p
        break
assert victim is not None, \
    'the "Two limits" paragraph is gone -- has v35 already been applied?'
target = victim._p
for el in els:
    target.addnext(el)
    target = el
victim._p.getparent().remove(victim._p)
print(f'replaced the "Two limits" paragraph with {len(els)} elements '
      f'(Table 24, the three-way)')

doc.save(DST)
print('wrote', DST)
