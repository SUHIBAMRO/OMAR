"""Fixes a phantom table citation found while auditing whether all nine of
Timon's points are actually present in the summary (not just the report).

The "Remaining work" paragraph (introduced in make_summary_v17.py, when
fixing a different stale sentence) cited "Table 12c closed the last of
them" for the B2 zero-shot results. But the summary never labels that
table "Table 12c" -- unlike the report, which does number it that way,
the summary introduces the B2 zero-shot table with plain prose ("Every
B1/B2 x material combination now has a valid zero-shot result:") and no
table number at all. A reader searching this document for "Table 12c"
would find nothing. The content itself was never missing -- only the
citation was wrong.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v20.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import shutil

from docx import Document

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v20.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v20.docx')
ORIGINAL = list(doc.paragraphs)


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


# This document never labels the B2 zero-shot table "Table 12c" -- confirm
# that before overwriting the citation, so this fix cannot introduce a new
# wrong claim in the other direction.
hits_12c = [p for p in ORIGINAL if p.text.strip().startswith('Table 12c')]
assert len(hits_12c) == 0, (
    'expected no "Table 12c" caption in the summary; found one -- '
    're-check this fix, the citation may now be correct')

retext('Remaining work.',
       'Remaining work. Zero-shot resolution invariance is reported for '
       'all six geometry–material cases — every B1/B2 × material '
       'combination has a valid result, above; nothing remains open '
       'there. The B2-geometry high-DOF, element-order study (Section '
       '4.4\'s deeper Q4-vs-Q9 check) was cancelled at the sponsor\'s '
       'request and is not planned, not merely postponed. Point 2\'s '
       'accuracy/cost Pareto is now measured for all six geometry × '
       'material combinations (Table 18e closed the last of them). The '
       'one genuine open item is the out-of-distribution attribution '
       'extension named in the report\'s Section 10. It does not affect '
       'the results reported above.')

doc.save(SRC)
print('summary v20: phantom "Table 12c" citation in the closing paragraph '
      'corrected')
