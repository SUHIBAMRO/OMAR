"""Mirrors report v42 into the summary: Table 18a (Mooney-Rivlin) and 18b
(Arruda-Boyce), completing the B1 Pareto sweep in the summary too. Same
JSONs and assertions as make_v42.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v15.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

MR = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B1_mooney_rivlin.json')))
AB = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B1_arruda_boyce.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]


def other_row(data, N):
    r = next(r for r in data['rows'] if r['N'] == N)
    return r['n_nodes'], r['fem_rel_L2'], r['fem_ms_per_sample'], \
        r['operator_rel_L2'], r['operator_ms_per_sample']


MR_pc = {N: other_row(MR, N) for N in NS}
AB_pc = {N: other_row(AB, N) for N in NS}
MR_speedup = {N: MR_pc[N][2] / MR_pc[N][4] for N in NS}
AB_speedup = {N: AB_pc[N][2] / AB_pc[N][4] for N in NS}
MR_op_min_N = min(NS, key=lambda N: MR_pc[N][3])
AB_op_min_N = min(NS, key=lambda N: AB_pc[N][3])
assert MR_op_min_N == 49 and AB_op_min_N == 37

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v15.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v15.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']


def table_rows(data):
    return [[str(N), f'{data[N][0]:,}', f'{data[N][1]*100:.3f}',
             f'{data[N][2]/1000:.1f}', f'{data[N][3]*100:.2f}',
             f'{data[N][4]:.3f}',
             f'{data[N][2]/data[N][4]:,.0f}×'] for N in NS]


retext('9. Accuracy/cost Pareto: operator vs. FEM (B1 × Neo-Hookean)',
       '9. Accuracy/cost Pareto: operator vs. FEM, all three B1 materials')

insert_after('Table 18. Accuracy and cost at nine resolutions', [
    para(
        f'Table 18a. Mooney-Rivlin: operator error falls monotonically to '
        f'N=49 ({MR_pc[49][3]*100:.2f}%), no minimum inside the range — '
        f'unlike Neo-Hookean. Speed-up '
        f'{min(MR_speedup.values()):,.0f}×–{max(MR_speedup.values()):,.0f}×, '
        f'~2.2× larger than Neo-Hookean\'s throughout (2.1–2.4× costlier '
        f'CPU assembly, Table 4a).')._p,
    new_table(HEADER, table_rows(MR_pc))._tbl,
    para(
        f'Table 18b. Arruda-Boyce: bottoms at N={AB_op_min_N} '
        f'({AB_pc[AB_op_min_N][3]*100:.2f}%) then rises to N=49 '
        f'({AB_pc[49][3]*100:.2f}%) — same shape as Neo-Hookean; '
        f'Mooney-Rivlin is the exception among the three. Speed-up '
        f'{min(AB_speedup.values()):,.0f}×–{max(AB_speedup.values()):,.0f}×, '
        f'between the other two. No material, at any mesh, has the '
        f'operator matching even the cheapest FEM solve on accuracy.')._p,
    new_table(HEADER, table_rows(AB_pc))._tbl,
])

doc.save(SRC)
print(f'summary v15: Table 18a/18b added -- MR speedup '
      f'{min(MR_speedup.values()):,.0f}x-{max(MR_speedup.values()):,.0f}x, '
      f'AB speedup {min(AB_speedup.values()):,.0f}x-{max(AB_speedup.values()):,.0f}x')
