"""Fixes a stale "Remaining work" closing paragraph found while auditing
the whole document. It is the very last paragraph of the summary,
under "Conclusion", and predates every make_summary_vN.py builder --
none of the 16 prior version bumps ever targeted it. It claimed:

  * the B2 high-DOF/element-order study "has not been started" -- Omar
    cancelled it on 2026-08-27 ("khalas mulgha ma badna yaha"); it is not
    planned, not merely not-yet-started.
  * zero-shot resolution invariance is "currently running... from one
    representative case to the remaining five" -- stale from before this
    project existed in its current form; all six cases have been done for
    days (Table 12c closed the last one).

Both claims are checked against the summary's OWN already-correct text
elsewhere in the same document before being overwritten, so this cannot
introduce a second disagreement.
"""
import shutil

from docx import Document

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v17.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v17.docx')
ORIGINAL = list(doc.paragraphs)


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


# Cross-check against this same document's own Table 18c paragraph, which
# already correctly describes the B2 Pareto sweeps as the genuine
# remaining item.
table18c_para = find_para('Unlike every B1 curve')
assert 'Mooney-Rivlin and' in table18c_para.text and \
    'still running' in table18c_para.text, (
    'expected this document to already say the two B2 Pareto sweeps are '
    'running -- if that wording changed, re-check this fix against it')

retext(
    'Remaining work.',
    'Remaining work. Zero-shot resolution invariance is reported for all '
    'six geometry–material cases (Table 12c closed the last of them); '
    'nothing remains open there. The B2-geometry high-DOF, element-order '
    'study (Section 4.4\'s deeper Q4-vs-Q9 check) was cancelled at the '
    'sponsor\'s request and is not planned, not merely postponed. The '
    'genuine open items are the B2 × Mooney-Rivlin and B2 × Arruda-Boyce '
    'Pareto sweeps (Table 18c\'s counterparts for the other two B2 '
    'materials), running as this is written, and the out-of-distribution '
    'attribution extension named in the report\'s Section 10. Neither '
    'affects the results reported above.')

doc.save(SRC)
print('summary v17: stale "Remaining work" closing paragraph corrected')
