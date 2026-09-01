"""Mirrors report v39 into the summary: Table 24d, the three norms (H1 semi,
stress, energy) on the 16-member MMS family that v11 left out, reporting L2
only. Same JSON, same assertions as make_v39.py, so the two documents cannot
disagree.

Run from the directory holding PFEM_Summary_Completed_Work.docx; it copies
the current file to .pre_v12.docx first, then writes the new content back to
PFEM_Summary_Completed_Work.docx, matching how every earlier summary builder
in this chain has worked.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

FAM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_family_fem_B1_neo_hookean.json')))

FROWS = {r['N']: r for r in FAM['rows']}
FNS = sorted(FROWS)
RATE = FAM['observed_rates_N9_to_N33_two_point']
PERI = FAM['per_interval_rates']
METRICS = ['L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel']
assert FNS == [9, 17, 33], FNS

assert abs(RATE['Q4']['L2_rel'] - 1.98) < 0.05, RATE['Q4']['L2_rel']
assert abs(RATE['Q4']['H1_semi_rel'] - 1.00) < 0.05, RATE['Q4']['H1_semi_rel']

RATIO = {N: {m: FROWS[N]['operator_over_Q4_on_the_family'][m] for m in METRICS}
         for N in FNS}
assert abs(RATIO[9]['L2_rel'] - 0.62) < 0.01, RATIO[9]['L2_rel']
assert abs(RATIO[17]['L2_rel'] - 2.59) < 0.01, RATIO[17]['L2_rel']
assert abs(RATIO[33]['L2_rel'] - 14.49) < 0.01, RATIO[33]['L2_rel']
for m in ('H1_semi_rel', 'stress_rel_L2'):
    assert RATIO[33][m] < 1.5, (m, RATIO[33][m])
for m in ('L2_rel', 'energy_rel'):
    assert RATIO[33][m] > 8, (m, RATIO[33][m])

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v12.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v12.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


insert_after('Limits: the operator\'s GPU training cost', [
    para(
        'That comparison was L2 only. The same family sweep also scores H1 '
        'semi-norm, stress and energy, and they do not tell the same story:')._p,
    new_table(
        ['N', 'L2', 'H1 semi', 'stress', 'energy'],
        [[str(N)] + [f'{RATIO[N][m]:.2f}×' for m in METRICS]
         for N in FNS])._tbl,
    para(
        'H1 and stress stay near the ceiling through the whole refinement — '
        f'{RATIO[9]["H1_semi_rel"]:.2f}× to {RATIO[33]["H1_semi_rel"]:.2f}× in '
        f'H1, {RATIO[9]["stress_rel_L2"]:.2f}× to '
        f'{RATIO[33]["stress_rel_L2"]:.2f}× in stress — while L2 reaches '
        f'{RATIO[33]["L2_rel"]:.2f}×. Fitted rates over N = 9 to 33: the '
        f'operator is {RATE["operator"]["H1_semi_rel"]:.2f} in H1 and '
        f'{RATE["operator"]["stress_rel_L2"]:.2f} in stress — still improving '
        f'with refinement, against Q4\'s 1.00 — but {RATE["operator"]["L2_rel"]:+.2f} '
        f'in L2. Energy is between the two, {RATE["operator"]["energy_rel"]:+.2f}. '
        'Per interval the direction holds throughout: '
        f'{PERI["operator"]["N9_to_N17"]["H1_semi_rel"]:+.2f} then '
        f'{PERI["operator"]["N17_to_N33"]["H1_semi_rel"]:+.2f} in H1, '
        f'{PERI["operator"]["N9_to_N17"]["stress_rel_L2"]:+.2f} then '
        f'{PERI["operator"]["N17_to_N33"]["stress_rel_L2"]:+.2f} in stress — '
        'the same inversion the single-member run showed (H1 and stress '
        'protected near the variational ceiling, L2 not), now confirmed on '
        '16 members at both refinement intervals.')._p,
])

doc.save(SRC)
print('summary v12: Table 24d added -- H1/stress protected near ceiling '
      f'({RATIO[33]["H1_semi_rel"]:.2f}x/{RATIO[33]["stress_rel_L2"]:.2f}x at '
      f'N=33), L2/energy diverge ({RATIO[33]["L2_rel"]:.2f}x/'
      f'{RATIO[33]["energy_rel"]:.2f}x)')
