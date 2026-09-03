"""v44 -> v45. B2 x Mooney-Rivlin's Pareto sweep finished (9/9). Adds
Table 18d, mirroring Table 18c's layout for Neo-Hookean, and updates the
running "one of three B2 materials so far" line to "two of three".

The finding: the training-resolution anchoring effect (local minima at
N=21/N=33, the two jointly-trained resolutions) REPLICATES here, at a
smaller magnitude than Neo-Hookean's (3.07x/3.17x against 4.30x/3.65x).
Two of three B2 materials now show it. Every number asserted against the
committed JSON before writing.
"""
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v44.docx', 'PFEM_Transolver_Report_v45.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2MR = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_mooney_rivlin.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
assert [r['N'] for r in B2MR['rows']] == NS

rows = {r['N']: r for r in B2MR['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = B2MR['shape']['training_resolution_anchoring']['anchor_ratios']
anchor = {int(k): v for k, v in anchor.items()}
nh_anchor = B2MR['shape']['training_resolution_anchoring'][
    'neo_hookean_anchor_ratios_for_comparison']
nh_anchor = {int(k): v for k, v in nh_anchor.items()}
assert anchor[21] > 2.5 and anchor[33] > 2.5

print(f'B2xMR Pareto: speedup {min(speedup.values()):,.0f}x-{max(speedup.values()):,.0f}x, '
      f'anchor ratios {anchor[21]:.2f}x / {anchor[33]:.2f}x '
      f'(NH: {nh_anchor[21]:.2f}x / {nh_anchor[33]:.2f}x)')

# ======================================================================
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, __import__('copy').deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']

table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

insert_after('Speed-up spans', [
    para(
        'B2 × Mooney-Rivlin\'s Pareto sweep finished next. Table 18d.')._p,
    para('Table 18d. Accuracy and cost of both methods, B2 × Mooney-Rivlin.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'The anchoring effect replicates: local minima at N=21 and N=33 '
        f'again, {anchor[21]:.2f}× and {anchor[33]:.2f}× lower than each '
        f'point\'s neighbours — smaller than Neo-Hookean\'s '
        f'{nh_anchor[21]:.2f}×/{nh_anchor[33]:.2f}× but the same shape, at '
        f'the same two resolutions. Two of the three B2 materials now show '
        f'it; Arruda-Boyce, the last one, decides whether this is a general '
        f'property of B2\'s two-resolution training protocol or '
        f'coincidental to these two materials specifically.')._p,
    para(
        f'Speed-up spans {min(speedup.values()):,.0f}× to '
        f'{max(speedup.values()):,.0f}×, the same order-of-magnitude band '
        'as every other case measured so far. Arruda-Boyce\'s B2 Pareto '
        'sweep had not finished as this was written.')._p,
])

retext('This comparison has been run for all three B1 materials',
       'This comparison has been run for all three B1 materials (Tables '
       '18, 18a, 18b) and for two of three B2 materials so far '
       '(Neo-Hookean, Table 18c; Mooney-Rivlin, Table 18d). '
       'Arruda-Boyce\'s B2 sweep was still running as this was written.')

doc.save(DST)
print(f'wrote {DST}')
