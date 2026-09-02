"""v42 -> v43. B2 x Neo-Hookean's Pareto sweep finished and is not yet in
the report. Adds Table 18c and retexts the closing line that said the B1
Pareto comparison "is not yet extended to B2".

The genuinely new finding: unlike every B1 Pareto curve (Tables 18/18a/18b),
this one is NOT smooth in N -- it has sharp local minima exactly at the two
resolutions (21, 33) this checkpoint was jointly trained on. B1 x
Neo-Hookean's own Pareto, from a checkpoint trained at N=21 only, shows no
such dip. Stated as a candidate explanation, not established, since the two
checkpoints differ in both geometry and training protocol at once.

Every number is read from the committed JSON and asserted before writing.
"""
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v42.docx', 'PFEM_Transolver_Report_v43.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2NH = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_neo_hookean.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
assert [r['N'] for r in B2NH['rows']] == NS

rows = {r['N']: r for r in B2NH['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = B2NH['shape']['training_resolution_anchoring']['anchor_ratios']
b1_ratio_21 = B2NH['shape']['training_resolution_anchoring'][
    'b1_neo_hookean_ratio_at_N21_for_comparison']
# anchor_ratios keys came back as strings through json.load
anchor = {int(k): v for k, v in anchor.items()}
assert anchor[21] > 3 and anchor[33] > 3
assert 0.8 < b1_ratio_21 < 1.3

print(f'B2xNH Pareto: speedup {min(speedup.values()):,.0f}x-{max(speedup.values()):,.0f}x, '
      f'anchor ratios {anchor[21]:.2f}x / {anchor[33]:.2f}x')

# ======================================================================
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, __import__('copy').deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']

table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

insert_after('This comparison has been run for all three B1 materials', [
    para(
        'B2 × Neo-Hookean\'s Pareto sweep is measured too — the first of '
        'the three B2 materials to finish. Table 18c.')._p,
    para('Table 18c. Accuracy and cost of both methods, B2 × Neo-Hookean.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'This curve is not smooth in N the way Tables 18/18a/18b are. The '
        f'operator error has sharp local minima exactly at N = 21 and '
        f'N = 33 — the two resolutions this checkpoint was jointly trained '
        f'on — {anchor[21]:.2f}× and {anchor[33]:.2f}× lower than the mean '
        f'of each point\'s immediate neighbours '
        f'({rows[21]["operator_rel_L2"]*100:.2f}% at N=21 against '
        f'{rows[17]["operator_rel_L2"]*100:.2f}% and '
        f'{rows[25]["operator_rel_L2"]*100:.2f}% either side; '
        f'{rows[33]["operator_rel_L2"]*100:.2f}% at N=33 against '
        f'{rows[29]["operator_rel_L2"]*100:.2f}% and '
        f'{rows[37]["operator_rel_L2"]*100:.2f}%). Table 18\'s own '
        f'B1 × Neo-Hookean curve, scored from a checkpoint trained at '
        f'N = 21 only, shows no such dip at N = 21 '
        f'({b1_ratio_21:.2f}×, effectively flat).')._p,
    para(
        'The candidate explanation — that joint training at two specific '
        'resolutions leaves two visible anchor points that '
        'single-resolution training does not — is offered as a candidate, '
        'not established. The two checkpoints being compared differ in '
        'geometry (B1 vs B2) and training protocol (one resolution vs two) '
        'at once, so this comparison alone cannot separate which '
        'difference causes the contrast; resolving it would need a B1 '
        'checkpoint trained jointly at N = 21 and 33, which was not run.')._p,
    para(
        f'Speed-up spans {min(speedup.values()):,.0f}× to '
        f'{max(speedup.values()):,.0f}×, in the same order-of-magnitude '
        'band as the three B1 materials. Mooney-Rivlin\'s and '
        'Arruda-Boyce\'s B2 Pareto sweeps had not finished as this was '
        'written.')._p,
])

retext('This comparison has been run for all three B1 materials',
       'This comparison has been run for all three B1 materials (Tables '
       '18, 18a, 18b) and for one of three B2 materials so far '
       '(Neo-Hookean, Table 18c). Mooney-Rivlin and Arruda-Boyce\'s B2 '
       'sweeps were still running as this was written.')

doc.save(DST)
print(f'wrote {DST}')
