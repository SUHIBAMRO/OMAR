"""Adds the GPU-FEM scaling sweep (report Table 20) to the parallel summary.

Reads the same JSON as make_v33.py and re-derives the same quantities, with
the same assertions, so the two documents cannot disagree.
"""
import copy
import json
import math
import os

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(HERE, '..', 'omar_pfem', 'point8_results',
                                'gpu_fem_scaling_B1_neo_hookean.json')))
ROWS = sorted(D['rows'], key=lambda r: r['N'])
BY_N = {r['N']: r for r in ROWS}
last = ROWS[-1]

lo_N, lo_us = min(((r['N'], r['us_per_dof']) for r in ROWS), key=lambda t: t[1])
assert lo_N == 501, f'the minimum moved to N={lo_N}'
fall = ROWS[0]['us_per_dof'] / lo_us
rise = last['us_per_dof'] / lo_us

big = [r for r in ROWS if r['N'] >= 501]
_x = [math.log(r['n_dof']) for r in big]
_y = [math.log(r['solve_s']) for r in big]
_mx, _my = sum(_x) / len(_x), sum(_y) / len(_y)
expo = (sum((a - _mx) * (b - _my) for a, b in zip(_x, _y))
        / sum((a - _mx) ** 2 for a in _x))
assert abs(expo - D['derived']['cost_exponent_vs_dof']) < 0.01

MEM_FIXED, MEM_PER_MDOF = 818.0, 607.0
pred = MEM_FIXED + MEM_PER_MDOF * last['n_dof'] / 1e6
err_pct = abs(pred - last['peak_gpu_mem_MB']) / last['peak_gpu_mem_MB'] * 100
assert err_pct < 5
mid = BY_N[701]
mid_pct = (mid['peak_gpu_mem_MB'] - (MEM_FIXED + MEM_PER_MDOF * mid['n_dof'] / 1e6)) \
    / mid['peak_gpu_mem_MB'] * 100

brk = [r['cost_breakdown_pct'] for r in ROWS if 'cost_breakdown_pct' in r]
assert len(brk) == 4 and all(b['cg'] > 99 for b in brk)
cg_lo, cg_hi = min(b['cg'] for b in brk), max(b['cg'] for b in brk)
asm = [b['residual'] + b['precond'] for b in brk]

doc = Document('PFEM_Summary_Completed_Work.pre_v6.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def fmt_time(s):
    return f'{s / 60:.1f} min' if s < 3600 else f'{s / 3600:.1f} h'


rows = []
for r in ROWS:
    b = r.get('cost_breakdown_pct')
    rows.append([str(r['N']), f"{r['n_dof']:,}", fmt_time(r['solve_s']),
                 f"{r['us_per_dof']:,.0f}",
                 f"{r['peak_gpu_mem_MB']:,}" if 'peak_gpu_mem_MB' in r else '—',
                 (f"{b['residual']:.1f}/{b['precond']:.1f}/{b['cg']:.1f}" if b else '—')])

hs = None
for p in doc.paragraphs:
    if p.text.strip().startswith('6. Out-of-distribution'):
        hs = p.style
        break
assert hs is not None

els = [
    para('Scaling study: the matrix-free Newton-CG solver (not the dense batched '
         'one above, which cannot reach these sizes — a 2M×2M FP64 tangent is '
         '~32 TB), eight resolutions on one A100-SXM4-80GB, FP64, 10 load '
         'steps.')._p,
    new_table(['N', 'DOF', 'Solve', 'µs/DOF', 'Peak GPU (MB)',
               'resid/precond/CG (%)'], rows)._tbl,
    para(f'Table 20. Largest case: {last["n_dof"]:,} DOF in '
         f'{last["solve_s"] / 3600:.1f} h using {last["peak_gpu_mem_MB"]:,} MB of '
         f'80 GB (~{last["peak_gpu_mem_MB"] / 1024 / 80 * 100:.0f}%). Memory is '
         'not the constraint.')._p,
    para(f'µs/DOF is U-shaped: falls {fall:.1f}× from '
         f'{ROWS[0]["us_per_dof"]:,.0f} at N=101 to {lo_us:,.0f} at N={lo_N}, '
         f'then rises {rise:.1f}× to {last["us_per_dof"]:,.0f} at N=1401. Below '
         'the minimum the GPU is under-occupied (launch overhead, same effect as '
         'bs=1 in Table 10a); above it the CG iteration count grows because the '
         'tangent\'s condition number scales as 1/h² and Jacobi only partly '
         f'offsets it. Cost on the large branch fits DOF^{expo:.2f}. So: O(DOF) '
         'in memory, not in time — and the solver has an operating range rather '
         'than simply degrading.')._p,
    para(f'Memory: {MEM_FIXED:,.0f} MB fixed + {MEM_PER_MDOF:,.0f} MB per '
         f'million DOF, a line through N=501 and N=1001 built before N=1401 ran. '
         f'It predicted {pred:,.0f} MB; the measured peak was '
         f'{last["peak_gpu_mem_MB"]:,} MB — {err_pct:.1f}% out of sample. Caveat: '
         f'N=701 sits {mid_pct:.0f}% above that line, so the linearity is good to '
         'about ten per cent, not better.')._p,
    para(f'Cost breakdown: explicit assembly (residual + Jacobi preconditioner) '
         f'{min(asm):.1f}–{max(asm):.1f}%, CG {cg_lo:.1f}–{cg_hi:.1f}%. This '
         'superficially confirms Timon\'s expectation that assembly should be '
         'minimal, and should NOT be quoted that way: matrix-free means every CG '
         'iteration IS a Hessian-vector product, i.e. an assembly-like pass over '
         'all elements. The assembly did not get cheap, it moved inside CG where '
         'this instrumentation cannot see it. The clean split exists only for '
         'assemble-once-and-factorise solvers — on the CPU reference (Table 4a) '
         'assembly outweighs the sparse solve by 290–692×.')._p,
]

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('5. GPU-native'):
        anchor = p._p
        break
assert anchor is not None, 'summary section 5 heading not found'
cur = last_el = anchor
while True:
    nxt = cur.getnext()
    if nxt is None:
        break
    if nxt.tag == qn('w:p'):
        pp = Paragraph(nxt, doc)
        if pp.style is not None and pp.style == hs and pp.text.strip():
            break
    last_el = nxt
    cur = nxt
target = last_el
for el in els:
    target.addnext(el)
    target = el

doc.save('PFEM_Summary_Completed_Work.docx')
print(f'inserted {len(els)} elements into summary section 5')
