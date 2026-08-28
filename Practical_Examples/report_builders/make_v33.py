"""v32 -> v33: the GPU-native FEM scaling sweep, the advisor's point 8.

"Test GPU-native FEM at finer discretizations up to a few million DOFs",
extended in round 6 with "smaller intermediate numbers and a breakdown of
the computational cost".

Goes at the end of section 8.5, which is where the GPU-native solver lives.
Table 20, continuing 15-19.

Every number in the prose is derived from the committed result JSON or read
back out of the source document, never typed in by hand, and the claims are
asserted before they are written. Two errors in the first draft of this file
were caught exactly that way: it quoted 3,215 us/DOF for N=501 where the
table prints the run's own 3,219, and it attributed the CPU assembly-versus-
solve split to a "Table 4b" that does not exist, at a factor of 74 where the
real one is 309.
"""
import copy
import json
import math
import os

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC, DST = 'PFEM_Transolver_Report_v32.docx', 'PFEM_Transolver_Report_v33.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(
    HERE, '..', 'omar_pfem', 'point8_results', 'gpu_fem_scaling_B1_neo_hookean.json')))
ROWS = sorted(D['rows'], key=lambda r: r['N'])
BY_N = {r['N']: r for r in ROWS}

doc = Document(SRC)

# ---- read Table 4a back out of the source, so the CPU assembly-vs-solve
# ---- factor quoted below is the document's own number and cannot drift ----
_items = []
for _el in doc.element.body:
    if _el.tag == qn('w:p'):
        _items.append(('p', Paragraph(_el, doc)))
    elif _el.tag == qn('w:tbl'):
        _items.append(('t', Table(_el, doc)))

t4a = None
for i, (k, v) in enumerate(_items):
    if k == 'p' and v.text.strip().startswith('Table 4a'):
        for j in range(i - 1, max(-1, i - 4), -1):
            if _items[j][0] == 't':
                t4a = _items[j][1]
                break
        break
assert t4a is not None, 'Table 4a (measured native CPU FEM cost) not found in ' + SRC

hdr = [c.text.strip() for c in t4a.rows[0].cells]
ci_case, ci_asm, ci_sol = hdr.index('Case'), hdr.index('Assembly (s)'), hdr.index('Solve (s)')
cpu = {}
for r in t4a.rows[1:]:
    c = [x.text.strip() for x in r.cells]
    cpu[c[ci_case]] = (float(c[ci_asm]), float(c[ci_sol]))
b1nh = next(k for k in cpu if k.startswith('B1') and 'Neo-Hookean' in k)
asm_b1nh, sol_b1nh = cpu[b1nh]
ratio_b1nh = asm_b1nh / sol_b1nh
ratios = [a / s for a, s in cpu.values()]
assert len(cpu) == 6, f'Table 4a has {len(cpu)} cases, expected 6'
assert ratio_b1nh > 1.0, 'assembly does not outweigh the linear solve on CPU'

# Which cases are dearer to assemble is a MATERIAL effect, not a geometry one:
# Neo-Hookean has an analytic PK1 and tangent (omar_pfem/data/materials.py),
# while Mooney-Rivlin and Arruda-Boyce go through jax.jacfwd(jax.grad(...))
# (omar_pfem/data/material_models_jax.py). An earlier draft of this paragraph
# blamed the B2 geometry; Table 4a says otherwise.
_asm = {k: a for k, (a, _) in cpu.items()}
_ad = [k for k in _asm if 'Mooney-Rivlin' in k or 'Arruda-Boyce' in k]
_an = [k for k in _asm if 'Neo-Hookean' in k]
assert len(_ad) == 4 and len(_an) == 2
ad_lo = min(_asm[a] / _asm[n] for a in _ad for n in _an if a[:2] == n[:2])
ad_hi = max(_asm[a] / _asm[n] for a in _ad for n in _an if a[:2] == n[:2])
b2_over_b1_nh = _asm[next(k for k in _an if k.startswith('B2'))] / asm_b1nh
assert ad_lo > 1.8, f'the autodiff materials are only {ad_lo:.2f}x dearer'
assert abs(b2_over_b1_nh - 1.0) < 0.15, \
    f'the geometry now matters for assembly cost ({b2_over_b1_nh:.2f}x)'

# ---- claims, checked before they are written --------------------------
us = [(r['N'], r['us_per_dof']) for r in ROWS]
lo_N, lo_us = min(us, key=lambda t: t[1])
assert us[0][1] > lo_us and us[-1][1] > lo_us, 'us/DOF is not U-shaped'
assert lo_N == 501, f'the minimum moved to N={lo_N}'
big = [r for r in ROWS if r['N'] >= 501]
assert all(big[i]['us_per_dof'] < big[i + 1]['us_per_dof'] for i in range(len(big) - 1)), \
    'the large branch is not monotonically rising'

fall = ROWS[0]['us_per_dof'] / lo_us
rise = ROWS[-1]['us_per_dof'] / lo_us
assert fall > 5.5 and rise > 3.0, f'the U-shape changed: fall {fall:.2f}, rise {rise:.2f}'

# cost exponent, refitted here rather than trusted from the JSON's derived block
_x = [math.log(r['n_dof']) for r in big]
_y = [math.log(r['solve_s']) for r in big]
_mx, _my = sum(_x) / len(_x), sum(_y) / len(_y)
expo = (sum((a - _mx) * (b - _my) for a, b in zip(_x, _y))
        / sum((a - _mx) ** 2 for a in _x))
pairwise = [math.log(big[i + 1]['solve_s'] / big[i]['solve_s'])
            / math.log(big[i + 1]['n_dof'] / big[i]['n_dof']) for i in range(len(big) - 1)]
assert abs(expo - D['derived']['cost_exponent_vs_dof']) < 0.01, \
    f"refitted exponent {expo:.3f} disagrees with the JSON's " \
    f"{D['derived']['cost_exponent_vs_dof']}"
assert pairwise[-1] == max(pairwise), 'the last interval is no longer the steepest'

# Memory: the model was built before N=1401 ran. It is a two-point line
# through N=501 and N=1001, NOT a fit to all three -- N=701 sits well above
# it, and the prose says so rather than calling it a clean fit.
MEM_FIXED, MEM_PER_MDOF = 818.0, 607.0
mem = [r for r in ROWS if 'peak_gpu_mem_MB' in r]
assert len(mem) == 4
_anchor = [r for r in mem if r['N'] in (501, 1001)]
assert len(_anchor) == 2, 'the memory model anchor points are missing'
for r in _anchor:
    fit = MEM_FIXED + MEM_PER_MDOF * r['n_dof'] / 1e6
    assert abs(fit - r['peak_gpu_mem_MB']) < 1.0, \
        f"{MEM_FIXED:.0f}+{MEM_PER_MDOF:.0f} is not the line through N={r['N']}"
_mid = BY_N[701]
mid_fit = MEM_FIXED + MEM_PER_MDOF * _mid['n_dof'] / 1e6
mid_resid = _mid['peak_gpu_mem_MB'] - mid_fit
mid_resid_pct = mid_resid / _mid['peak_gpu_mem_MB'] * 100
assert mid_resid > 0 and mid_resid_pct > 5, \
    'N=701 no longer sits above the line; rewrite the caveat'
last = ROWS[-1]
pred_1401 = MEM_FIXED + MEM_PER_MDOF * last['n_dof'] / 1e6
err_pct = abs(pred_1401 - last['peak_gpu_mem_MB']) / last['peak_gpu_mem_MB'] * 100
assert err_pct < 5, f'the memory prediction was off by {err_pct:.1f}%'
mem_40M = (MEM_FIXED + MEM_PER_MDOF * 40) / 1024

# cost breakdown
brk = [r['cost_breakdown_pct'] for r in ROWS if 'cost_breakdown_pct' in r]
assert len(brk) == 4 and all(b['cg'] > 99 for b in brk), 'CG is not dominant'
cg_lo, cg_hi = min(b['cg'] for b in brk), max(b['cg'] for b in brk)
asm = [b['residual'] + b['precond'] for b in brk]
asm_lo, asm_hi = min(asm), max(asm)

print(f'claims verified: U-shape, minimum {lo_us:,.0f} us/DOF at N={lo_N} '
      f'(falls {fall:.2f}x, rises {rise:.2f}x); exponent {expo:.2f}; '
      f'memory prediction {pred_1401:,.0f} vs {last["peak_gpu_mem_MB"]:,} MB '
      f'({err_pct:.1f}%); CG {cg_lo}-{cg_hi}%, assembly {asm_lo:.1f}-{asm_hi:.1f}%; '
      f'CPU assembly/solve for {b1nh} = {ratio_b1nh:.0f}x '
      f'(range {min(ratios):.0f}-{max(ratios):.0f}x)')

ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def fmt_time(s):
    return f'{s / 60:.1f} min' if s < 3600 else f'{s / 3600:.1f} h'


def u(N):
    return f'{BY_N[N]["us_per_dof"]:,.0f}'


rows = []
for r in ROWS:
    b = r.get('cost_breakdown_pct')
    rows.append([
        str(r['N']), f"{r['n_dof']:,}", fmt_time(r['solve_s']),
        f"{r['us_per_dof']:,.0f}",
        f"{r['peak_gpu_mem_MB']:,}" if 'peak_gpu_mem_MB' in r else '—',
        (f"{b['residual']:.1f} / {b['precond']:.1f} / {b['cg']:.1f}" if b else '—')])

els = []
h = doc.add_paragraph(style='Heading 3')
h.add_run('Scaling to a few million degrees of freedom')
els.append(h._p)

els.append(para(
    'The batched solver benchmarked above forms its tangent densely and calls '
    'torch.linalg.solve, which is the right choice at the study\'s own mesh but '
    'cannot be scaled: a dense two-million-by-two-million tangent in double '
    'precision is roughly 32 terabytes. The scaling study therefore uses the '
    'matrix-free Newton–Raphson solver with a Jacobi-preconditioned conjugate '
    'gradient inner loop — the same solver that produced the 10- and '
    '40-million-DOF references of Section 4.4. It never forms the tangent at '
    'all, only its action on a vector, so its memory requirement is O(DOF). '
    'Eight resolutions were run on a single A100-SXM4-80GB in FP64 with ten '
    f'load steps, spanning {ROWS[0]["n_dof"] / 1e6:.2f} to '
    f'{last["n_dof"] / 1e6:.2f} million degrees of freedom.')._p)

els.append(new_table(
    ['N', 'DOF', 'Solve time', 'µs/DOF', 'Peak GPU (MB)',
     'Residual / precond. / CG (%)'], rows)._tbl)
els.append(para(
    'Table 20. GPU-native matrix-free solver, B1 × Neo-Hookean, eight '
    'resolutions on one A100. µs/DOF is solve time divided by degrees of '
    'freedom, the figure that shows whether cost grows linearly in problem '
    'size. The last column is the wall-clock split between residual assembly, '
    'building the Jacobi preconditioner, and the CG solve; it was added after '
    'the four larger resolutions had already run, which is why they carry no '
    'breakdown. Peak memory was not recorded for the four smaller ones.')._p)

els.append(para(
    f'The headline is that the solver reaches {last["n_dof"]:,} degrees of '
    f'freedom on a single GPU, in {last["solve_s"] / 3600:.1f} hours, using '
    f'{last["peak_gpu_mem_MB"]:,} MB of the card\'s 80 GB — about '
    f'{last["peak_gpu_mem_MB"] / 1024 / 80 * 100:.0f}% of it. Memory is not the '
    'limiting resource at these sizes and is not close to becoming one.')._p)

els.append(para(
    'Cost per degree of freedom is not constant, and it is not monotone '
    f'either. It falls {fall:.1f}-fold from {u(101)} µs/DOF at N=101 to '
    f'{u(501)} at N=501, then rises {rise:.1f}-fold to {u(1401)} at N=1401 — a '
    'U-shape with its minimum near half a million degrees of freedom. The two '
    'branches have different causes. Below the minimum the problem is too '
    'small to occupy the GPU: each CG iteration is a small amount of '
    'arithmetic behind a fixed kernel-launch cost, so most of the time is '
    'overhead, which is the same effect the batch-size-1 measurements of '
    'Table 10a show. Above it, the number of CG iterations required grows with '
    'refinement, because the tangent\'s condition number scales with the '
    'inverse square of the element size and Jacobi preconditioning only partly '
    'offsets that. Fitting the four points from N=501 upward gives a cost '
    f'scaling of DOF^{expo:.2f}, with pairwise exponents of '
    f'{", ".join(f"{p:.2f}" for p in pairwise)} — the last and largest interval '
    'is the steepest, so the exponent has not settled. The solver is therefore '
    'O(DOF) in memory but not in time, and reporting only the large end would '
    'have made it look like a method that simply degrades, when in fact it has '
    'an operating range.')._p)

els.append(para(
    'The memory scaling deserves a stronger statement than the others, because '
    'it was tested out of sample rather than only fitted. A linear model of '
    f'{MEM_FIXED:,.0f} MB of fixed overhead plus {MEM_PER_MDOF:,.0f} MB per '
    'million degrees of freedom was constructed before the N=1401 case had '
    f'run, and it predicted {pred_1401:,.0f} MB for its '
    f'{last["n_dof"] / 1e6:.2f} million degrees of freedom. The measured peak '
    f'was {last["peak_gpu_mem_MB"]:,} MB, an out-of-sample error of '
    f'{err_pct:.1f}% — a prediction that could have failed and did not, on a '
    'case more than twice as large as anything the model had seen. Two '
    'qualifications belong with it. The line is drawn through two points, '
    'N=501 and N=1001, and it is not a fit to all three of the resolutions '
    f'that had run: N=701 lies {mid_resid:,.0f} MB, or {mid_resid_pct:.0f}%, '
    'above it, so peak memory is linear in problem size only to within about '
    'ten per cent, which is unsurprising given that the figure being measured '
    'is a peak over an allocator\'s behaviour rather than a count of stored '
    'values. A least-squares line through all three fitted points would '
    'predict the N=1401 case about equally well. And the extrapolation below '
    'is an extrapolation: applying the same model to the 40-million-DOF '
    f'reference of Section 4.4 gives roughly {mem_40M:.0f} GB, which is still '
    'within a single card, but that number is a projection ten times beyond '
    'the largest case measured here and has not been verified.')._p)

els.append(para(
    'On the cost breakdown, the measured split is unambiguous and its '
    'interpretation is not. Explicit assembly — evaluating the residual, and '
    f'building the Jacobi preconditioner — accounts for {asm_lo:.1f}% to '
    f'{asm_hi:.1f}% of the solve, and the CG loop for {cg_lo:.1f}% to '
    f'{cg_hi:.1f}%, with the CG share rising as the problem grows. Read '
    'literally that says the solver dominates and assembly is negligible. It '
    'should not be read literally. A matrix-free solver never forms the '
    'tangent, so every CG iteration is a Hessian-vector product, and a '
    'Hessian-vector product is itself a pass over every element performing '
    f'assembly-like work. The {cg_hi:.1f}% is solver time that contains the '
    'assembly by construction: the assembly has not become cheap, it has moved '
    'inside the CG loop, where this instrumentation cannot separate it. A '
    'clean assembly-versus-solve split exists only for a solver that assembles '
    'the tangent once and factorises it, which is what the batched GPU-native '
    'solver of Table 10 does and what Table 4a measures for the CPU reference '
    f'— where, at the study\'s own small mesh, assembly outweighs the sparse '
    f'linear solve by a factor of {ratio_b1nh:.0f} for this case, and by '
    f'{min(ratios):.0f} to {max(ratios):.0f} across the six. At the sizes in '
    'Table 20 the question simply does not have the clean answer it has at '
    'small scale.')._p)

els.append(para(
    'Three limits of this study should be stated. It covers one geometry and '
    'one material, B1 × Neo-Hookean. What that leaves out is mostly the '
    'material: Neo-Hookean is the one model here with an analytic first '
    'Piola–Kirchhoff stress and tangent, while Mooney-Rivlin and Arruda-Boyce '
    'obtain both by automatic differentiation, which on the CPU reference '
    f'costs {ad_lo:.1f} to {ad_hi:.1f} times as much to assemble (Table 4a). '
    'The geometry, by contrast, is nearly free — B2 × Neo-Hookean assembles '
    f'within {abs(b2_over_b1_nh - 1) * 100:.0f}% of B1 × Neo-Hookean — so the '
    'expectation, untested here, is that the other five cases shift the '
    'absolute times, chiefly through the material, without changing the shape '
    'of the curve. Second, the exponent is fitted on four points and its last '
    'interval is its steepest, so it should be read as a description of the '
    'range measured rather than an asymptotic rate. And third, the timings are '
    'hardware-specific: every row was measured on the same A100, and the GPU '
    'is recorded in the run manifest alongside them.')._p)

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('8.6 Out-of-distribution'):
        anchor = p
        break
assert anchor is not None, 'section 8.6 heading not found'
target = anchor._p.getprevious()
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements at the end of section 8.5')

doc.save(DST)
print('wrote', DST)
