"""v47 -> v48. Full numeric audit of every table against its source JSON
(requested explicitly: "check everything and fix the mistake"), covering
the 27 tables not yet checked cell-by-cell in a prior session (mesh
convergence, hyperparameters/protocol, batch-size sweep, GPU-native FEM
timing/latency/speed-up/break-even, GPU memory, solver agreement, OOD,
zero-shot, physics-informed-vs-data-driven, and the MMS family tables).

Found exactly one real error, in Table 24c (Q4/Q9/operator over the
16-member test family): the "operator/Q4, single member" ratio at N=33
is printed as 13.32x. The correct value, computed from the single
(non-family) operator and Q4 L2 errors at N=33 -- 0.011362 / 0.00085254,
both taken from mms_operator_rate_B1_neo_hookean.json, the same source
Table 24a/24b already draw from -- is 13.33x, not 13.32x. Table 24b (the
very next table, same report) already states this ratio correctly as
13.33x, and the source JSON's own precomputed `ratios_operator_over_Q4`
field says 13.33 too. Table 24c's 13.32x was a transcription slip against
its neighbor.

Every other cross-checked table (18-18e, 19, 19a, 20, 20a, 20b, 22, 23,
24, 24a, 24b, 24d, 24e, 11, 12, 12b, 12c, 21, plus the cross-table
identities Table 4a<->7's Native-FEM column, Table 5<->7's step counts
for the three B1 cases, Table 5<->11's in-distribution column, Table
5<->13's final recipe row, and Table 7's own Speed-up arithmetic) matched
their source exactly -- no other numeric error found. Tables with no
locally committed source JSON (1, 2, 1a, 1b, 2a, 2b, 6a, 3, 4, 6, 8, 9,
10, 10a-10d) predate the record_*.py/JSON convention; they were checked
for internal and cross-table arithmetic consistency, which all passed,
but could not be independently re-derived from a source file in this
session.
"""
from docx import Document

SRC, DST = 'PFEM_Transolver_Report_v47.docx', 'PFEM_Transolver_Report_v48.docx'

doc = Document(SRC)

# Table 24c is doc.tables[46]: header row + N=9,17,33. Last cell of the
# N=33 row is the "operator/Q4, single member" ratio.
t = doc.tables[46]
assert [c.text for c in t.rows[0].cells] == [
    'N', 'Q4 (family mean)', 'Q9 (family mean)', 'operator (family mean)',
    'operator/Q4, family', 'operator/Q4, single member'], 'not Table 24c'
row33 = t.rows[3]
cells = row33.cells
assert cells[0].text == '33', 'not the N=33 row'
assert cells[-1].text == '13.32×', f'expected stale 13.32×, found {cells[-1].text!r}'
cells[-1].text = '13.33×'
assert cells[-1].text == '13.33×'

doc.save(DST)
print('wrote', DST)
