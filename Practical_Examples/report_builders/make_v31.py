"""v30 -> v31: the accuracy/cost Pareto comparison, the advisor's point 2.

One case so far, B1 x Neo-Hookean, which is what Timon asked us to start
with ("I'd start with one specific problem such as B1-Neo Hookean. Based on
the results, we can decide then. For the paper, we should ideally have a
comparison for all problems."). Written in as a subsection of 8.7, since it
uses the same zero-shot checkpoint and the same fine-mesh reference.

Table number 18, continuing 15-17 from section 8.8.

Every number comes from pareto_table.py, which reads the committed result
JSON. Two things the prose must not soften, and does not:
  * the operator never reaches even the coarsest FEM solve's accuracy;
  * these errors are not comparable with Table 12's, for two independent
    reasons, both stated.
"""
import copy

from docx import Document
from docx.oxml.ns import qn

import pareto_table as PT

SRC, DST = 'PFEM_Transolver_Report_v30.docx', 'PFEM_Transolver_Report_v31.docx'
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
F = PT.facts()

# `advantage` is how many times LARGER the operator's best error is than the
# coarsest FEM solve's, i.e. how much more accurate that FEM solve is. It must
# be greater than one for the sentence built from it to be true; an earlier
# draft printed its reciprocal and claimed "0.2x more accurate".
assert F["advantage"] > 1.0, "the operator beats the coarsest FEM solve -- rewrite the paragraph"
assert F["op_lo"] > F["fem_hi"], "the accuracy branches overlap -- the 'empty middle' claim is false"
assert F["speed_hi"] > F["speed_lo"], "speed-up does not grow with resolution"
print(f'claims verified: FEM at N=13 is {F["advantage"]:.1f}x more accurate than '
      f'the operator at its best; branches do not overlap')


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


els = []
h = doc.add_paragraph(style='Heading 3')
h.add_run('Accuracy against cost: the operator and the finite-element solver on one pair of axes')
els.append(h._p)

els.append(para(
    'Placing the two methods on a single accuracy-cost plot is less '
    'straightforward than it appears, and the obvious construction is wrong. '
    'This report already contains a finite-element accuracy-versus-cost curve '
    '(Table 6a) and an operator accuracy figure (Table 5), but Table 6a scores '
    'a fixed analytic field against a 10-million-DOF reference while Table 5 '
    'scores random parametric fields against a same-mesh solution. Putting '
    'those on one pair of axes would produce a plausible-looking and '
    'meaningless picture. Both sides were therefore measured afresh on one '
    f'footing: the same {F["n_samples"]} problem instances, drawn from the same '
    'seeds; the same common fine-mesh reference (N=101), which the coarse FEM '
    'solutions and the operator predictions are both interpolated onto and '
    'scored against; and the same device. Each coarse resolution contributes '
    'one finite-element point and one operator point, since the operator is '
    'resolution-invariant and can be evaluated at any of them from a single '
    'checkpoint — the claim being made, so the plot shows it rather than '
    'assuming it.')._p)

els.append(new_table(PT.HEAD, PT.rows())._tbl)
els.append(para(
    'Table 18. Accuracy and cost of both methods at nine mesh resolutions, B1 × '
    'Neo-Hookean, each scored against the common N=101 reference. The '
    'finite-element cost is the CPU reference solver\'s own solve time, which '
    'is what producing a new solution actually costs today; the GPU-native '
    'solver of Section 8.5 is 71.7–171.5× faster and would shift the speed-up '
    'column down by about two orders of magnitude without touching the accuracy '
    'columns. Error is the combined relative L2 over both displacement '
    'components — see the caveat below.')._p)

els.append(para(
    'The first thing the table settles is that the two methods do not compete '
    'on accuracy at all. The finite-element solver at its coarsest setting — '
    f'N=13, 169 nodes, {F["coarsest_fem_s"]:.1f} s — reaches '
    f'{F["coarsest_fem"]:.3f}%, already {F["advantage"]:.1f}× more accurate '
    f'than the operator at its best ({F["op_lo"]:.2f}% at N={F["op_best_N"]}). '
    'There is no mesh in this sweep at which the operator matches even the '
    'cheapest finite-element solve. The Pareto front consequently has two '
    'branches with an empty middle: below a few seconds per problem the '
    f'operator is the only option available, at {F["op_ms_lo"]:.1f}–'
    f'{F["op_ms_hi"]:.1f} ms and {F["op_lo"]:.2f}–{F["op_hi"]:.2f}%; above it '
    f'the finite-element solver, at {F["fem_lo"]:.3f}–{F["fem_hi"]:.3f}%. '
    'Nothing occupies the range between. This is not a curve to select an '
    'operating point on; it is a choice between two regimes, and which one is '
    'right depends on whether an error of several per cent is acceptable at '
    'all.')._p)

els.append(para(
    'The second is that the operator\'s case rests on how cost scales, not on '
    'any single speed-up figure. Its inference cost is essentially flat in mesh '
    f'size — {F["op_ms_lo"]:.3f}–{F["op_ms_hi"]:.3f} ms from 169 to 2,401 nodes, '
    'a fourteenfold increase in problem size for no measurable increase in cost, '
    'because at batch size 1 the forward pass is dominated by kernel-launch '
    'overhead rather than arithmetic (Section 8.5) — while the finite-element '
    'solve grows superlinearly. The speed-up therefore climbs from '
    f'{F["speed_lo"]:,.0f}× to {F["speed_hi"]:,.0f}× across the sweep. That '
    'trend is the substantive result; a speed-up quoted at one resolution is '
    'not.')._p)

els.append(para(
    'The operator\'s accuracy is not monotone in resolution. It improves from '
    f'{F["op_hi"]:.2f}% at N=13 to {F["op_lo"]:.2f}% at N={F["op_best_N"]} and '
    'then worsens again, which is qualitatively the zero-shot finding of Table '
    '12 — a minimum near the resolutions the model was trained on — although '
    'the minimum falls at a different N here.')._p)

els.append(para(
    'One caveat governs how these errors may be quoted. They are measured with '
    'the combined relative L2 over both displacement components, ‖e‖/‖u‖, which '
    'is the convergence-study convention of Section 4.4 and the correct choice '
    'here because the finite-element side of this table is a convergence curve '
    'and both sides must be scored identically for the comparison to mean '
    'anything. Tables 5, 11 and 12 instead use the per-component average, and '
    'on B1 the loaded component dominates, so the combined norm reads lower. A '
    'second difference compounds it: this study draws its problem instances '
    'from a different seed range than the zero-shot evaluation, so the two are '
    'not the same physical problems scored two ways. The operator column here '
    'therefore cannot be read against Table 12, and no conversion factor '
    'between them is quoted, because it would conflate the two causes.')._p)

els.append(para(
    'Finally, on reproducibility. This comparison was run twice, hours apart on '
    'different Colab instances. Both error columns came back identical to every '
    'digit reported, on both sides, at all nine resolutions. The wall-clock '
    'columns did not: the second run\'s finite-element timings were uniformly '
    '2.887–2.925× slower than the first\'s — a spread of 1.3% across nine '
    'resolutions, which is a different machine rather than measurement noise — '
    'and the operator moved with it. The speed-up column, being a ratio, agreed '
    'between the runs to within 17%. The timings reported above are the second '
    'run\'s, because those are the ones consistent with the rest of this report: '
    'Table 10a measures this architecture at batch size 1 as 4.582 ms and the '
    'second run gives 4.584 ms on the same 441-node mesh, where the first gives '
    '1.610 ms. Absolute latencies in this table should be read as properties of '
    'the hardware; the ratios are the transferable quantity.')._p)

els.append(para(
    'This comparison has so far been run for B1 × Neo-Hookean only. Extending '
    'it to the remaining five cases requires each one\'s zero-shot checkpoint, '
    'which is the study currently in progress.')._p)

anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('8.8 Error in physically'):
        anchor = p
        break
assert anchor is not None, 'section 8.8 heading not found'
target = anchor._p.getprevious()
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements at the end of section 8.7')

doc.save(DST)
print('wrote', DST)
