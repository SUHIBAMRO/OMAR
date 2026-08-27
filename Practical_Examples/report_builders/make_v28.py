"""v27 -> v28: the operator-vs-GPU-FEM comparison at identical batch sizes.

Two changes, both in section 8.5.

Section 8.5's closing paragraph already gives a GPU-to-GPU speed-up of
"73-80x", but it compares the FEM solver at batch size 128 against the
operator at batch size 1. That is not like-for-like: batching is exactly
what the FEM solver needs to amortise its kernel launches, and the operator
was denied it. Measured at the same batch sizes, the operator gains about
16x (4.58-4.83 ms/sample at bs=1 down to 0.291-0.292 at bs=128), so the
matched figure is 1,215-1,297x. The sentence is corrected and the matched
tables added.

Break-even against the GPU baseline is added at the same time, and the
finding worth stating is that it is not one number: it ranges 1,133-95,038
across the six cases and four batch sizes, because the batch size the
comparison assumes matters more than anything else in it.

New tables are numbered 10a/10b/10c rather than renumbering Tables 11-14
and every cross-reference to them. The document already uses "Table 4a", so
the convention is its own.
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC, DST = 'PFEM_Transolver_Report_v27.docx', 'PFEM_Transolver_Report_v28.docx'
doc = Document(SRC)
body = doc.element.body

CASES = ['B1 × Neo-Hookean', 'B1 × Mooney-Rivlin', 'B1 × Arruda-Boyce',
         'B2 × Neo-Hookean', 'B2 × Mooney-Rivlin', 'B2 × Arruda-Boyce']

LATENCY = [['4.582', '0.573', '0.336', '0.292'],
           ['4.832', '0.607', '0.336', '0.292'],
           ['4.753', '0.599', '0.336', '0.292'],
           ['4.680', '0.585', '0.336', '0.291'],
           ['4.714', '0.594', '0.336', '0.291'],
           ['4.799', '0.602', '0.336', '0.292']]

SPEEDUP = [['360×', '833×', '1,134×', '1,215×'],
           ['424×', '812×', '1,154×', '1,232×'],
           ['531×', '878×', '1,221×', '1,297×'],
           ['356×', '819×', '1,135×', '1,217×'],
           ['426×', '830×', '1,158×', '1,235×'],
           ['544×', '879×', '1,223×', '1,297×']]

BREAKEVEN = [['1,745', '6,021', '7,543', '8,112'],
             ['1,363', '5,663', '7,179', '7,751'],
             ['1,133', '5,441', '6,956', '7,554'],
             ['19,410', '67,391', '84,627', '90,990'],
             ['17,033', '69,404', '87,884', '95,038'],
             ['9,530', '46,993', '60,490', '65,698']]

# ---------------------------------------------------------------- helpers
ref_tbl = None
for e in body.iter(qn('w:tbl')):
    ref_tbl = e
    break


def new_table(header, rows):
    """A table matching the document's existing ones, built by cloning the
    first table's properties rather than relying on a named style that may
    not carry the same borders."""
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    if ref_tbl is not None:
        pr = ref_tbl.find(qn('w:tblPr'))
        if pr is not None:
            old = t._tbl.find(qn('w:tblPr'))
            if old is not None:
                t._tbl.remove(old)
            t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        cell = t.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


# --------------------------------------------- 1. correct the 73-80x claim
target = doc.paragraphs[226]
old = ('a 73–80× GPU-to-GPU speed-up in the operator’s favor')
new = ('a 73–80× speed-up in the operator’s favor. That particular pair of '
       'numbers is not, however, a like-for-like comparison: it takes the FEM '
       'solver at batch size 128 and the operator at batch size 1. Batching is '
       'precisely what allows the FEM solver to amortise its kernel launches, '
       'and the operator is denied it. Measured at identical batch sizes '
       '(below), the operator gains roughly 16× itself and the matched '
       'speed-up is 1,215–1,297×')
found = False
for run in target.runs:
    if old in run.text:
        run.text = run.text.replace(old, new)
        found = True
        break
if not found:
    # the phrase may be split across runs; rebuild the paragraph's text
    full = target.text
    assert old in full, 'the sentence being corrected was not found'
    for r in list(target.runs)[1:]:
        r._r.getparent().remove(r._r)
    target.runs[0].text = full.replace(old, new)
print('corrected the unmatched speed-up sentence')

# ------------------------------------------------- 2. the new subsection
elements = []

h = doc.add_paragraph(style='Heading 3')
h.add_run('Operator vs. GPU-native FEM at identical batch sizes')
elements.append(h._p)

elements.append(para(
    'The GPU-FEM benchmark above is reported at batch sizes 1, 8, 32 and 128, '
    'because batching independent problem instances is the natural way to use a '
    'GPU for many small solves and is what amortises the fixed cost of a kernel '
    'launch. The trained operator was previously benchmarked at batch size 1 '
    'only — the realistic deployment case, where new problem instances arrive '
    'one at a time, and the right number for a deployment claim. Comparing the '
    'two as measured therefore flatters whichever side happens to benefit from '
    'the mismatch. The operator was accordingly re-benchmarked at exactly the '
    'batch sizes the FEM solver was measured at, under the same protocol: '
    'untimed warm-up calls first, CUDA synchronisation on both sides of the '
    'timed region, the median over 50 repeats rather than the mean, and each '
    'batch built from distinct test samples so that nothing can be cached '
    'across it.')._p)

elements.append(new_table(
    ['Case', 'bs=1', 'bs=8', 'bs=32', 'bs=128'],
    [[c] + r for c, r in zip(CASES, LATENCY)])._tbl)
elements.append(para(
    'Table 10a. Trained-operator inference latency in milliseconds per sample, '
    'at the same batch sizes as Table 10, all six cases (median of 50 timed '
    'repeats after 10 warm-up calls, same GPU).')._p)

elements.append(para(
    'Batching benefits the operator substantially, and almost all of the gain '
    'arrives immediately: a batch of 8 takes essentially the same wall-clock '
    'time as a batch of 1 (4.83 ms versus 4.79 ms per batch for B1 × '
    'Mooney-Rivlin), so the eightfold increase in throughput is free — at batch '
    'size 1 the GPU is idle for most of the call, waiting on launch overhead '
    'rather than computing. Per-sample cost then falls from 4.58–4.83 ms to '
    '0.291–0.292 ms at batch size 128, a factor of about 16, and is essentially '
    'identical across all six cases, as expected: every case uses the same '
    'architecture on the same 441-node mesh, and the forward pass does not '
    'depend on the material model.')._p)

elements.append(new_table(
    ['Case', 'bs=1', 'bs=8', 'bs=32', 'bs=128'],
    [[c] + r for c, r in zip(CASES, SPEEDUP)])._tbl)
elements.append(para(
    'Table 10b. Speed-up of the trained operator over the GPU-native FEM solver '
    'with both measured at the same batch size, obtained by dividing Table 10 by '
    'Table 10a row by row. This is a genuinely hardware-matched comparison: both '
    'sides run on the same GPU at the same batch size.')._p)

elements.append(para(
    'At matched batch sizes the operator is 1,215–1,297× faster than the '
    'GPU-native solver, not the 73–80× obtained when the FEM solver is batched '
    'and the operator is not. The gap is unsurprising in kind — the FEM solver '
    'still performs the full nonlinear iteration, ten load steps of '
    'Newton–Raphson each requiring an assembly and a linear solve per sample, '
    'whereas the network replaces all of it with one forward pass — but its '
    'size depends entirely on whether the comparison is fair, which is why both '
    'figures are given here rather than only the larger one.')._p)

elements.append(para(
    'The same matched timings give the break-even point: the number of new '
    'problems that must be solved before the operator\'s one-off training cost '
    '(Table 7) is repaid by the per-sample saving over the GPU solver.')._p)

elements.append(new_table(
    ['Case', 'bs=1', 'bs=8', 'bs=32', 'bs=128'],
    [[c] + r for c, r in zip(CASES, BREAKEVEN)])._tbl)
elements.append(para(
    'Table 10c. Break-even against the GPU-native FEM solver, in new problem '
    'instances, at each matched batch size. Computed as the case\'s total '
    'training wall-clock time divided by the per-sample saving (Table 10 minus '
    'Table 10a) at that batch size.')._p)

elements.append(para(
    'Break-even against a GPU baseline is far less favourable than against the '
    'CPU one used in Section 8.3, where it is 52–1,245 samples; both are '
    'reported here rather than only the favourable figure. More importantly, it '
    'is not a single number. Across the six cases and four batch sizes it spans '
    '1,133 to 95,038, a factor of 84, and the batch size assumed matters more '
    'than the case does. The batch-size-128 figures assume 128 problems are '
    'available to solve simultaneously — but a user with 128 problems in hand '
    'would batch the FEM solver too, which is what makes that column the '
    'least favourable one. In the deployment setting the operator is actually '
    'aimed at, where problems arrive one at a time, break-even is 1,133–19,410. '
    'Any break-even figure quoted for this work should therefore name the batch '
    'size it assumes; quoted alone, the number is close to meaningless.')._p)

elements.append(para(
    'Two limitations of these figures should be stated. Both sides are measured '
    'at the study\'s own mesh (N=21, 441 nodes); the balance would shift with '
    'resolution, since the FEM solve grows superlinearly in the number of '
    'degrees of freedom while the operator\'s forward pass grows close to '
    'linearly. And the training cost entering Table 10c is wall-clock time on '
    'the hardware each case actually used, which for the three B2 cases is the '
    'corrected loss-normalized recipe of Section 9.1 at batch size 1 — an order '
    'of magnitude more expensive than the B1 runs, and the reason the B2 rows '
    'break even an order of magnitude later.')._p)

# move everything from the end of the document to just after paragraph 226
anchor = doc.paragraphs[226]._p
for el in elements:
    anchor.addnext(el)
    anchor = el
print(f'inserted {len(elements)} elements after section 8.5')

doc.save(DST)
print('wrote', DST)
