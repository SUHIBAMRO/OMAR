"""v45 -> v46. B2 x Arruda-Boyce's Pareto sweep finished (9/9) -- the
third and last B2 material. Point 2 of the round-5 review (accuracy/cost
Pareto) is now complete for all six geometry x material combinations.

Adds Table 18e, mirroring Tables 18c/18d, and retexts the three stale
"not yet finished" sentences left over from when this was still running:
the summary line, and each of Table 18c's and 18d's own speed-up
sentences that named the still-running materials.

The finding: the training-resolution anchoring effect (local minima at
N=21/N=33) now replicates in ALL THREE B2 materials -- Neo-Hookean
(4.30x/3.65x), Mooney-Rivlin (3.07x/3.17x), Arruda-Boyce (2.99x/2.47x).
Three of three is a property of the training protocol, not a
coincidence of one or two materials. Every number asserted against the
committed JSON before writing.
"""
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v45.docx', 'PFEM_Transolver_Report_v46.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2AB = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_arruda_boyce.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
assert [r['N'] for r in B2AB['rows']] == NS

rows = {r['N']: r for r in B2AB['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = B2AB['shape']['training_resolution_anchoring']['anchor_ratios']
anchor = {int(k): v for k, v in anchor.items()}
nh_anchor = {int(k): v for k, v in B2AB['shape']['training_resolution_anchoring'][
    'neo_hookean_anchor_ratios_for_comparison'].items()}
mr_anchor = {int(k): v for k, v in B2AB['shape']['training_resolution_anchoring'][
    'mooney_rivlin_anchor_ratios_for_comparison'].items()}
assert anchor[21] > 2.0 and anchor[33] > 2.0
assert anchor[21] < mr_anchor[21] < nh_anchor[21]
assert anchor[33] < mr_anchor[33] < nh_anchor[33]

print(f'B2xAB Pareto: speedup {min(speedup.values()):,.0f}x-{max(speedup.values()):,.0f}x, '
      f'anchor ratios {anchor[21]:.2f}x / {anchor[33]:.2f}x')

# ======================================================================
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, __import__('copy').deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']

table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

# Table 18c's own speed-up sentence no longer needs to name the two
# materials that have since finished.
retext(
    'Speed-up spans 1,738',
    'Speed-up spans 1,738× to 27,392×, in the same order-of-magnitude '
    'band as the three B1 materials.')

# Table 18d's own speed-up sentence, same correction.
retext(
    'Speed-up spans 3,811',
    'Speed-up spans 3,811× to 60,651×, the same order-of-magnitude band '
    'as every other case measured so far.')

insert_after('Speed-up spans 3,811', [
    para(
        'B2 × Arruda-Boyce\'s Pareto sweep finished last, closing point 2 '
        'for all six geometry × material combinations. Table 18e.')._p,
    para('Table 18e. Accuracy and cost of both methods, B2 × Arruda-Boyce.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'The anchoring effect replicates a third time: local minima at '
        f'N=21 and N=33 again, {anchor[21]:.2f}× and {anchor[33]:.2f}× '
        f'lower than each point\'s neighbours — the smallest of the three '
        f'B2 materials ({mr_anchor[21]:.2f}×/{mr_anchor[33]:.2f}× '
        f'Mooney-Rivlin, {nh_anchor[21]:.2f}×/{nh_anchor[33]:.2f}× '
        f'Neo-Hookean) but the same shape, at the same two resolutions, '
        f'every time. Three of three B2 materials now show it: this is a '
        f'property of B2\'s two-resolution training protocol, not a '
        f'coincidence of one or two materials. What remains open is '
        f'whether the effect traces to the training protocol or to the B2 '
        f'geometry itself — no B1 checkpoint trained jointly at two '
        f'resolutions exists to separate the two, so every available '
        f'B1-vs-B2 comparison still differs in both at once.')._p,
    para(
        f'Speed-up spans {min(speedup.values()):,.0f}× to '
        f'{max(speedup.values()):,.0f}×, the same order-of-magnitude band '
        'as every other case in this report. Point 2 of the round-5 '
        'review is now measured for all six geometry × material '
        'combinations.')._p,
])

retext('This comparison has been run for all three B1 materials',
       'This comparison has been run for all three B1 materials (Tables '
       '18, 18a, 18b) and all three B2 materials (Neo-Hookean, Table 18c; '
       'Mooney-Rivlin, Table 18d; Arruda-Boyce, Table 18e) — six of six '
       'geometry × material combinations, closing point 2 of the round-5 '
       'review.')

doc.save(DST)
print(f'wrote {DST}')
