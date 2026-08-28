"""Adds the point-5 results to the parallel summary document.

Same three tables as report section 8.8, built from the same module, so the
two documents cannot disagree — last round they were typed twice and then
compared cell by cell to prove they matched.

Kept to the summary's own register, which the user set explicitly: results
only, one line of lead per table, caption underneath, no discussion. The
reasoning lives in the report; this document is what the numbers are.

Placed as a new numbered section 8, after "7. Resolution invariance
(zero-shot)" and before the "Error metric definition" block, so the numbered
sections continue to track the advisor's request list.
"""
import copy

from docx import Document
from docx.oxml.ns import qn

import point5_tables as P5

# Reads the pre-point-5 copy and writes the live one, so re-running this
# script replaces section 8 rather than appending a second one.
SRC = 'PFEM_Summary_Completed_Work.pre_v3.docx'
DST = 'PFEM_Summary_Completed_Work.docx'
doc = Document(SRC)
body = doc.element.body
ref_tbl = next(body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style=None):
    p = doc.add_paragraph() if style is None else doc.add_paragraph(style=style)
    p.add_run(text)
    return p


# the heading style the other numbered sections use
heading_style = None
for p in doc.paragraphs:
    if p.text.strip().startswith('7. Resolution invariance'):
        heading_style = p.style
        anchor_heading = p
        break
assert heading_style is not None, 'section 7 heading not found'

els = []
h = doc.add_paragraph(style=heading_style)
h.add_run('8. Error in physically important quantities beyond displacement')
els.append(h._p)

els.append(para('Trained operator scored on 50 held-out samples per case, N=21, '
                'against each sample\'s own same-mesh FE solution, in FP64. '
                'Mean over the 50 samples, worst sample in parentheses; '
                'percentages.')._p)
els.append(new_table(P5.NORMS_HEAD, P5.norms_rows())._tbl)
els.append(para('Table 15. Displacement and integral-norm error, all six cases')._p)

els.append(para('First Piola-Kirchhoff stress, P = dW/dF by autodiff at the '
                'mesh\'s Gauss points. Means only; percentages.')._p)
els.append(new_table(P5.STRESS_HEAD, P5.stress_rows())._tbl)
els.append(para('Table 16. PK1 stress error. The P11 column on B1 and both shear '
                'columns on B1 divide by a near-zero reference — read the '
                '‖P‖F column, not those.')._p)

els.append(para('Reaction = internal force assembled from the same Gauss-point '
                'stresses on the constrained nodes. Means; percentages.')._p)
els.append(new_table(P5.REACTION_HEAD, P5.reaction_rows())._tbl)
els.append(para('Table 17. Reaction-force error. B2 is constrained on two '
                'symmetry edges and contributes a row for each.')._p)

anchor = anchor_heading._p
# walk forward to the last element belonging to section 7, i.e. stop just before
# the next heading
cur = anchor
last = anchor
while True:
    nxt = cur.getnext()
    if nxt is None:
        break
    if nxt.tag == qn('w:p'):
        from docx.text.paragraph import Paragraph
        pp = Paragraph(nxt, doc)
        if pp.style is not None and pp.style == heading_style and pp.text.strip():
            break
    last = nxt
    cur = nxt

target = last
for el in els:
    target.addnext(el)
    target = el

# ------------------------------------------------------------- the Conclusion
# The Conclusion quotes accuracy in displacement terms only. That is now known
# to be the most forgiving quantity available, so it needs the qualification.
concl = None
for p in doc.paragraphs:
    if p.text.strip().startswith('On efficiency, the picture is consistent'):
        concl = p._p
        break
assert concl is not None, 'the efficiency paragraph of the Conclusion was not found'

extra = para(
    'What the accuracy numbers do and do not cover. Every accuracy figure above '
    'is a displacement error. Scoring the same six checkpoints on quantities '
    'that depend on the displacement gradient shows that figure is a lower '
    f'bound: the H1 semi-norm error is {P5.span("H1_semi_rel")}%, the tangent-'
    f'energy error {P5.span("energy_rel")}% and the aggregate PK1 stress error '
    f'{P5.span("P_rel_L2")}%, each larger than the displacement error of the '
    'same case in all six cases. Integrated quantities behave better — the '
    f'reaction resultant is {P5.reaction_span("reaction_resultant_rel_err")}%, '
    'comparable to the displacement — and '
    'peak stress splits the two benchmarks sharply, at '
    f'{P5.span("P_peak_rel_err", ("B2",))}% on B2 but '
    f'{P5.span("P_peak_rel_err", ("B1",))}% on B1, where the network '
    'consistently overshoots the peak. Section 8 has the detail.')
concl.addnext(extra._p)

doc.save(DST)
print(f'inserted {len(els)} elements as summary section 8; '
      f'{DST} now has {len(Document(DST).tables)} tables')
