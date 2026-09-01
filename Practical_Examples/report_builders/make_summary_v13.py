"""Mirrors report v40 into the summary: Table 24e, the operator's per-member
consistency (coefficient of variation) against Q4's, closing the open
question v12 left -- same JSON, same assertions as make_v40.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; it copies
the current file to .pre_v13.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

PM = json.load(open(os.path.join(PF, 'point9_results',
                                 'mms_operator_per_member_B1_neo_hookean.json')))
FAM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_family_fem_B1_neo_hookean.json')))

METRICS = ['L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel']
SUM = {int(N): PM['per_member_summary'][str(N)] for N in (9, 17, 33)}
Q4 = {int(N): PM['q4_std_over_mean_same_family'][str(N)] for N in (9, 17, 33)}
NS = sorted(SUM)
assert NS == [9, 17, 33], NS

FROWS = {r['N']: r for r in FAM['rows']}
for N in NS:
    for m in METRICS:
        assert abs(SUM[N][m]['mean'] / FROWS[N]['operator_family_mean'][m] - 1) < 2e-4

assert SUM[9]['H1_semi_rel']['std_over_mean'] < 0.01
assert SUM[33]['H1_semi_rel']['std_over_mean'] > 0.1
for N in NS:
    assert SUM[N]['L2_rel']['std_over_mean'] > 0.2
    assert SUM[N]['energy_rel']['std_over_mean'] > 0.3

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v13.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v13.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


insert_after('H1 and stress stay near the ceiling', [
    para(
        'A different question than the mean-ratio table above: is the '
        'operator itself as consistent from member to member as Q4 is? The '
        'same three checkpoints were re-scored per test member (no '
        'retraining) and their coefficient of variation (std/mean) compared '
        'to Q4\'s own, already negligible:')._p,
    new_table(
        ['N', 'L2', 'H1 semi', 'stress', 'energy'],
        [[str(N)] + [f'{SUM[N][m]["std_over_mean"]:.3f} (Q4 {Q4[N][m]:.3f})'
                    for m in METRICS] for N in NS])._tbl,
    para(
        'No, in every norm at every mesh, and in L2/energy not close: the '
        'operator\'s std/mean is '
        f'{min(SUM[N]["L2_rel"]["std_over_mean"] for N in NS):.3f}–'
        f'{max(SUM[N]["L2_rel"]["std_over_mean"] for N in NS):.3f} in L2 and '
        f'{min(SUM[N]["energy_rel"]["std_over_mean"] for N in NS):.3f}–'
        f'{max(SUM[N]["energy_rel"]["std_over_mean"] for N in NS):.3f} in '
        'energy at every mesh — against Q4\'s 0.000–0.007 everywhere. H1 and '
        'stress start indistinguishable from Q4\'s own spread at N = 9 '
        f'({SUM[9]["H1_semi_rel"]["std_over_mean"]:.3f} and '
        f'{SUM[9]["stress_rel_L2"]["std_over_mean"]:.3f}) and grow away from '
        f'it by N = 33 ({SUM[33]["H1_semi_rel"]["std_over_mean"]:.3f} and '
        f'{SUM[33]["stress_rel_L2"]["std_over_mean"]:.3f}) — the same '
        'ceiling-proximity effect the mean ratio shows, now visible in '
        'per-member reliability too.')._p,
])

doc.save(SRC)
print('summary v13: Table 24e added -- operator std/mean at N=33: '
      f'L2 {SUM[33]["L2_rel"]["std_over_mean"]:.3f}, '
      f'H1 {SUM[33]["H1_semi_rel"]["std_over_mean"]:.3f}, '
      f'stress {SUM[33]["stress_rel_L2"]["std_over_mean"]:.3f}, '
      f'energy {SUM[33]["energy_rel"]["std_over_mean"]:.3f}')
