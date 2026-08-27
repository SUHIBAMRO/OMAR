"""Adds the matched-batch-size comparison to the parallel summary document.

Kept to the summary's own register, which the user set explicitly: results
only, one line of lead per table, caption underneath, no discussion. The
reasoning lives in the report; this document is what the numbers are.

Placed in section 5 directly after Table 10, since Tables 10a-c are read
against Table 10 row by row.
"""
import copy
from docx import Document
from docx.oxml.ns import qn

SRC = 'PFEM_Summary_Completed_Work.docx'
DST = 'PFEM_Summary_Completed_Work.docx'
doc = Document(SRC)
body = doc.element.body

CASES = ['B1 × Neo-Hookean', 'B1 × Mooney-Rivlin', 'B1 × Arruda-Boyce',
         'B2 × Neo-Hookean', 'B2 × Mooney-Rivlin', 'B2 × Arruda-Boyce']
LATENCY = [['4.582', '0.573', '0.336', '0.292'],
           ['4.832', '0.607', '0.336', '0.292'],
           ['4.753', '0.599', '0.336', '0.292'],
           ['4.680', '0.585', '0.336', '0.291'],
           ['4.714', '0.594', '0.336', '0.291'],
           ['4.799', '0.602', '0.336', '0.292']]
SPEEDUP = [['360×', '833×', '1,134×', '1,215×'],
           ['424×', '812×', '1,154×', '1,232×'],
           ['531×', '878×', '1,221×', '1,297×'],
           ['356×', '819×', '1,135×', '1,217×'],
           ['426×', '830×', '1,158×', '1,235×'],
           ['544×', '879×', '1,223×', '1,297×']]
BREAKEVEN = [['1,745', '6,021', '7,543', '8,112'],
             ['1,363', '5,663', '7,179', '7,751'],
             ['1,133', '5,441', '6,956', '7,554'],
             ['19,410', '67,391', '84,627', '90,990'],
             ['17,033', '69,404', '87,884', '95,038'],
             ['9,530', '46,993', '60,490', '65,698']]

ref_tbl = next(body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


# anchor: the "Table 10." caption paragraph in section 5
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Table 10.'):
        anchor = p._p
        break
assert anchor is not None, 'Table 10 caption not found'

HEAD = ['Case', 'bs=1', 'bs=8', 'bs=32', 'bs=128']
els = []

els.append(para('Transolver re-benchmarked at the same batch sizes as the GPU '
                'FEM solver above, so the two are directly comparable.')._p)
els.append(new_table(HEAD, [[c] + r for c, r in zip(CASES, LATENCY)])._tbl)
els.append(para('Table 10a. Transolver inference latency, ms/sample')._p)

els.append(para('Operator vs. GPU FEM, both at the same batch size '
                '(Table 10 ÷ Table 10a).')._p)
els.append(new_table(HEAD, [[c] + r for c, r in zip(CASES, SPEEDUP)])._tbl)
els.append(para('Table 10b. Operator speed-up over GPU-native FEM, matched batch '
                'sizes. Unmatched (FEM at bs=128 vs. operator at bs=1) gives '
                '73–80×.')._p)

els.append(para('New problems that must be solved before training cost (Table 7) '
                'is repaid by the per-sample saving over the GPU solver.')._p)
els.append(new_table(HEAD, [[c] + r for c, r in zip(CASES, BREAKEVEN)])._tbl)
els.append(para('Table 10c. Break-even vs. GPU-native FEM, in problem instances. '
                'Range 1,133–95,038: the assumed batch size matters more than the '
                'case. Against the CPU baseline it is 52–1,245.')._p)

for el in els:
    anchor.addnext(el)
    anchor = el

doc.save(DST)
print(f'inserted {len(els)} elements; {DST} now has '
      f'{len(Document(DST).tables)} tables')
