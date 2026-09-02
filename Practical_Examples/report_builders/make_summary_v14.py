"""Mirrors report v41 into the summary: Mooney-Rivlin and Arruda-Boyce's B2
results, closing point 7 across all six cases. Same JSONs and assertions as
make_v41.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v14.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

NH = json.load(open(os.path.join(PF, 'point7a_results',
                                 'B2_zeroshot_fixedselection.json')))
MR = json.load(open(os.path.join(PF, 'point7a_results',
                                 'B2_mooney_rivlin_zeroshot_fixedselection.json')))
AB = json.load(open(os.path.join(PF, 'point7a_results',
                                 'B2_arruda_boyce_zeroshot_fixedselection.json')))
B1 = {m: json.load(open(os.path.join(PF, 'point7a_results',
                                     f'zeroshot_B1_{m}.json')))
      for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')}

NS = [13, 17, 25, 29, 37, 41, 49]
NH_pc = {r['N']: r['mean_rel_L2_vs_fine_reference'] for r in NH['rows']}
MR_pc = {r['N']: r['mean_rel_L2_vs_fine_reference'] for r in MR['rows']}
AB_pc = {r['N']: r['mean_rel_L2_vs_fine_reference'] for r in AB['rows']}
assert set(NH_pc) == set(MR_pc) == set(AB_pc) == set(NS)

B1_at_NS = {m: {r['N']: r['mean_rel_L2_vs_fine_reference']
                for r in B1[m]['rows'] if r['N'] in NS} for m in B1}
B1_ALL = [v for d in B1_at_NS.values() for v in d.values()]
B1_LO, B1_HI = min(B1_ALL), max(B1_ALL)

NH_SPREAD = max(NH_pc.values()) / min(NH_pc.values())
MR_SPREAD = max(MR_pc.values()) / min(MR_pc.values())
AB_SPREAD = max(AB_pc.values()) / min(AB_pc.values())
B1_SPREAD = max(max(d.values()) / min(d.values()) for d in B1_at_NS.values())
assert MR_SPREAD > NH_SPREAD > AB_SPREAD > B1_SPREAD

WORST_AT_MESH = {N: max((NH_pc[N], 'Neo-Hookean'), (MR_pc[N], 'Mooney-Rivlin'),
                        (AB_pc[N], 'Arruda-Boyce'))[1] for N in NS}
AB_WORST_COUNT = sum(1 for v in WORST_AT_MESH.values() if v == 'Arruda-Boyce')
assert AB_WORST_COUNT == 6, WORST_AT_MESH

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v14.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v14.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def replace(prefix, els):
    v = find(prefix)
    t = v._p
    for el in els:
        t.addnext(el)
        t = el
    v._p.getparent().remove(v._p)


replace('Only Neo-Hookean has been re-run', [
    para(
        'Mooney-Rivlin and Arruda-Boyce were re-run immediately after, same '
        'protocol. Every B1/B2 × material combination now has a valid '
        'zero-shot result:')._p,
    new_table(['N', 'Neo-Hookean', 'Mooney-Rivlin', 'Arruda-Boyce'],
              [[str(N), f'{NH_pc[N]:.4f}', f'{MR_pc[N]:.4f}', f'{AB_pc[N]:.4f}']
               for N in NS])._tbl,
    para(
        f'All three B2 materials keep the pattern Neo-Hookean set: real '
        f'resolution invariance, weaker than B1\'s. Each material\'s spread '
        f'— {MR_SPREAD:.2f}× Mooney-Rivlin, {NH_SPREAD:.2f}× Neo-Hookean, '
        f'{AB_SPREAD:.2f}× Arruda-Boyce — exceeds B1\'s worst case '
        f'{B1_SPREAD:.2f}×, no exception. Arruda-Boyce is least accurate at '
        f'{AB_WORST_COUNT} of 7 meshes but also most consistent (narrowest '
        f'spread of the three) — different properties, and this case '
        f'separates them. Mooney-Rivlin and Neo-Hookean cross rather than '
        f'one dominating: Mooney-Rivlin wins at every mesh coarser than the '
        f'training resolutions (13, 17, 25) and loses at every mesh finer '
        f'(37, 41, 49).')._p,
    para(
        'One more data point on the open patience question: Arruda-Boyce\'s '
        'early stop actually fired (the other two ran the full 4,000-epoch '
        f'budget) — best at epoch {AB["training"]["best_epoch"]:,}, stopped '
        f'at {AB["training"]["stopped_epoch"]:,}, exactly best_epoch + '
        f'patience({AB["training"]["early_stop_patience"]}) × '
        f'validate_every({AB["training"]["validate_every"]}). Patience 8 '
        'would have stopped this run before its actual best.')._p,
])

doc.save(SRC)
print(f'summary v14: B2 complete -- spreads NH {NH_SPREAD:.2f}x, '
      f'MR {MR_SPREAD:.2f}x, AB {AB_SPREAD:.2f}x, B1 worst {B1_SPREAD:.2f}x')
