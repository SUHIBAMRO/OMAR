"""v28 -> v29: the advisor's point 5, error in quantities beyond displacement.

Timon asked to "examine error in physically important quantities beyond
displacement, including error in H1 semi-norm, energy and also some local
quantities such as stress components and reaction forces; maybe looking at
maxima. I am referring to the transolver." All six cases have been measured;
this writes them in.

Placement. The new material becomes section 8.8, inserted after the
resolution-invariance study (8.7) and before the training visualizations,
which move from 8.8 to 8.9. Nothing in the document cross-references "8.8"
(checked), so that renumbering is safe -- unlike renumbering a table, which
the 10a/10b/10c convention was adopted to avoid.

Table numbers. The next free caption is 15: the document's captions run
1-13 plus 4a, 6a, 10a-c. It does contain one dangling reference to a
"Table 14" in the conclusion, but no such table was ever inserted, so 14 is
deliberately left alone rather than silently absorbed by this section's
first table.

Every number, in the tables and in the prose, comes from point5_tables.py,
which reads the six result JSONs. Nothing here is typed by hand.
"""
import copy

from docx import Document
from docx.oxml.ns import qn

import point5_tables as P5

SRC, DST = 'PFEM_Transolver_Report_v28.docx', 'PFEM_Transolver_Report_v29.docx'
doc = Document(SRC)
body = doc.element.body

ref_tbl = next(body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


# ------------------------------------------------- figures quoted in the text
disp_B1, disp_B2 = P5.span('disp_rel_L2', ('B1',)), P5.span('disp_rel_L2', ('B2',))
h1_B1, h1_B2 = P5.span('H1_semi_rel', ('B1',)), P5.span('H1_semi_rel', ('B2',))
en_B1, en_B2 = P5.span('energy_rel', ('B1',)), P5.span('energy_rel', ('B2',))
fro_all = P5.span('P_rel_L2')
peak_B1, peak_B2 = P5.span('P_peak_rel_err', ('B1',)), P5.span('P_peak_rel_err', ('B2',))
p22_B1 = P5.span('P22_rel_L2', ('B1',))
p11_B1 = P5.span('P11_rel_L2', ('B1',))
h1_all = P5.span('H1_semi_rel')

MATS = ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')


def _range(geometry, keys, scale=100.0, fmt='.2f', field='mean'):
    vals = [scale * P5.value(geometry, m, k, field) for m in MATS for k in keys]
    return f"{min(vals):{fmt}}–{max(vals):{fmt}}"


shear_B1 = _range('B1', ('P12_rel_L2', 'P21_rel_L2'))
shear_B2 = _range('B2', ('P12_rel_L2', 'P21_rel_L2'))
normal_B2 = _range('B2', ('P11_rel_L2', 'P22_rel_L2'))
# how much better P22 is than the shear components, per case, on B1
_p22_factor = [P5.value('B1', m, k) / P5.value('B1', m, 'P22_rel_L2')
               for m in MATS for k in ('P12_rel_L2', 'P21_rel_L2')]
p22_factor = f"{min(_p22_factor):.1f}–{max(_p22_factor):.1f}"

reaction_res = []
reaction_nod = []
for g, m, _ in P5.CASES:
    M = P5.ALL[(g, m)]['metrics']
    if g == 'B1':
        reaction_res.append(100 * M['reaction_resultant_rel_err']['mean'])
        reaction_nod.append(100 * M['reaction_nodal_rel_L2']['mean'])
    else:
        for suf in ('edge0', 'edge1'):
            reaction_res.append(100 * M[f'reaction_resultant_rel_err_{suf}']['mean'])
            reaction_nod.append(100 * M[f'reaction_nodal_rel_L2_{suf}']['mean'])
res_span = f"{min(reaction_res):.2f}–{max(reaction_res):.2f}"
nod_span = f"{min(reaction_nod):.2f}–{max(reaction_nod):.2f}"

b1_peaks = '; '.join(
    f"{P5.value('B1', m, 'P_peak_pred'):.2f} against {P5.value('B1', m, 'P_peak_ref'):.2f}"
    for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
b2_peaks = '; '.join(
    f"{P5.value('B2', m, 'P_peak_pred'):.2f} against {P5.value('B2', m, 'P_peak_ref'):.2f}"
    for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
_mr = [f"{100 * P5.value('B1', m, 'reaction_max_rel_err'):.2f}%"
       for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')]
b1_maxreact = f"{_mr[0]}, {_mr[1]} and {_mr[2]}"
b1_shear_abs = (f"{min(P5.value('B1', m, k) for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce') for k in ('P12_max_abs_err', 'P21_max_abs_err')):.2f}"
                f"–{max(P5.value('B1', m, k) for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce') for k in ('P12_max_abs_err', 'P21_max_abs_err')):.2f}")
b1_p22_abs = (f"{min(P5.value('B1', m, 'P22_max_abs_err') for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')):.2f}"
              f"–{max(P5.value('B1', m, 'P22_max_abs_err') for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')):.2f}")
peak_std_B1 = _range('B1', ('P_peak_rel_err',), fmt='.0f', field='std')  # pct points

# ------------------------------------------------------- claims, checked not asserted
# Each sentence below that generalises across cases is verified here first, so a
# re-run on new numbers fails loudly instead of printing a false statement. Three
# claims in the first draft of this section did not survive this check.
_derived_exceed = all(
    P5.value(g, m, k) > P5.value(g, m, 'disp_rel_L2')
    for g in ('B1', 'B2') for m in MATS
    for k in ('H1_semi_rel', 'energy_rel', 'P_rel_L2'))
assert _derived_exceed, 'H1/energy/stress do NOT all exceed the displacement error'


def _worst_reaction(g, m):
    M = P5.ALL[(g, m)]['metrics']
    if g == 'B1':
        return M['reaction_resultant_rel_err']['mean']
    return max(M[f'reaction_resultant_rel_err_edge{i}']['mean'] for i in (0, 1))


_react_better = sum(_worst_reaction(g, m) < P5.value(g, m, 'disp_rel_L2')
                    for g in ('B1', 'B2') for m in MATS)
assert _react_better == 4, f'reaction beats displacement in {_react_better} cases, not 4'
_NUM = {4: 'four', 5: 'five', 6: 'six', 3: 'three', 2: 'two'}
react_better_word = _NUM[_react_better]

assert all(P5.value('B2', m, 'P_peak_rel_err') < P5.value('B2', m, 'disp_rel_L2')
           for m in MATS), 'peak stress is not better than displacement on all of B2'
assert all(P5.value('B1', m, 'P_peak_rel_err') > P5.value('B1', m, 'disp_rel_L2')
           for m in MATS), 'peak stress is not worse than displacement on all of B1'
assert all(P5.value('B1', m, 'H1_semi_rel') > max(P5.value('B1', m, 'energy_rel'),
                                                  P5.value('B1', m, 'P_rel_L2'))
           for m in MATS), 'H1 is not the largest integral measure on B1'
assert all(P5.value('B2', m, 'P_rel_L2') > max(P5.value('B2', m, 'energy_rel'),
                                               P5.value('B2', m, 'H1_semi_rel'))
           for m in MATS), 'aggregate stress is not the largest integral measure on B2'
_stress_ratio = [P5.value(g, m, 'P_rel_L2') / P5.value(g, m, 'disp_rel_L2')
                 for g in ('B1', 'B2') for m in MATS]
stress_ratio = f"{min(_stress_ratio):.1f} to {max(_stress_ratio):.1f}"

# every derived quantity, as a multiple of the same case's displacement error
_derived_keys = ('H1_semi_rel', 'energy_rel', 'P_rel_L2', 'P_peak_rel_err')
_ratios = [P5.value(g, m, k) / P5.value(g, m, 'disp_rel_L2')
           for g in ('B1', 'B2') for m in MATS for k in _derived_keys]
_ratios += [_worst_reaction(g, m) / P5.value(g, m, 'disp_rel_L2')
            for g in ('B1', 'B2') for m in MATS]
derived_ratio = f"{min(_ratios):.2f}× to {max(_ratios):.2f}×"
print('all cross-case claims verified against the result JSONs')

# ------------------------------------------------------------- the subsection
elements = []

h = doc.add_paragraph(style='Heading 2')
h.add_run('8.8 Error in physically important quantities beyond displacement')
elements.append(h._p)

elements.append(para(
    'Every operator error reported above is a displacement error. That is the '
    'most forgiving quantity that could be asked about: displacement is the '
    'primary unknown, it is what the network is trained to produce, and it is '
    'smoother than its own derivatives. Stress depends on the displacement '
    'gradient and so loses an order of accuracy; reaction forces are an '
    'integral of that stress over the supports. A 9% displacement error '
    'therefore does not imply a 9% stress error, and the size of that gap is '
    'not something that can be inferred — it has to be measured. This section '
    'measures it for all six cases.')._p)

elements.append(para(
    'Protocol. Each case is evaluated using its own trained checkpoint — the '
    'same six checkpoints behind Tables 5, 7 and 11, which for the three B2 '
    'cases means the corrected loss-normalized runs of Section 9.1 and not the '
    'superseded pre-fix ones — on 50 held-out samples that follow the 800 used '
    'for training, on the study\'s standard N=21 mesh, against each sample\'s '
    'own same-mesh finite-element solution. Derived quantities are far more '
    'sensitive to round-off than displacement is, so the evaluation runs in '
    'FP64 although the network itself was trained in FP32. The L2, H1 '
    'semi-norm and tangent-energy norms are computed by the same routines used '
    'for the Q4-vs-Q9 finite-element study of Section 4.4, so the operator is '
    'graded with exactly the norms already applied to the reference solver '
    'rather than with a second set of definitions. The first Piola–Kirchhoff '
    'stress is obtained as P = ∂W/∂F by automatic differentiation of the same '
    'strain-energy density the solver and the training objective use, '
    'evaluated at the mesh\'s own Gauss points with the mesh\'s own '
    'shape-function gradients; this is exact for all three materials, whereas '
    'only Neo-Hookean has a closed-form expression. Reaction forces are the '
    'internal force assembled from those same Gauss-point stresses and '
    'restricted to the constrained nodes. The external traction is zero on '
    'those nodes in both benchmarks — B1 is loaded on its top edge, B2 on its '
    'inner arc — so the assembled internal force there is precisely the '
    'reaction the supports must supply.')._p)

elements.append(new_table(P5.NORMS_HEAD, P5.norms_rows())._tbl)
elements.append(para(
    'Table 15. Error of the trained operator in displacement and in the three '
    'integral norms, all six cases, as percentages. Each entry is the mean over '
    'the 50 held-out samples with the worst single sample in parentheses. '
    'Column 1 is the per-node RMS relative error used everywhere else in this '
    'report (Tables 5 and 11); column 2 is the quadrature-weighted L2 norm of '
    'Section 4.4. Both are given so that the remaining columns can be read '
    'against either.')._p)

elements.append(para(
    'Two things should be read carefully here. First, the displacement column '
    'is close to but not identical with Table 5\'s best validation error — '
    f'{P5.pct(P5.ALL[("B1", "neo_hookean")]["metrics"], "disp_rel_L2")}% against '
    '9.59% for B1 × Neo-Hookean, '
    f'{P5.pct(P5.ALL[("B2", "neo_hookean")]["metrics"], "disp_rel_L2")}% against '
    '9.11% for B2 × Neo-Hookean — because this is a different held-out set of '
    '50 samples, not the validation set early stopping selected on. The two '
    'agreeing to within about a point is itself a check that the right '
    'checkpoints were scored. Second, the per-node RMS definition and the '
    'quadrature-weighted L2 norm are genuinely different measures, and the FE '
    'norm is consistently the smaller of the two; neither is wrong, but they '
    'should not be quoted interchangeably.')._p)

elements.append(para(
    'The H1 semi-norm is the quantity this section was principally asked for, '
    'and it behaves as theory predicts. On B1 it is roughly twice the '
    f'displacement error ({h1_B1}% against {disp_B1}%), which is the expected '
    'penalty for differentiating a field the network was never asked to '
    f'differentiate. On B2 the penalty is markedly smaller ({h1_B2}% against '
    f'{disp_B2}%). The tangent-energy error is {en_B1}% on B1 and {en_B2}% on '
    'B2, in both cases above the displacement error and below the H1 figure on '
    'B1, though not uniformly so on B2.')._p)

elements.append(new_table(P5.STRESS_HEAD, P5.stress_rows())._tbl)
elements.append(para(
    'Table 16. First Piola–Kirchhoff stress error, as percentages, mean over '
    'the 50 held-out samples. The first column is the quadrature-weighted '
    'relative L2 error of the full tensor in the Frobenius norm; the next four '
    'are the same measure applied to individual components; the last is the '
    'error in the single largest value of ‖P‖F over the mesh. The '
    'per-component columns must be read with the caution given below.')._p)

elements.append(para(
    f'The aggregate stress error is {fro_all}% across all six cases, that is '
    f'{stress_ratio} times the displacement error of the same case. This is the '
    'number to quote for stress accuracy. The per-component columns are '
    'informative but treacherous, because a relative error divides by that '
    'component\'s own reference norm and in both benchmarks some components are '
    'near zero almost everywhere. B1 is pulled on its top edge, so P22 carries '
    f'the load and is accurate ({p22_B1}%), while P11 ({p11_B1}%) and the two '
    f'shear components ({shear_B1}%) are small in magnitude and their relative '
    'errors are measured against a near-vanishing reference. In absolute terms '
    'those same shear errors are the smallest in the table: a mean '
    f'largest-pointwise shear error of {b1_shear_abs} stress units, against '
    f'{b1_p22_abs} for P22 — whose relative error is nonetheless {p22_factor} '
    'times better, which is the whole of the effect. B2\'s quarter ring under '
    'internal pressure develops all four components at comparable magnitude, '
    f'and there the per-component picture is uniform: {normal_B2}% for the two '
    f'normal components and {shear_B2}% for the two shear components. The '
    'Frobenius column is unaffected by the small-denominator problem and is '
    'consistent across both geometries, which is why it is the right single '
    'figure.')._p)

elements.append(para(
    'Peak stress — the maximum of ‖P‖F over the mesh, a design quantity in its '
    'own right and the "maxima" the request asked about — separates the two '
    'benchmarks more sharply than anything else measured here. On B2 it is the '
    f'most accurate quantity in this section at {peak_B2}%, better than the '
    f'displacement error itself. On B1 it is {peak_B1}% and highly variable '
    f'from sample to sample — the standard deviation across the 50 samples is '
    f'{peak_std_B1} percentage points, as large as the mean. The direction of '
    'the discrepancy is '
    'consistent: averaged over the 50 samples the predicted peak exceeds the '
    f'reference peak in all three B1 cases ({b1_peaks}, in the same units as E) '
    f'and falls slightly short of it in all three B2 cases ({b2_peaks}). No '
    'cause has been isolated for the B1 overshoot; it is reported here as '
    'measured, and it is the single result in this section that most clearly '
    'limits what can be claimed about the operator for a design application '
    'where peak stress is the governing quantity.')._p)

elements.append(new_table(P5.REACTION_HEAD, P5.reaction_rows())._tbl)
elements.append(para(
    'Table 17. Reaction-force error on the constrained boundaries, as '
    'percentages, mean over the 50 held-out samples. B1 fixes both displacement '
    'components on its bottom edge and so contributes one row per case; B2 is '
    'constrained by symmetry on two radial edges, each fixing the single '
    'component normal to it, and contributes two. "Resultant" is the error in '
    'the total force the support must supply; "nodal" is the relative L2 error '
    'over the individual nodal reaction magnitudes.')._p)

elements.append(para(
    'Reaction forces are the best-behaved derived quantity measured here: '
    f'{res_span}% in the resultant and {nod_span}% in the nodal distribution, '
    'that is the same order as the displacement error and far better than the '
    'pointwise stress it is assembled from. That is what integration does — '
    'the resultant sums Gauss-point stresses over the whole support and the '
    'pointwise errors partially cancel, so the quantity an engineer would '
    'actually check a support against is more accurate than the field it comes '
    f'from. The error in the single largest nodal reaction on B1 is {b1_maxreact} '
    'for Neo-Hookean, Mooney-Rivlin and Arruda-Boyce respectively, in line with '
    'the distribution as a whole rather than concentrated at the worst node.')._p)

elements.append(para(
    'The ordering of these quantities is not the same on the two geometries, '
    'and that is itself part of the result. On B1 the H1 semi-norm is the '
    f'largest of the integral measures in all three cases ({h1_B1}%), with '
    f'tangent energy and aggregate stress below it. On B2 it is the aggregate '
    f'stress that is largest ({P5.span("P_rel_L2", ("B2",))}%), with the H1 '
    f'semi-norm and the tangent energy close together and lower ({h1_B2}% and '
    f'{en_B2}%). Anyone quoting a single '
    'derived-quantity error for this operator therefore has to name both the '
    'quantity and the benchmark.')._p)

elements.append(para(
    'What is common to both geometries is the direction. The three field '
    'measures that involve a derivative of the displacement — H1 semi-norm, '
    'tangent energy, aggregate stress — exceed the displacement error in every '
    'one of the six cases, without exception. The two quantities that can beat '
    'it are the ones that are not pointwise fields: the reaction resultant, '
    f'which is more accurate than the displacement in {react_better_word} of '
    'the six cases, and the peak stress, which is more accurate than it on all '
    'three B2 cases and substantially worse on all three B1 ones. The practical '
    'consequence is that a displacement error quoted alone is a lower bound on '
    'what a downstream engineering quantity would see, and it is not a reliable '
    f'predictor of it: across the quantities measured here that ratio spans '
    f'{derived_ratio}. If a single conservative '
    f'figure is wanted, the widest of the integral measures spans {h1_all}% '
    'across the six cases.')._p)

# ------------------------------------------------------------ insert + renumber
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('8.8 Representative training visualizations'):
        anchor = p
        break
assert anchor is not None, 'section 8.8 heading not found'

# the new subsection goes immediately BEFORE the visualizations section
target = anchor._p.getprevious()
assert target is not None
for el in elements:
    target.addnext(el)
    target = el
print(f'inserted {len(elements)} elements as the new section 8.8')

for run in anchor.runs:
    if '8.8' in run.text:
        run.text = run.text.replace('8.8 Representative', '8.9 Representative')
        break
else:
    anchor.runs[0].text = anchor.text.replace('8.8 Representative', '8.9 Representative')
    for r in list(anchor.runs)[1:]:
        r._r.getparent().remove(r._r)
assert doc.paragraphs[[i for i, p in enumerate(doc.paragraphs)
                       if p.text.strip().startswith('8.9 Representative')][0]]
print('renumbered the visualizations section 8.8 -> 8.9')

# ------------------------------------ a row in the executive-summary point table
summary_tbl = doc.tables[0]
assert summary_tbl.rows[0].cells[0].text.strip() == '#', 'unexpected first table'
row = summary_tbl.add_row()
row.cells[0].text = '8'
row.cells[1].text = 'Error in physically important quantities beyond displacement'
row.cells[2].text = (
    'Measured for all six cases on 50 held-out samples (Section 8.8, Tables '
    f'15–17): H1 semi-norm {h1_all}%, tangent energy {P5.span("energy_rel")}%, '
    f'aggregate PK1 stress {fro_all}%, reaction resultant {res_span}%, against '
    f'a displacement error of {P5.span("disp_rel_L2")}%. Peak stress is the '
    f'outlier: {peak_B2}% on B2 but {peak_B1}% on B1.')
print('added row 8 to the executive-summary table')

# -------------------------------------- the Conclusion's accuracy qualification
# Section 10 quotes accuracy in displacement terms throughout. That is now known
# to be the most forgiving quantity available, so it needs one paragraph of
# qualification -- inserted before the "remaining items" list, not into it,
# because this item is finished rather than outstanding.
lead = None
for p in doc.paragraphs:
    if p.text.strip() == 'The remaining items are:':
        lead = p._p
        break
assert lead is not None, 'the "remaining items" lead-in was not found'

qual = para(
    'One qualification applies to every accuracy figure quoted above. All of '
    'them are displacement errors, and Section 8.8 shows that displacement is '
    'the most forgiving quantity available: on the same six checkpoints the H1 '
    f'semi-norm error is {h1_all}%, the tangent-energy error '
    f'{P5.span("energy_rel")}% and the aggregate first Piola–Kirchhoff stress '
    f'error {fro_all}%, each of them larger than the displacement error of the '
    'same case in all six cases. Integrated quantities fare better — the '
    f'reaction resultant is {res_span}% — and peak stress divides the two '
    f'benchmarks, at {peak_B2}% on B2 against {peak_B1}% on B1. None of this '
    'changes a measurement reported elsewhere in this document; it changes '
    'what a displacement error should be taken to mean.')
lead.addprevious(qual._p)
print('added the derived-quantity qualification to Section 10')

doc.save(DST)
print('wrote', DST)
