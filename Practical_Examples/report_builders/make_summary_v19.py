"""Mirrors report v46 into the summary: Table 18e, B2 x Arruda-Boyce's
Pareto -- the third and last B2 material, closing point 2 for all six
geometry x material combinations. Same JSON and assertions as make_v46.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v19.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2AB = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_arruda_boyce.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
rows = {r['N']: r for r in B2AB['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = {int(k): v for k, v in B2AB['shape']['training_resolution_anchoring']['anchor_ratios'].items()}
nh_anchor = {int(k): v for k, v in B2AB['shape']['training_resolution_anchoring'][
    'neo_hookean_anchor_ratios_for_comparison'].items()}
mr_anchor = {int(k): v for k, v in B2AB['shape']['training_resolution_anchoring'][
    'mooney_rivlin_anchor_ratios_for_comparison'].items()}
assert anchor[21] > 2.0 and anchor[33] > 2.0
assert anchor[21] < mr_anchor[21] < nh_anchor[21]
assert anchor[33] < mr_anchor[33] < nh_anchor[33]

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v19.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v19.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def retext(prefix, text):
    p = find(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']
table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

# Table 18d's own paragraph named Arruda-Boyce as still running; correct
# that before inserting Table 18e right after it.
retext('The anchoring effect replicates',
       'The anchoring effect replicates: local minima at N=21 and N=33 '
       'again, 3.07× and 3.17× lower than each point\'s neighbours — '
       'smaller than Neo-Hookean\'s 4.30×/3.65× but the same shape, at '
       'the same two resolutions.')

insert_after('The anchoring effect replicates', [
    para(
        f'Table 18e. B2 × Arruda-Boyce (third and last B2 material). '
        f'Speed-up {min(speedup.values()):,.0f}×–'
        f'{max(speedup.values()):,.0f}×.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'The anchoring effect replicates a third time: '
        f'{anchor[21]:.2f}× and {anchor[33]:.2f}× at N=21/N=33 — the '
        f'smallest of the three B2 materials but the same shape, every '
        f'time. Three of three now show it: a property of B2\'s '
        f'two-resolution training protocol, not a coincidence of one or '
        f'two materials. Whether it traces to the protocol or to the B2 '
        f'geometry itself remains open — no B1 checkpoint trained '
        f'jointly at two resolutions exists to separate the two. Point 2 '
        f'of the round-5 review is now measured for all six geometry × '
        f'material combinations.')._p,
])

retext('Remaining work.',
       'Remaining work. Zero-shot resolution invariance is reported for '
       'all six geometry–material cases (Table 12c closed the last of '
       'them); nothing remains open there. The B2-geometry high-DOF, '
       'element-order study (Section 4.4\'s deeper Q4-vs-Q9 check) was '
       'cancelled at the sponsor\'s request and is not planned, not '
       'merely postponed. Point 2\'s accuracy/cost Pareto is now measured '
       'for all six geometry × material combinations (Table 18e closed '
       'the last of them). The one genuine open item is the '
       'out-of-distribution attribution extension named in the report\'s '
       'Section 10. It does not affect the results reported above.')

doc.save(SRC)
print(f'summary v19: Table 18e added -- anchor ratios {anchor[21]:.2f}x/'
      f'{anchor[33]:.2f}x -- point 2 complete for all six cases')
