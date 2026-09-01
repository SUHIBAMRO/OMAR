"""v39 -> v40. Table 24e: is the operator consistent across the family, or
only on average? v39's Table 24d answered this for the MEAN ratio across
four norms. This is the same question for each mesh's own 16-member spread
-- std/mean of the operator's per-member error, next to Q4's (already known
to be negligible from mms_family_fem's own Q4_spread_stdev_over_mean).

Nothing new measured in this document's sense: the three operator
checkpoints (N=9, 17, 33) were re-scored per member with no retraining
(mms_operator_per_member.py), recorded and asserted against the family
sweep's own means in point9_results/mms_operator_per_member_B1_neo_hookean.json.
This builder reads that file and asserts the same agreement independently.
"""
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v39.docx', 'PFEM_Transolver_Report_v40.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

PM = json.load(open(os.path.join(PF, 'point9_results',
                                 'mms_operator_per_member_B1_neo_hookean.json')))
FAM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_family_fem_B1_neo_hookean.json')))

METRICS = ['L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel']
SUM = {int(N): PM['per_member_summary'][str(N)] if str(N) in PM['per_member_summary']
       else PM['per_member_summary'][N] for N in (9, 17, 33)}
Q4 = {int(N): PM['q4_std_over_mean_same_family'][str(N)]
      if str(N) in PM['q4_std_over_mean_same_family']
      else PM['q4_std_over_mean_same_family'][N] for N in (9, 17, 33)}
NS = sorted(SUM)
assert NS == [9, 17, 33], NS

FROWS = {r['N']: r for r in FAM['rows']}
for N in NS:
    for m in METRICS:
        got = SUM[N][m]['mean']
        want = FROWS[N]['operator_family_mean'][m]
        assert abs(got / want - 1) < 2e-4, (N, m, got, want)

# The finding this table exists to carry.
assert SUM[9]['H1_semi_rel']['std_over_mean'] < 0.01
assert SUM[33]['H1_semi_rel']['std_over_mean'] > 0.1
for N in NS:
    assert SUM[N]['L2_rel']['std_over_mean'] > 0.2
    assert SUM[N]['energy_rel']['std_over_mean'] > 0.3

print('operator std/mean at N=33:',
      {m: SUM[33][m]['std_over_mean'] for m in METRICS})

# ======================================================================
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, __import__('copy').deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


insert_after('Per interval, the direction does not reverse', [
    para(
        'A different question than Table 24d\'s: not whether the operator\'s '
        'mean error tracks Q4\'s, but whether the operator itself is as '
        'consistent from member to member as Q4 is. The three checkpoints '
        'above were re-scored per test member rather than only by their '
        'mean — no retraining — and Table 24e reports each mesh\'s '
        'coefficient of variation (std/mean) next to Q4\'s own, already '
        'shown negligible.')._p,
    para(
        'Table 24e. Coefficient of variation (std/mean) across the '
        '16-member family, operator against Q4, all four norms. The means '
        'behind this table reproduce Table 24c/24d\'s to four significant '
        'figures, confirming the same checkpoints and the same members.')._p,
    new_table(
        ['N', 'L2', 'H1 semi', 'stress', 'energy'],
        [[str(N)] + [f'{SUM[N][m]["std_over_mean"]:.3f} (Q4 {Q4[N][m]:.3f})'
                    for m in METRICS] for N in NS])._tbl,
    para(
        'The answer is no, in every norm, at every mesh — and in two of '
        'the four it is not close. Q4\'s std/mean is 0.000–0.007 across '
        'every metric and mesh: essentially member-independent. The '
        f'operator\'s is {min(SUM[N]["L2_rel"]["std_over_mean"] for N in NS):.3f}'
        f'–{max(SUM[N]["L2_rel"]["std_over_mean"] for N in NS):.3f} in L2 and '
        f'{min(SUM[N]["energy_rel"]["std_over_mean"] for N in NS):.3f}–'
        f'{max(SUM[N]["energy_rel"]["std_over_mean"] for N in NS):.3f} in '
        'energy at every mesh tested — roughly two orders of magnitude more '
        'variable than Q4, already at the coarsest mesh.')._p,
    para(
        'H1 and stress tell a third story, and it is the same '
        'ceiling-proximity effect Table 24d reports for the mean ratio, now '
        'visible in per-member reliability. At N = 9 the operator\'s H1 and '
        f'stress std/mean — {SUM[9]["H1_semi_rel"]["std_over_mean"]:.3f} and '
        f'{SUM[9]["stress_rel_L2"]["std_over_mean"]:.3f} — are '
        'indistinguishable from Q4\'s own '
        f'({Q4[9]["H1_semi_rel"]:.3f} and {Q4[9]["stress_rel_L2"]:.3f}). By '
        f'N = 33 they have grown to {SUM[33]["H1_semi_rel"]["std_over_mean"]:.3f} '
        f'and {SUM[33]["stress_rel_L2"]["std_over_mean"]:.3f}, while Q4\'s have '
        'not moved. The reading is the same one given above for the mean: '
        'when the operator sits close to the Q4 optimum in a norm it '
        'inherits some of Q4\'s member-independence there; as optimisation '
        'error comes to dominate at finer meshes, both the mean error and '
        'its spread across the family grow together. L2 and energy show no '
        'such coarse-mesh grace period — they are already far from Q4\'s '
        'consistency at N = 9, the same two norms Table 24d shows '
        'diverging outright in the mean.')._p,
])

doc.save(DST)
print(f'wrote {DST}')
