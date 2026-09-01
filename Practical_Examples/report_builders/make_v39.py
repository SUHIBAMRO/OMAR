"""v38 -> v39. Table 24c only reported L2 on the 16-member family; the same
JSON already carries H1 semi-norm, stress and energy on the same family, at
the same three meshes. Table 24d adds them, so the operator/Q4 comparison
that Table 24c makes in L2 is made in all four norms the rest of Section
8.11 uses.

Nothing new is measured here. `mms_family_fem_B1_neo_hookean.json` (built by
`record_mms_family.py`, already committed) has always had these four fields
per row; v38 simply did not read three of them into the document. This
builder reads the same file v38 read and asserts the same Q4 control before
writing, so it cannot silently drift from Table 24c's numbers.
"""
import json
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v38.docx', 'PFEM_Transolver_Report_v39.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

FAM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_family_fem_B1_neo_hookean.json')))

FROWS = {r['N']: r for r in FAM['rows']}
FNS = sorted(FROWS)
RATE = FAM['observed_rates_N9_to_N33_two_point']
PERI = FAM['per_interval_rates']
METRICS = ['L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel']
LABEL = {'L2_rel': 'L2', 'H1_semi_rel': 'H1 semi', 'stress_rel_L2': 'stress',
         'energy_rel': 'energy'}
assert FNS == [9, 17, 33], FNS

# Same control Table 24c already asserts -- repeated here because this
# builder reads the JSON independently and must not trust v38's prose.
assert abs(RATE['Q4']['L2_rel'] - 1.98) < 0.05, RATE['Q4']['L2_rel']
assert abs(RATE['Q4']['H1_semi_rel'] - 1.00) < 0.05, RATE['Q4']['H1_semi_rel']

RATIO = {N: {m: FROWS[N]['operator_over_Q4_on_the_family'][m] for m in METRICS}
         for N in FNS}
# Table 24c already published the L2 ratios 0.62x / 2.59x / 14.49x; this
# table must reproduce them exactly, since it is the same field of the same
# file, or the two tables in the same section would disagree.
assert abs(RATIO[9]['L2_rel'] - 0.62) < 0.01, RATIO[9]['L2_rel']
assert abs(RATIO[17]['L2_rel'] - 2.59) < 0.01, RATIO[17]['L2_rel']
assert abs(RATIO[33]['L2_rel'] - 14.49) < 0.01, RATIO[33]['L2_rel']

# The finding this table adds: H1 and stress stay near the ceiling (ratio
# close to 1) across both refinements, while L2 and energy do not.
for m in ('H1_semi_rel', 'stress_rel_L2'):
    assert RATIO[33][m] < 1.5, (m, RATIO[33][m])
for m in ('L2_rel', 'energy_rel'):
    assert RATIO[33][m] > 8, (m, RATIO[33][m])

print('H1/stress ratio at N=33:', RATIO[33]['H1_semi_rel'],
      RATIO[33]['stress_rel_L2'])
print('L2/energy ratio at N=33:', RATIO[33]['L2_rel'], RATIO[33]['energy_rel'])
print('rates:', {m: (RATE['Q4'][m], RATE['Q9'][m], RATE['operator'][m])
                  for m in METRICS})

# ======================================================================
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, __import__('copy').deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


# ======================================================================
# Section 8.11: Table 24d, the three norms Table 24c left out.
# ======================================================================
insert_after('Q4\'s own spread across the family is negligible', [
    para(
        'Table 24c reports L2 only. The same file records the H1 semi-norm, '
        'stress and energy on the identical 16-member family at the same '
        'three meshes, and the pattern is not the same in every norm. Table '
        '24d.')._p,
    para(
        'Table 24d. operator/Q4 on the 16-member family, all four norms. '
        'Values above 1 mean the operator is further from the manufactured '
        'solution than Q4 at the same mesh; the L2 column repeats Table 24c.')._p,
    new_table(
        ['N', 'L2', 'H1 semi', 'stress', 'energy'],
        [[str(N)] + [f'{RATIO[N][m]:.2f}×' for m in METRICS] for N in FNS])._tbl,
    para(
        'H1 and stress do not diverge the way L2 does. Their operator/Q4 '
        f'ratio moves {RATIO[9]["H1_semi_rel"]:.2f}×, '
        f'{RATIO[17]["H1_semi_rel"]:.2f}×, {RATIO[33]["H1_semi_rel"]:.2f}× in '
        f'H1 and {RATIO[9]["stress_rel_L2"]:.2f}×, '
        f'{RATIO[17]["stress_rel_L2"]:.2f}×, {RATIO[33]["stress_rel_L2"]:.2f}× '
        'in stress — staying within 45% of the ceiling through the same '
        f'refinement over which the L2 ratio reaches {RATIO[33]["L2_rel"]:.2f}×. '
        'The fitted rates over N = 9 to 33 confirm it: the operator is '
        f'{RATE["operator"]["H1_semi_rel"]:.2f} in H1 and '
        f'{RATE["operator"]["stress_rel_L2"]:.2f} in stress — positive, '
        'meaning the family-mean error still falls as the mesh refines, '
        'against Q4\'s theoretical 1.00 in both — while the same operator is '
        f'{RATE["operator"]["L2_rel"]:+.2f} in L2. Energy sits between the two '
        f'patterns: ratio {RATIO[9]["energy_rel"]:.2f}×, '
        f'{RATIO[17]["energy_rel"]:.2f}×, {RATIO[33]["energy_rel"]:.2f}×, '
        f'rate {RATE["operator"]["energy_rel"]:+.2f} — worse than H1 and '
        'stress, better than L2\'s outright divergence. This is the same '
        'inversion the single-member run reported in Table 24a — H1 and '
        'stress protected near the variational ceiling, L2 and energy not — '
        'now confirmed on 16 members rather than one, at both refinement '
        'intervals rather than assumed from three points.')._p,
    para(
        'Per interval, the direction does not reverse. From N = 9 to 17 the '
        f'operator\'s rate is {PERI["operator"]["N9_to_N17"]["H1_semi_rel"]:+.2f} '
        f'in H1 and {PERI["operator"]["N9_to_N17"]["stress_rel_L2"]:+.2f} in '
        f'stress; from 17 to 33 it is '
        f'{PERI["operator"]["N17_to_N33"]["H1_semi_rel"]:+.2f} and '
        f'{PERI["operator"]["N17_to_N33"]["stress_rel_L2"]:+.2f}. Both stay '
        'positive and both slow down on the finer half, which is consistent '
        'with the same optimisation-error explanation given above: refining '
        'shrinks Q4\'s own error fastest, so even a network making no further '
        'progress would see its ratio to Q4 grow, and the ratio growing '
        'slower in H1 and stress than in L2 says the network\'s absolute '
        'error in the derivative norms is falling too, just not as fast as '
        'Q4\'s.')._p,
])

doc.save(DST)
print(f'wrote {DST}')
