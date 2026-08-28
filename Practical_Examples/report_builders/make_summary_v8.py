"""Mirrors report v35 into the summary: the operator third of the MMS.

Summary section 11 currently ends on a "NOT here" paragraph whose first
clause is that the operator third does not exist. That paragraph is
REPLACED, exactly as make_v35.py replaces the report's "Two limits", so the
two documents cannot end up asserting opposite things.

Same JSONs, same assertions as make_v35.py.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
P9D = os.path.join(HERE, '..', 'omar_pfem', 'point9_results')
P9 = json.load(open(os.path.join(P9D, 'mms_B1_neo_hookean.json')))
OP = json.load(open(os.path.join(P9D, 'mms_operator_B1_neo_hookean.json')))
DEMO = json.load(open(os.path.join(P9D, 'operator_demo_N9_undertrained.json')))

KEYS = ('L2_rel', 'H1_semi_rel', 'stress_rel_L2', 'energy_rel')
N = OP['N']
op = OP['operator_on_the_reference_member']
ref = OP['fem_reference_same_mesh']
T = OP['training']
FV = OP['functional_verified']
MEMBER = OP['family']['scored_member']

for order in ('Q4', 'Q9'):
    row = next(r for r in P9['rows'] if r['order'] == order and r['N'] == N)
    for k in KEYS:
        assert abs(ref[order][k] - row[k]) < 1e-15 * max(1.0, abs(row[k])), \
            f'{order} {k}: the summary would contradict its own Table 22'
Q4_DOF, Q9_DOF = ref['Q4']['n_dof'], ref['Q9']['n_dof']
assert Q4_DOF == OP['n_dof']

RATIO = {k: op[k] / ref['Q4'][k] for k in KEYS}
assert RATIO['L2_rel'] > 1.0, 'operator/Q4 below 1.0 is a bug, not a result'
assert (RATIO['H1_semi_rel'] < RATIO['L2_rel']
        and RATIO['stress_rel_L2'] < RATIO['L2_rel']), \
    'the norm ordering is no longer inverted -- rewrite the reading'
d_op = DEMO['operator_on_the_reference_member']
d_q4 = DEMO['fem_reference_same_mesh']['Q4']
assert (d_op['H1_semi_rel'] / d_q4['H1_semi_rel']
        < d_op['L2_rel'] / d_q4['L2_rel'])
assert DEMO['N'] != N and DEMO['device'] != OP['device']

H = OP['history']
B_HALF = min(x['L2_rel'] for x in H[:len(H) // 2])
B_END = min(x['L2_rel'] for x in H)
GAIN = (1 - B_END / B_HALF) * 100
assert GAIN > 0
assert T['label_cost'].startswith('zero')

doc = Document('PFEM_Summary_Completed_Work.pre_v8.docx')
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


els = []
els.append(para(
    'The third leg, now measured. The report\'s trained operators cannot be '
    'pointed at a manufactured problem — no body-force term in the energy '
    'functional, no input channel to carry one — so a separate operator was '
    'trained for this study on the same Q4 mesh, minimising the same discrete '
    'potential energy and scored by the same error routine as the two '
    'solvers.')._p)
els.append(para(
    'Read the table with the ceiling in hand: the operator minimises the same '
    'functional over the same Q4 space that the Q4 solver solves, so the Q4 '
    'solution is the minimiser and the operator cannot beat it at this mesh. '
    'A ratio below one would be a defect, not an advance. The functional was '
    f'verified against the solver before training — Π(u_FEM) = {FV["Pi_at_u_FEM"]:.6f} '
    'is a true minimum, the excess grows quadratically with ratio '
    f'{FV["quadratic_excess_ratio"]:.3f}, and a work term mis-scaled by eight '
    f'moves that minimum to {FV["best_scale_W_divided_by_8"]}.')._p)
els.append(new_table(
    ['Method', 'DOF', 'L2', 'H1 semi', 'Stress', 'Energy'],
    [['Q4 (same mesh)', f'{Q4_DOF:,}'] + [f'{ref["Q4"][k]:.3e}' for k in KEYS],
     ['Q9 (same N)', f'{Q9_DOF:,}'] + [f'{ref["Q9"][k]:.3e}' for k in KEYS],
     ['Physics-informed operator', f'{Q4_DOF:,}'] +
     [f'{op[k]:.3e}' for k in KEYS]])._tbl)
els.append(para(
    f'Table 24. Three-way at N = {N}, α = {MEMBER["alpha"]}, β = {MEMBER["beta"]}. '
    f'The FEM rows are Table 22\'s N = {N} rows, not a second measurement. '
    f'Operator: {T["opt_steps"]:,} steps, {T["train_wall_clock_min"]} min on an '
    f'A100, {OP["family"]["ntrain"]}-member family, no labels.')._p)
els.append(para(
    f'operator / Q4 = {RATIO["L2_rel"]:.2f}× in L2, so the ceiling holds. The '
    f'finding is that the four norms disagree: {RATIO["H1_semi_rel"]:.2f}× in '
    f'H1 and {RATIO["stress_rel_L2"]:.2f}× in stress — effectively at the Q4 '
    f'optimum — against {RATIO["L2_rel"]:.2f}× in L2 and '
    f'{RATIO["energy_rel"]:.2f}× in energy. That inverts the usual ordering. '
    'The loss is built from the deformation gradient, so strain and stress are '
    'what it constrains hardest and the displacement is pinned only through '
    f'them; the same inversion appears in an independent N = {DEMO["N"]} CPU run '
    f'({d_op["H1_semi_rel"] / d_q4["H1_semi_rel"]:.2f}× against '
    f'{d_op["L2_rel"] / d_q4["L2_rel"]:.2f}×). For a physics-informed operator '
    'an L2 displacement error overstates how wrong the mechanics are.')._p)
els.append(para(
    f'What is left is optimisation error, not discretisation error: best '
    f'held-out L2 went {B_HALF:.3e} → {B_END:.3e} over the second half of '
    f'training, a further {GAIN:.0f}%, still falling slowly. Limits: one mesh, '
    'so the operator has no convergence rate of its own; its GPU training cost '
    'is not commensurable with the CPU Newton solves and is not forced onto a '
    'common axis; and the whole section rests on one geometry, one material '
    'and one scored member of the family.')._p)
els.append(para(
    'One decision to raise: the body-force-versus-homogeneous fork was settled '
    'here, not confirmed by Timon. A body-force-free exact solution on this '
    'domain is a homogeneous deformation, which Q4 reproduces to machine '
    'precision and which would separate nothing — but if he wants the study '
    'shaped differently, this is the choice to revisit.')._p)

victim = None
for p in doc.paragraphs:
    if p.text.strip().startswith('NOT here: the operator third'):
        victim = p
        break
assert victim is not None, \
    'the "NOT here" paragraph is gone -- has v8 already been applied?'
target = victim._p
for el in els:
    target.addnext(el)
    target = el
victim._p.getparent().remove(victim._p)

doc.save('PFEM_Summary_Completed_Work.docx')
print(f'replaced the "NOT here" paragraph with {len(els)} elements; '
      f'op/Q4 L2 {RATIO["L2_rel"]:.2f}x, H1 {RATIO["H1_semi_rel"]:.2f}x')
