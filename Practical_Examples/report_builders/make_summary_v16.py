"""Mirrors report v43 into the summary: Table 18c, B2 x Neo-Hookean's
Pareto and the training-resolution anchoring finding. Same JSON and
assertions as make_v43.py.

Run from the directory holding PFEM_Summary_Completed_Work.docx; copies the
current file to .pre_v16.docx first, then writes back to
PFEM_Summary_Completed_Work.docx.
"""
import copy
import json
import os
import shutil

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

B2NH = json.load(open(os.path.join(PF, 'point2_results', 'pareto_B2_neo_hookean.json')))

NS = [13, 17, 21, 25, 29, 33, 37, 41, 49]
rows = {r['N']: r for r in B2NH['rows']}
speedup = {N: rows[N]['fem_ms_per_sample'] / rows[N]['operator_ms_per_sample']
           for N in NS}
anchor = {int(k): v for k, v in B2NH['shape']['training_resolution_anchoring']['anchor_ratios'].items()}
b1_ratio_21 = B2NH['shape']['training_resolution_anchoring'][
    'b1_neo_hookean_ratio_at_N21_for_comparison']
assert anchor[21] > 3 and anchor[33] > 3

SRC = 'PFEM_Summary_Completed_Work.docx'
shutil.copy2(SRC, 'PFEM_Summary_Completed_Work.pre_v16.docx')
doc = Document('PFEM_Summary_Completed_Work.pre_v16.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))
ORIGINAL = list(doc.paragraphs)


def new_table(header, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(data_rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, f'{len(hits)} matches for {prefix!r}'
    return hits[0]


def insert_after(prefix, els):
    target = find(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


HEADER = ['N', 'Nodes', 'FEM rel. L2 (%)', 'FEM cost (s)',
          'Operator rel. L2 (%)', 'Operator cost (ms)', 'Speed-up']
table_data = [[str(N), f'{rows[N]["n_nodes"]:,}', f'{rows[N]["fem_rel_L2"]*100:.3f}',
               f'{rows[N]["fem_ms_per_sample"]/1000:.1f}',
               f'{rows[N]["operator_rel_L2"]*100:.2f}',
               f'{rows[N]["operator_ms_per_sample"]:.3f}',
               f'{speedup[N]:,.0f}×'] for N in NS]

insert_after('Table 18b. Arruda-Boyce', [
    para(
        f'Table 18c. B2 × Neo-Hookean (first of three B2 materials to '
        f'finish). Speed-up {min(speedup.values()):,.0f}×–'
        f'{max(speedup.values()):,.0f}×.')._p,
    new_table(HEADER, table_data)._tbl,
    para(
        f'Unlike every B1 curve, this one is not smooth in N: sharp local '
        f'minima at N=21 and N=33 — the two resolutions this checkpoint '
        f'was jointly trained on — {anchor[21]:.2f}× and {anchor[33]:.2f}× '
        f'lower than each point\'s neighbours. Table 18\'s B1 × '
        f'Neo-Hookean curve, from a checkpoint trained at N=21 only, shows '
        f'no such dip ({b1_ratio_21:.2f}×, flat) — candidate explanation '
        f'(joint two-resolution training leaves visible anchor points), '
        f'not established, since geometry and protocol differ at once '
        f'between the two checkpoints compared. Mooney-Rivlin and '
        f'Arruda-Boyce\'s B2 sweeps were still running as this was '
        f'written.')._p,
])

doc.save(SRC)
print(f'summary v16: Table 18c added -- anchor ratios {anchor[21]:.2f}x/'
      f'{anchor[33]:.2f}x, B1 comparison {b1_ratio_21:.2f}x')
