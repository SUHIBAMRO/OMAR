"""Mirrors report v36 into the summary.

Three edits, matching make_v36.py's three and reading the same JSONs so the
documents cannot disagree:

  * section 7 / Table 12 -> three materials, seven resolutions, corrected
    training description;
  * section 6 -> the tested normalization mitigation;
  * the Table 20 line -> the CG finding, which changes what that number means.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')
P8 = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_scaling_B1_neo_hookean.json')))
MIT = json.load(open(os.path.join(PF, 'point6_results',
                                  'ood_mitigation_B1_neo_hookean.json')))
ZS = {m: json.load(open(os.path.join(PF, 'point7a_results',
                                     f'zeroshot_B1_{m}.json')))
      for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')}

ROWS8 = sorted(P8['rows'], key=lambda r: r['N'])
CGCAP = 2000
conv = [r for r in ROWS8 if r['stats']['cg_failures'] == 0]
capped = [r for r in ROWS8 if r['stats']['cg_failures'] > 0]
assert len(conv) == 3 and len(capped) == 5
for r in capped:
    assert r['stats']['cg_failures'] == r['stats']['newton_iters_total']
K = [r['stats']['cg_iters_total'] / r['stats']['newton_iters_total'] / r['N']
     for r in conv]
KBAR = sum(K) / len(K)
FRAC1401 = CGCAP / (KBAR * 1401)
NSCOST = {r['N']: r['stats'].get('t_cg_s', r['solve_s_in_source'])
          / r['stats']['cg_iters_total'] * 1e9 / r['n_dof'] for r in ROWS8}
BIG = [NSCOST[n] for n in (701, 1001, 1401)]
assert max(BIG) / min(BIG) - 1 < 0.01

NS_LIST = ZS['neo_hookean']['test_resolutions']
TRAIN_RES = ZS['neo_hookean']['protocol']['train_resolutions']
for d in ZS.values():
    assert d['test_resolutions'] == NS_LIST
    assert d['protocol']['train_resolutions'] == TRAIN_RES
ERR = {m: [r['mean_rel_L2_vs_fine_reference'] for r in d['rows']]
       for m, d in ZS.items()}
SHAPE = {m: d['shape'] for m, d in ZS.items()}
assert SHAPE['mooney_rivlin']['monotone_decreasing_in_N']
assert not SHAPE['neo_hookean']['monotone_decreasing_in_N']
assert not SHAPE['arruda_boyce']['monotone_decreasing_in_N']

IND = MIT['in_distribution']
MROWS = MIT['rows']
worse = [r for r in MROWS if r['normalized'] >= r['raw']]
better = [r for r in MROWS if r['normalized'] < r['raw']]
assert len(worse) > len(better)
mat = [r for r in MROWS if r['factor'] == 'material']
TURN_K = mat[[r['normalized'] for r in mat].index(
    max(r['normalized'] for r in mat))]['k']
BUD = MIT['training_budget_resolved']
MORE = (BUD['normalized']['opt_steps_at_end']
        / BUD['baseline']['opt_steps_at_end'] - 1) * 100

doc = Document('PFEM_Summary_Completed_Work.pre_v9.docx')
ref = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise AssertionError(f'not found: {prefix!r}')


def replace(prefix, els):
    v = find(prefix)
    t = v._p
    for el in els:
        t.addnext(el)
        t = el
    v._p.getparent().remove(v._p)


def after(prefix, els):
    t = find(prefix)._p
    for el in els:
        t.addnext(el)
        t = el


# ---- Table 20 line: what the CG counters change ----------------------
replace('Table 20. Largest case:', [
    para('Table 20. Largest case: 3,925,602 DOF in 11.0 h using 3,280 MB of '
         '80 GB (~4%). Memory is not the constraint, and the memory model '
         'held out of sample to 2.4%.')._p,
    para(
        f'The timings in that table need a caveat the report now states in '
        f'full. The solver\'s own counters show CG did not converge at N ≥ '
        f'401: it hit its {CGCAP:,}-iteration cap on EVERY Newton step, so '
        f'those wall clocks are the cost of a truncated budget rather than of '
        f'a converged solve. Accuracy is unaffected — Newton\'s test is on the '
        f'absolute residual and its counts stay far below their limit — but '
        f'the cost analysis changes. The three resolutions where CG did '
        f'converge fix the real requirement at {KBAR:.2f} × N iterations per '
        f'Newton solve, the textbook O(1/h) rate, which means N=1401 completed '
        f'only {FRAC1401 * 100:.0f}% of the iterations it needed.')._p,
    para(
        f'What survives cleanly, and is the better result: dividing CG time by '
        f'iteration count gives {", ".join(f"{NSCOST[n]:.0f}" for n in (701, 1001, 1401))} '
        f'nanoseconds per iteration per degree of freedom at N = 701, 1001 and '
        f'1401 — flat to {(max(BIG) / min(BIG) - 1) * 100:.1f}% across a '
        f'fourfold size change. The matrix-free product is O(DOF) in time, '
        f'measured rather than fitted. The previously reported DOF^1.54 '
        f'exponent was not measuring the solver\'s scaling: those runs each '
        f'ran a constant {CGCAP:,} iterations per Newton step and what grew '
        f'was the Newton count, which grows only because a truncated CG '
        f'returns an inexact direction.')._p,
])

# ---- section 6: the mitigation, tested -------------------------------
els = [para(
    f'Mitigation, tested rather than proposed. Section 8.6 named normalizing '
    f'the material channel as the cheapest candidate fix. It was run: the '
    f'identical protocol on the identical dataset with input normalization as '
    f'the only change, then both sweeps repeated on a shared cache of FEM '
    f'references. It does not work.')._p]
els.append(new_table(
    ['Factor', 'k', 'Raw', 'Normalized', 'Raw ×', 'Norm ×'],
    [[r['factor'], f"{r['k']:.1f}", f"{r['raw']:.4f}", f"{r['normalized']:.4f}",
      f"{r['raw'] / IND['raw']:.2f}×",
      f"{r['normalized'] / IND['normalized']:.2f}×"] for r in MROWS])._tbl)
els.append(para(
    f'Table 19a. Both models over the same sweep. Each "×" divides by that '
    f'model\'s own in-distribution error ({IND["raw"]:.4f} raw, '
    f'{IND["normalized"]:.4f} normalized).')._p)
els.append(para(
    f'Why it is a negative result and not a win: normalization costs '
    f'{IND["change_pct"]:.1f}% in distribution — confirmed on two independent '
    f'metrics and despite {MORE:.0f}% MORE optimizer steps — and makes the '
    f'absolute error worse in {len(worse)} of {len(MROWS)} shifted cells, '
    f'every one at or below 1.5σ among them. Its one striking gain, the '
    f'material ratio at 3σ, sits on a curve that peaks at k = {TURN_K:.1f} and '
    f'then FALLS: an error that stops growing as the shift grows is a '
    f'prediction collapsing toward something input-independent, not '
    f'extrapolation. The loading rows are the control and stay flat for both '
    f'models, so the sensitivity did not move. The §8.6 mechanism stands — '
    f'standardizing is affine, so a shifted E is still outside the trained '
    f'range. The untested candidate that remains is predicting a '
    f'stiffness-scaled quantity instead of the displacement.')._p)
after('6. Out-of-distribution generalization', els)

# ---- section 7: Table 12, three materials ----------------------------
old_tbl = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Table 12.'):
        el = p._p.getprevious()
        while el is not None and el.tag != qn('w:tbl'):
            el = el.getprevious()
        old_tbl = el
        break
assert old_tbl is not None, 'summary Table 12 not found'
t12 = new_table(
    ['N', 'Neo-Hookean', 'Mooney-Rivlin', 'Arruda-Boyce'],
    [[str(N)] + [f'{ERR[m][i]:.4f}' for m in
                 ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')]
     for i, N in enumerate(NS_LIST)])
old_tbl.addnext(t12._tbl)
old_tbl.getparent().remove(old_tbl)

nh, mr, ab = SHAPE['neo_hookean'], SHAPE['mooney_rivlin'], SHAPE['arruda_boyce']
replace('Table 12. Zero-shot resolution invariance', [
    para(
        f'Table 12 (revised). Zero-shot resolution invariance, all three B1 '
        f'materials. One checkpoint per material, trained JOINTLY at N = '
        f'{TRAIN_RES[0]} and {TRAIN_RES[1]}, evaluated with no retraining at '
        f'seven unseen resolutions against a common N=101 reference, twenty '
        f'realizations each. An earlier revision described the Neo-Hookean '
        f'checkpoint as trained at N=21 alone; it was not, and the correction '
        f'does not move any of its numbers.')._p,
    para(
        f'All three stay between '
        f'{min(min(v) for v in ERR.values()) * 100:.1f}% and '
        f'{max(max(v) for v in ERR.values()) * 100:.1f}% across every unseen '
        f'resolution, with no retraining. But reading down the columns shows '
        f'what one material could not: two of three reach their best near the '
        f'training range and then get WORSE on the finest meshes — '
        f'Neo-Hookean bottoms at N = {nh["best_resolution_N"]} and rises '
        f'{nh["rise_after_minimum_pct"]:.1f}% by N = {NS_LIST[-1]}, '
        f'Arruda-Boyce bottoms at N = {ab["best_resolution_N"]} and rises '
        f'{ab["rise_after_minimum_pct"]:.1f}%, while Mooney-Rivlin keeps '
        f'improving to {mr["best_error"]:.4f}. Zero-shot transfer to much '
        f'finer meshes is not free and how much it costs is '
        f'material-dependent.')._p,
    para(
        'The three B2 cases are NOT here. Their sample caches carried an '
        'applied load overstated by a mesh-dependent factor — about 13× at '
        'N=21 against 21× at N=33 — which for a study that measures transfer '
        'across meshes invalidates them outright. The caches have been '
        'repaired and checked (one fixed pressure field now assembles to '
        'within 0.008% across the two meshes) and the affected models '
        'deleted; the cases are being regenerated.')._p,
])

doc.save('PFEM_Summary_Completed_Work.docx')
print('summary v9: Table 20 caveat, Table 19a, Table 12 with three materials')
