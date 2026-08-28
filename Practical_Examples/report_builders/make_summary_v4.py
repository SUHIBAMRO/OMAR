"""Adds the two things the summary is missing: the side-by-side break-even
(report Table 10d) and the Pareto comparison (report Table 18).

Reads the pre-v4 copy and writes the live file, so re-running replaces these
sections rather than appending duplicates -- the same pattern make_summary_v3
uses.

Register is the summary's own, set by the user: results only, one line of
lead per table, caption underneath, no discussion.
"""
import copy

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

import pareto_table as PT

SRC = 'PFEM_Summary_Completed_Work.pre_v4.docx'
DST = 'PFEM_Summary_Completed_Work.docx'
doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
F = PT.facts()

CASES = ['B1 × Neo-Hookean', 'B1 × Mooney-Rivlin', 'B1 × Arruda-Boyce',
         'B2 × Neo-Hookean', 'B2 × Mooney-Rivlin', 'B2 × Arruda-Boyce']
# recomputed the same way make_v30.py does, from Tables 4a and 7
TRAIN_S = [2873.8, 2785.3, 2855.6, 32244.0, 34164.1, 24846.8]
CPU_FEM_S = [25.432, 53.735, 52.542, 25.909, 61.712, 60.285]
INFER_MS = [4.586, 4.693, 4.745, 4.809, 4.908, 4.984]
GPU_BE = [[1745, 6021, 7543, 8112], [1363, 5663, 7179, 7751],
          [1133, 5441, 6956, 7554], [19410, 67391, 84627, 90990],
          [17033, 69404, 87884, 95038], [9530, 46993, 60490, 65698]]
cpu_be = [t / (f - ms / 1000.0) for t, f, ms in zip(TRAIN_S, CPU_FEM_S, INFER_MS)]
assert 51 <= min(cpu_be) <= 53 and 1240 <= max(cpu_be) <= 1250, 'CPU break-even drifted'


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style=None):
    p = doc.add_paragraph() if style is None else doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def last_element_of_section(heading_prefix, heading_style):
    """The last body element belonging to the section that starts with
    `heading_prefix`, i.e. everything up to but not including the next
    heading at the same level."""
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(heading_prefix):
            anchor = p._p
            break
    assert anchor is not None, f'{heading_prefix!r} not found'
    cur = last = anchor
    while True:
        nxt = cur.getnext()
        if nxt is None:
            break
        if nxt.tag == qn('w:p'):
            pp = Paragraph(nxt, doc)
            if pp.style is not None and pp.style == heading_style and pp.text.strip():
                break
        last = nxt
        cur = nxt
    return last


heading_style = None
for p in doc.paragraphs:
    if p.text.strip().startswith('5. GPU-native finite-element solver'):
        heading_style = p.style
        break
assert heading_style is not None, 'section 5 heading not found'

# ---------------------------------------------- Table 10d, into section 5
els = [para('Break-even against both baselines together. CPU column from '
            'Tables 4a and 7; GPU columns are Table 10c.')._p,
       new_table(['Case', 'vs. CPU FEM', 'vs. GPU FEM, bs=1', 'bs=8', 'bs=32', 'bs=128'],
                 [[c, f'{be:,.0f}'] + [f'{v:,}' for v in g]
                  for c, be, g in zip(CASES, cpu_be, GPU_BE)])._tbl,
       para('Table 10d. Break-even in new problem instances, both baselines. '
            f'{min(cpu_be):.0f}–{max(cpu_be):.0f} against CPU; '
            f'{min(v for g in GPU_BE for v in g):,}–'
            f'{max(v for g in GPU_BE for v in g):,} against GPU, of which '
            f'{min(g[0] for g in GPU_BE):,}–{max(g[0] for g in GPU_BE):,} at '
            'batch size 1, the deployment case.')._p]
target = last_element_of_section('5. GPU-native finite-element solver', heading_style)
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements into section 5')

# ------------------------------------- the Pareto, as a new section 9
els = []
h = doc.add_paragraph(style=heading_style)
h.add_run('9. Accuracy/cost Pareto: operator vs. FEM (B1 × Neo-Hookean)')
els.append(h._p)
els.append(para(
    f'Both sides on the same {F["n_samples"]} instances, same N=101 fine '
    'reference, same device, batch size 1. FEM cost is the CPU reference '
    'solver. Error is the combined relative L2 over both components — NOT '
    'Table 12\'s per-component average, and drawn from a different seed range, '
    'so this column is not comparable with Table 12.')._p)
els.append(new_table(PT.HEAD, PT.rows())._tbl)
els.append(para(
    'Table 18. Accuracy and cost at nine resolutions. FEM at its coarsest '
    f'(N=13) is {F["coarsest_fem"]:.3f}%, {F["advantage"]:.1f}× more accurate '
    f'than the operator at its best ({F["op_lo"]:.2f}% at N={F["op_best_N"]}); '
    'the two branches do not overlap. Operator cost is flat in mesh size '
    f'({F["op_ms_lo"]:.3f}–{F["op_ms_hi"]:.3f} ms over a 14× node increase) '
    f'while FEM grows superlinearly, so the speed-up climbs {F["speed_lo"]:,.0f}× '
    f'→ {F["speed_hi"]:,.0f}×. Run twice: errors identical to every digit, '
    'timings 2.9× apart between Colab instances; the run consistent with '
    'Table 10a is reported. One of six cases.')._p)

target = last_element_of_section('8. Error in physically important quantities',
                                 heading_style)
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements as section 9')

doc.save(DST)
print(f'{DST} now has {len(Document(DST).tables)} tables')
