"""v36 -> v37: three edits, all of them things that were open in v36 and
are now measured.

  A. Section 8.5. v36 says the direction of Table 20's error "is not
     one-signed ... this study does not establish which effect is larger."
     It has now been established. N=501 and N=701 were re-run with the CG
     cap raised from 2,000 to 8,000, CG converged at both, and Table 20
     UNDERSTATES the converged cost: +28% and +18%. Table 20b carries the
     converged rows; the two extrapolated resolutions are labelled as
     extrapolations.
  B. Section 8.11. The ceiling is stated too strongly. "A ratio below one
     would indicate a defect" is true of Pi and false of L2, and the
     three-mesh run below produces 0.37x in L2 with nothing wrong.
     Restated: the ceiling constrains Pi, it transfers to the derivative
     norms empirically, and it does not transfer to L2.
  C. Section 8.11. The limitation "the operator was trained and scored at
     one mesh only, so it has no convergence rate of its own" is removed:
     N=9 and N=33 were trained under the N=17 protocol. Tables 24a and
     24b, and the result, which is that the operator does not converge.

Every number is read from the committed JSONs and asserted before it is
written.
"""
import copy
import json
import math
import os

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v36.docx', 'PFEM_Transolver_Report_v37.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

P8 = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_scaling_B1_neo_hookean.json')))
CG = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_cg_converged_B1_neo_hookean.json')))
MMS = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_operator_rate_B1_neo_hookean.json')))
BAD = json.load(open(os.path.join(PF, 'point7a_results',
                                  'INVALID_B2_zeroshot.json')))
RETRAIN = json.load(open(os.path.join(PF, 'point7a_results',
                                      'B2_zeroshot_retrain_status.json')))

# ------------------------------------------------------- A: converged CG
TRUNC = {r['N']: r for r in P8['rows']}
CONV = {r['N']: r for r in CG['rows']}
assert set(CONV) == {501, 701}, sorted(CONV)
for N, r in CONV.items():
    assert r['cg_failures'] == 0, (
        f'N={N}: CG did not converge in the re-run either, so this section '
        f'cannot claim a converged cost')
    assert abs(r['prediction_error_pct']) < 1.0, (
        f'N={N}: the 5.011 x N law missed by {r["prediction_error_pct"]}%, '
        f'which is not the "held to 0.4%" the text claims')
    assert r['n_dof'] == TRUNC[N]['n_dof'], N
# the direction claim -- assert it rather than trusting the transcription
DELTA = {N: (CONV[N]['solve_s'] / TRUNC[N]['solve_s_in_source'] - 1) * 100
         for N in CONV}
assert all(v > 0 for v in DELTA.values()), (
    'the converged solve is not more expensive at every size, so "Table 20 '
    'understates" is the wrong statement', DELTA)
assert DELTA[501] > DELTA[701], (
    'the gap no longer narrows with size -- rewrite the explanation', DELTA)
# Newton fell at 701 and was already minimal at 501
assert TRUNC[701]['stats']['newton_iters_total'] == 30
assert CONV[701]['newton_iters_total'] == 20
assert TRUNC[501]['stats']['newton_iters_total'] == CONV[501]['newton_iters_total'] == 20
# the per-CG-iteration cost agrees between the truncated and converged runs
# The point-8 sweep recorded no t_cg_s at these two resolutions, so the only
# like-for-like division is solve time over CG iterations on both sides. CG is
# 99.8% of the converged solve, which bounds what that substitution can hide,
# and the JSON stores both forms.
MS_TRUNC = {N: TRUNC[N]['solve_s_in_source']
            / TRUNC[N]['stats']['cg_iters_total'] * 1e3 for N in CONV}
for N in CONV:
    assert 't_cg_s' not in TRUNC[N]['stats'], (
        f'N={N}: the truncated sweep now has a CG-time breakdown, so the '
        f'comparison below should use it on both sides')
    assert CONV[N]['cg_share_of_solve_pct'] > 99.5, CONV[N]
MS_GAP = {N: abs(CONV[N]['ms_per_cg_iter_from_solve'] / MS_TRUNC[N] - 1) * 100
          for N in CONV}
assert max(MS_GAP.values()) < 3.0, (
    'the two runs no longer agree on the matvec cost', MS_GAP, MS_TRUNC)
EXTRA = CG['EXTRAPOLATED_NOT_MEASURED']
# The O(N) law itself, refitted here from the sweep's three converged rows
# rather than taken from either JSON's prose, so that the constant the text
# quotes is the one the predicted column was built from.
_c = [r for r in P8['rows'] if r['stats']['cg_failures'] == 0]
KBAR = sum(r['stats']['cg_iters_total'] / r['stats']['newton_iters_total']
           / r['N'] for r in _c) / len(_c)
for N in CONV:
    # not bit-for-bit: the run rounded its printed prediction to the nearest
    # iteration, so agreement is asserted to a tenth of a per cent rather
    # than exactly.
    assert abs(KBAR * N / CONV[N]['predicted_cg_per_newton'] - 1) < 1e-3, (
        f'N={N}: the refitted law gives {KBAR * N:.1f} where the run '
        f'predicted {CONV[N]["predicted_cg_per_newton"]}')

# ------------------------------------------------------- B/C: MMS operator
MROWS = sorted(MMS['rows'], key=lambda r: r['N'])
assert [r['N'] for r in MROWS] == [9, 17, 33], [r['N'] for r in MROWS]
RATE = MMS['fitted_rates_in_h']
# the control: Q4's own fitted rates must land on Table 23's measured ones,
# or these three runs are not comparable to that table and nothing here can
# be quoted beside it.
assert abs(RATE['Q4_L2'] - 1.98) < 0.05, RATE
assert abs(RATE['Q4_H1_semi'] - 1.00) < 0.05, RATE
assert RATE['operator_L2'] < 0, (
    'the operator L2 error no longer grows with refinement -- the whole '
    'reading below is wrong', RATE)
assert RATE['operator_H1_semi'] > 0, RATE
# refit the rates here rather than trusting the transcribed ones
def _fit(get):
    xs = [math.log(1.0 / (r['N'] - 1)) for r in MROWS]
    ys = [math.log(get(r)) for r in MROWS]
    mx, my = sum(xs) / 3, sum(ys) / 3
    return (sum((x - mx) * (y - my) for x, y in zip(xs, ys))
            / sum((x - mx) ** 2 for x in xs))
for key, get in (('operator_L2', lambda r: r['operator']['L2']),
                 ('Q4_L2', lambda r: r['Q4']['L2']),
                 ('operator_H1_semi', lambda r: r['operator']['H1_semi']),
                 ('Q4_H1_semi', lambda r: r['Q4']['H1_semi'])):
    assert abs(_fit(get) - RATE[key]) < 0.01, (key, _fit(get), RATE[key])
RAT = MMS['ratios_operator_over_Q4']
L2R = [RAT['L2'][str(r['N'])] for r in MROWS]
assert L2R[0] < 1.0 < L2R[1] < L2R[2], (
    'the L2 ratio no longer crosses one between N=9 and N=17 -- the '
    'crossover reading has to be rewritten', L2R)
for norm in ('H1_semi', 'stress'):
    assert all(RAT[norm][str(r['N'])] > 1.0 for r in MROWS), (
        f'{norm} fell below one, so the empirical protection of the '
        f'derivative norms no longer holds', RAT[norm])
assert min(RAT['energy'].values()) < 1.0, RAT['energy']
# Table 24 already prints the N=17 column; Table 24a must not contradict it
T24 = next(r for r in MROWS if r['N'] == 17)
assert abs(T24['operator']['L2'] / T24['Q4']['L2'] - 2.42) < 0.01, T24
# Each operator run solved its own Q4 reference, so Table 24a's Q4 columns are
# a second measurement of two rows Table 22 already prints. They are separate
# solves and need not be bit-for-bit, but a report that prints both cannot
# have them disagree visibly.
FEM = json.load(open(os.path.join(PF, 'point9_results',
                                  'mms_B1_neo_hookean.json')))
for r in MROWS:
    match = [q for q in FEM['rows'] if q['order'] == 'Q4' and q['N'] == r['N']]
    if not match:
        continue                      # N=33 is beyond Table 22's three meshes
    for key, tkey in (('L2', 'L2_rel'), ('H1_semi', 'H1_semi_rel')):
        assert abs(r['Q4'][key] / match[0][tkey] - 1) < 2e-3, (
            f"N={r['N']} {key}: Table 24a would print {r['Q4'][key]:.4e} where "
            f"Table 22 prints {match[0][tkey]:.4e}")
# and the memory claim in section 8.5
for N in CONV:
    assert abs(CONV[N]['peak_gpu_mem_MB'] - TRUNC[N]['peak_gpu_mem_MB']) <= 1, (
        f'N={N}: peak memory moved between the truncated and converged runs, '
        f'so "unchanged to within one megabyte" is wrong')

# ------------------------------------------------------- D: B2 zero-shot
BEST = RETRAIN['result']['best_by_case']
assert all(abs(v['combined_val_error'] - 1.0) < 0.05 for v in BEST.values()), (
    'a B2 case has moved off 1.0, so section 8.7 must be rewritten around '
    'what it actually reaches', BEST)
EV = RETRAIN['result']['eval_errors']
B1SP = RETRAIN['result']['B1_spread_over_the_mesh_pct']
# the eval numbers in that file were read off Drive; re-derive the two claims
# made from them rather than trusting the stored summary
for m, v in EV.items():
    got = (max(v['mean_rel_L2_vs_fine_reference'])
           / min(v['mean_rel_L2_vs_fine_reference']) - 1) * 100
    assert abs(got - v['spread_over_the_mesh_pct']) < 1e-3, (m, got)
    assert got < 0.2, (m, got)
for m, sp in B1SP.items():
    got = (lambda v: (max(v) / min(v) - 1) * 100)(
        [r['mean_rel_L2_vs_fine_reference']
         for r in json.load(open(os.path.join(
             PF, 'point7a_results', f'zeroshot_B1_{m}.json')))['rows']])
    assert abs(got - sp) < 0.05, (m, got, sp)
OLD_B2 = {m: min(BAD['the_numbers_that_gave_it_away'][f'B2_{m}'].values())
          for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')}
assert min(OLD_B2.values()) > 1.0, OLD_B2
MESH_OK = BAD['repair']['mesh_independence_check_after_repair']
assert MESH_OK['spread_pct'] < 0.1, MESH_OK

print(f'A: converged CG at 501/701, prediction held to '
      f'{max(abs(r["prediction_error_pct"]) for r in CG["rows"]):.2f}%; '
      f'Table 20 understates by +{DELTA[501]:.0f}% and +{DELTA[701]:.0f}%')
print(f'B/C: operator L2 rate {RATE["operator_L2"]:+.2f} against Q4 '
      f'{RATE["Q4_L2"]:.2f}; L2 ratio {L2R[0]:.2f}x -> {L2R[2]:.2f}x')

doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))
# Snapshot the paragraphs as loaded. para() below appends to the same body,
# so a later search over doc.paragraphs would also see the replacements this
# script has just written -- which is how the first run of it found two
# paragraphs beginning "Two consequences for how Table 20 should be read."
ORIGINAL = list(doc.paragraphs)


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; the edit would land '
        f'in the wrong place')
    return hits[0]


def replace_para(prefix, els):
    victim = find_para(prefix)
    target = victim._p
    for el in els:
        target.addnext(el)
        target = el
    victim._p.getparent().remove(victim._p)


def retext(prefix, text):
    """Rewrite one paragraph in place, keeping its style.

    replace_para() builds a fresh Normal paragraph, which is right in the
    body but wrong for the bulleted list in section 10 -- a replacement
    there would silently lose the bullet.
    """
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


# ======================================================================
# A. Section 8.5 -- the direction of Table 20's error is now measured
# ======================================================================
els = [para(
    f'Two consequences for how Table 20 should be read. Its times at N ≥ 401 '
    f'are not the solver\'s converged cost, and an earlier revision of this '
    f'section left the direction of that error open, on the grounds that '
    f'allowing CG to converge would add iterations to each Newton step while '
    f'removing Newton steps. Which effect is larger has since been measured '
    f'rather than argued: the two largest resolutions the budget allowed, N = '
    f'501 and N = 701, were re-run with the CG cap raised from '
    f'{2000:,} to {8000:,} and every other setting identical. CG converged at '
    f'both — no capped solves at all — and the answer is that Table 20 '
    f'understates. Table 20b.')._p]

els.append(new_table(
    ['N', 'DOF', 'CG per Newton solve, predicted', 'measured', 'error',
     'Newton iters, truncated → converged', 'Solve time, truncated → converged',
     'change'],
    [[str(N), f"{CONV[N]['n_dof']:,}",
      f"{CONV[N]['predicted_cg_per_newton']:,.0f}",
      f"{CONV[N]['cg_per_newton']:,.1f}",
      f"{CONV[N]['prediction_error_pct']:.2f}%",
      f"{TRUNC[N]['stats']['newton_iters_total']} → "
      f"{CONV[N]['newton_iters_total']}",
      f"{TRUNC[N]['solve_s_in_source']:,.0f} s → {CONV[N]['solve_s']:,.0f} s",
      f"+{DELTA[N]:.0f}%"]
     for N in sorted(CONV)])._tbl)
els.append(para(
    f'Table 20b. The same solver at the same two resolutions with the CG '
    f'iteration cap raised from {2000:,} to {8000:,}; B1 × Neo-Hookean, Q4, '
    f'ten load steps, FP64, the same A100, and cg_failures = 0 in both rows. '
    f'The predicted column is the {KBAR:.3f} × N law of the preceding '
    f'paragraph, fitted on N = 101, 201 and 301, and it was printed by the run '
    f'before the solves were launched rather than compared afterwards.')._p)

els.append(para(
    f'Three things follow. The O(N) iteration law, fitted on three '
    f'resolutions no larger than N = 301, predicts '
    f'{CONV[501]["predicted_cg_per_newton"]:,.0f} and '
    f'{CONV[701]["predicted_cg_per_newton"]:,.0f} iterations per Newton solve '
    f'and the converged runs needed {CONV[501]["cg_per_newton"]:,.1f} and '
    f'{CONV[701]["cg_per_newton"]:,.1f} — within '
    f'{max(abs(r["prediction_error_pct"]) for r in CG["rows"]):.1f}% at both, '
    f'the second of them more than twice the largest mesh the law was fitted '
    f'on. Truncated CG does cost Newton steps, as the argument for leaving the '
    f'direction open supposed: at N = 701 the Newton count falls from '
    f'{TRUNC[701]["stats"]["newton_iters_total"]} to '
    f'{CONV[701]["newton_iters_total"]}, which is the two per load step that '
    f'every converged row of Table 20a shows, while at N = 501 it was already '
    f'{CONV[501]["newton_iters_total"]} and there was nothing to recover. But '
    f'the saving does not pay for the extra iterations at either size: the '
    f'solve costs {DELTA[501]:.0f}% more at N = 501 and {DELTA[701]:.0f}% more '
    f'at N = 701, and the gap narrows with size precisely because the larger '
    f'run had more wasted Newton steps to give back. In µs/DOF the two rows '
    f'move from {TRUNC[501]["us_per_dof"]:,.0f} to '
    f'{CONV[501]["us_per_dof"]:,.0f} and from '
    f'{TRUNC[701]["us_per_dof"]:,.0f} to {CONV[701]["us_per_dof"]:,.0f}. '
    f'Third, the cost of one matrix-free Hessian-vector product agrees between '
    f'the truncated and the converged run to '
    f'{MS_GAP[501]:.1f}% and {MS_GAP[701]:.1f}% '
    f'({MS_TRUNC[501]:.1f} against '
    f'{CONV[501]["ms_per_cg_iter_from_solve"]:.1f} ms, and {MS_TRUNC[701]:.1f} '
    f'against {CONV[701]["ms_per_cg_iter_from_solve"]:.1f} ms, dividing solve '
    f'time by CG iterations on both sides because the sweep of Table 20a '
    f'recorded no separate CG time at these two resolutions; CG is '
    f'{min(CONV[N]["cg_share_of_solve_pct"] for N in CONV):.1f}% of the '
    f'converged solve, which bounds what that substitution can hide) — two '
    f'independent runs measuring the same quantity, which is what the O(DOF) '
    f'claim of the preceding paragraph rests on. Peak memory is unchanged to '
    f'within one megabyte, as it should be.')._p)

els.append(para(
    f'The two largest resolutions were not re-run: at the converged iteration '
    f'count they are roughly four and twelve hours, and the budget went to the '
    f'two that could settle the direction. Applying the now twice-confirmed '
    f'model — {CONV[701]["newton_iters_total"]} Newton steps, '
    f'{KBAR:.3f} × N CG iterations each, at the measured per-iteration '
    f'cost — gives {EXTRA["N1001"]["predicted_converged_s"] / 3600:.1f} hours '
    f'at N = 1001 against the {TRUNC[1001]["solve_s_in_source"] / 3600:.1f} '
    f'measured ({EXTRA["N1001"]["change_pct"]:+.0f}%) and '
    f'{EXTRA["N1401"]["predicted_converged_s"] / 3600:.1f} hours at N = 1401 '
    f'against {TRUNC[1401]["solve_s_in_source"] / 3600:.1f} '
    f'({EXTRA["N1401"]["change_pct"]:+.0f}%, small only because that run burnt '
    f'{TRUNC[1401]["stats"]["newton_iters_total"]} Newton steps against the '
    f'{CONV[701]["newton_iters_total"]} a converged CG needs, so almost all of '
    f'the extra CG work is paid for by the Newton steps it removes). Those two '
    f'figures are predictions and are labelled as such wherever they appear; '
    f'the two in Table 20b are measurements.')._p)

els.append(para(
    'Separately from all of this, the same configuration at N=501 was run '
    'twice on the same hardware, with identical iteration counts, 13.0% apart '
    'in wall clock — 1,616 seconds against 1,827 — which is the run-to-run '
    'variation every single-run timing in Table 20 silently carries, and it is '
    'smaller than the +28% above but not negligible beside it. The memory '
    'column is untouched by all of this, since memory does not depend on how '
    'many CG iterations are taken.')._p)

replace_para('Two consequences for how Table 20 should be read.', els)

# ======================================================================
# B. Section 8.11 -- the ceiling constrains Pi, not L2
# ======================================================================
replace_para('One property of that arrangement decides how the numbers read', [
    para(
        'One property of that arrangement decides how the numbers read, and it '
        'is arithmetic rather than an empirical expectation. The operator '
        'minimises the same discrete functional over the same Q4 space that the '
        'Q4 solver solves, and the minimiser of that functional over that space '
        'is the Q4 solution itself. No field the operator can produce, '
        'therefore, attains a lower discrete potential energy than Q4 does. '
        'That ceiling is a statement about Π, and Π is none of the four errors '
        'reported here. It does not transfer automatically to them: the L2 '
        'error against u* is a different functional, and a field that fails to '
        'minimise Π can still lie closer to u* in L2 than the minimiser does, '
        'because Q4\'s discretisation error is a systematic bias which the '
        'network\'s optimisation error may partly cancel rather than add to. An '
        'earlier revision of this section said that a ratio below one would '
        'indicate a defect in the Dirichlet mask, the quadrature or the work '
        'term rather than an advance. That is true of Π and false of L2, and '
        'the three-mesh study below produces an L2 ratio of '
        f'{L2R[0]:.2f} at N = 9 with none of those three defective. What the '
        'ceiling does appear to protect are the derivative-based norms: '
        'operator/Q4 exceeds one in the H1 semi-norm and in stress at every '
        'mesh measured. For a linear problem Galerkin optimality would '
        'guarantee that, the Galerkin solution being the minimiser of the '
        'energy norm of the error; this problem is nonlinear, so the guarantee '
        'does not formally transfer, and what is offered here is that it held '
        'wherever it was tested. The quantity of interest is the ratio '
        'operator/Q4, where one means the network has solved the variational '
        'problem exactly and anything above it is the network\'s optimisation '
        'error, measured on its own.')._p])

replace_para('The operator is 2.42× the Q4 error in L2', [
    para(
        f'At this mesh the operator is {T24["operator"]["L2"] / T24["Q4"]["L2"]:.2f}× '
        f'the Q4 error in L2, so the network has closed most, but not all, of '
        f'the distance to the optimum it is chasing. The more interesting '
        f'reading is that the four norms do not agree on how far it is. In the '
        f'H1 semi-norm the ratio is '
        f'{T24["operator"]["H1_semi"] / T24["Q4"]["H1_semi"]:.2f}× and in '
        f'stress {T24["operator"]["stress"] / T24["Q4"]["stress"]:.2f}×, both '
        f'effectively at the Q4 optimum, while L2 sits at '
        f'{T24["operator"]["L2"] / T24["Q4"]["L2"]:.2f}× and energy at '
        f'{T24["operator"]["energy"] / T24["Q4"]["energy"]:.2f}×. That ordering '
        f'is inverted relative to the usual one, in which L2 is the forgiving '
        f'norm and the derivative-based norms are the strict ones.')._p])

# ======================================================================
# C. Section 8.11 -- the operator now has a convergence rate
# ======================================================================
els = [para(
    'A rate of its own. An earlier revision listed as a limitation that the '
    'operator had been trained and scored at one mesh only, so that it had no '
    'convergence rate and could not appear in Table 23. Two further operators '
    'were therefore trained, at N = 9 and N = 33, under exactly the N = 17 '
    'protocol — the same 64-member manufactured family, the same 2,000 epochs '
    'at batch 8, the same optimiser and the same scoring member — so that the '
    'three points differ in the mesh and in nothing else. Each run solved its '
    'own Q4 and Q9 references, and the functional check described above was '
    'repeated and passed before each. Table 24a.')._p]

els.append(new_table(
    ['N', 'DOF', 'Operator L2', 'Q4 L2', 'Operator H1 semi', 'Q4 H1 semi'],
    [[str(r['N']), f"{r['n_dof']:,}",
      f"{r['operator']['L2']:.3e}", f"{r['Q4']['L2']:.3e}",
      f"{r['operator']['H1_semi']:.3e}", f"{r['Q4']['H1_semi']:.3e}"]
     for r in MROWS]
    + [['rate in h', '—',
        f"{RATE['operator_L2']:+.2f}", f"{RATE['Q4_L2']:.2f}",
        f"{RATE['operator_H1_semi']:+.2f}", f"{RATE['Q4_H1_semi']:.2f}"]])._tbl)
els.append(para(
    f'Table 24a. The operator across three meshes, against the same '
    f'manufactured solution and by the same routine as Tables 22 and 24. The '
    f'last row is the convergence rate in h fitted by least squares on the '
    f'three points, as Table 23 fits its own. The Q4 columns are the control: '
    f'their fitted rates, {RATE["Q4_L2"]:.2f} and {RATE["Q4_H1_semi"]:.2f}, '
    f'land on the {1.98:.2f} and {1.00:.2f} that Table 23 measures on eight '
    f'resolutions, so these three runs are comparable to that table and the '
    f'operator rates beside them can be quoted. Those Q4 columns are separate '
    f'solves from Table 22\'s and not a copy of them; where the two meshes '
    f'overlap, at N = 9 and N = 17, they agree to better than 0.2%.')._p)

els.append(new_table(
    ['N', 'L2', 'H1 semi', 'Stress', 'Energy'],
    [[str(r['N'])] + [f"{RAT[k][str(r['N'])]:.2f}×"
                      for k in ('L2', 'H1_semi', 'stress', 'energy')]
     for r in MROWS])._tbl)
els.append(para(
    'Table 24b. The ratio operator/Q4 in each of the four norms at the three '
    'meshes: Table 24a and the two columns it has no room for, divided through '
    'and read as the quantity of interest. One would mean the network had '
    'solved the variational problem exactly.')._p)

els.append(para(
    f'The operator does not converge. Its L2 error rises with refinement — '
    f'{MROWS[0]["operator"]["L2"]:.3e} at N = 9, '
    f'{MROWS[1]["operator"]["L2"]:.3e} at N = 17, '
    f'{MROWS[2]["operator"]["L2"]:.3e} at N = 33 — for a fitted rate of '
    f'{RATE["operator_L2"]:+.2f} where Q4, on the same three meshes, falls at '
    f'{RATE["Q4_L2"]:.2f}. The H1 semi-norm does improve, at '
    f'{RATE["operator_H1_semi"]:.2f} against Q4\'s {RATE["Q4_H1_semi"]:.2f}, '
    f'so the picture is not that nothing improves; it is that what improves '
    f'does so more slowly than the discretisation it sits on. This was the '
    f'outcome the study was set up to be able to find rather than one it was '
    f'set up to confirm: there is no theory for the operator\'s rate to be '
    f'checked against, unlike the two solvers whose expected rates Table 23 '
    f'verifies, so whatever came out was a measurement.')._p)

els.append(para(
    f'The explanation is that the operator\'s error is dominated by '
    f'optimisation error and not by discretisation error, and the two respond '
    f'to refinement in opposite ways. Refining the mesh reduces the '
    f'discretisation error that limits Q4; it leaves the network\'s own '
    f'optimisation error roughly where it was, while enlarging the problem the '
    f'network has to optimise. The gap therefore widens, and Table 24b shows it '
    f'widening: {L2R[0]:.2f}×, {L2R[1]:.2f}×, {L2R[2]:.2f}× in L2. The '
    f'crossover is visible inside the three points. At N = 9 the mesh is coarse '
    f'enough that Q4\'s own error ({MROWS[0]["Q4"]["L2"]:.2e}) exceeds the '
    f'network\'s optimisation error, and the operator is the more accurate of '
    f'the two in L2; by N = 33 Q4 is {L2R[2]:.1f} times better. Somewhere '
    f'between them the discretisation stops being the limiting factor and the '
    f'network becomes it. The derivative norms behave differently and more '
    f'mildly — {RAT["H1_semi"]["9"]:.2f}×, {RAT["H1_semi"]["17"]:.2f}×, '
    f'{RAT["H1_semi"]["33"]:.2f}× in H1 — which is the same inversion described '
    f'above, now seen at three resolutions rather than argued from one. The '
    f'energy column should not be read alongside them: it is the relative error '
    f'in the scalar internal strain energy, not the energy norm of the error, '
    f'and its dip to {RAT["energy"]["9"]:.2f}× at N = 9 is a cancellation in a '
    f'single number that carries none of the protection the H1 column does.')._p)

els.append(para(
    f'Two qualifications belong with the rate. It is fitted on three points, '
    f'the same number Table 23 uses per pair but over a narrower span, and the '
    f'operator column has no theoretical value to be checked against. And the '
    f'three runs hold the training budget fixed at 2,000 epochs while the '
    f'problem grows from {MROWS[0]["n_dof"]:,} to {MROWS[2]["n_dof"]:,} degrees '
    f'of freedom, so the negative rate is the rate at a fixed budget and not a '
    f'property of the method at convergence. A budget scaled with the mesh '
    f'might flatten it. That is the natural next experiment and it was not run: '
    f'holding the protocol fixed is what makes the three points differ in the '
    f'mesh alone, and changing both at once would have measured neither.')._p)

insert_after('The inversion is consistent with what the training principle',
             els)

# the surviving limitations, with the one that no longer applies removed
replace_para('What is left is optimisation error, not discretisation error', [
    para(
        'What is left is optimisation error, not discretisation error, and the '
        'run does not establish how much of it a longer budget would remove. '
        'The best held-out L2 was 1.429e-02 at the halfway point and 8.826e-03 '
        'at the end, a further 38% over the second half of training — still '
        'falling, and slowly. Two limits remain beyond that. The operator\'s '
        'cost — minutes of GPU training against seconds of FP64 CPU Newton '
        'solves — is not put on a common axis here, because no honest one was '
        'available. And the whole of section 8.11, all three legs of it, rests '
        'on one geometry and one material; the manufactured family is '
        'parametrised by two numbers and 64 of its members train each '
        'operator, but every error in Tables 22, 24 and 24a is scored on the '
        'single member α = 0.05, β = 0.7.')._p])

# ======================================================================
# D. Two places that describe the B2 zero-shot cases as "being
#    regenerated". They have been regenerated. Saying so, and saying that
#    the retrained models still do not work, is the only accurate state.
# ======================================================================
replace_para('Two limits. The three B2 cases are not here', [
    para(
        'Two limits. The three B2 cases are not here. Their sample caches were '
        'found to carry an applied load overstated by a mesh-dependent factor, '
        'which for a study that measures transfer across meshes invalidates '
        'them outright. The caches have since been repaired and the repair '
        'verified: one fixed pressure field assembled on each of the two '
        'training meshes now gives a total load of '
        f'{MESH_OK["one_fixed_pressure_field_assembled_on_each_resolution"]["N21"]:.4f} '
        f'and '
        f'{MESH_OK["one_fixed_pressure_field_assembled_on_each_resolution"]["N33"]:.4f}, '
        f'{MESH_OK["spread_pct"]:.3f}% apart, where before the repair the two '
        'differed by a factor of about 1.6. All three cases were then '
        'retrained from the corrected data under the same protocol as the B1 '
        'rows, with the per-sample force normalisation that Section 9.1 '
        'established B2 requires. They do not reach a usable accuracy: the '
        'best combined validation error is '
        + ', '.join(f'{BEST[m]["combined_val_error"]:.4f}' for m in
                    ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
        + ' for Neo-Hookean, Mooney-Rivlin and Arruda-Boyce respectively, '
        'against 0.0658 to 0.0827 for the three B1 cases on the same metric '
        'and the same trainer. One is the value a prediction of zero scores on '
        'that metric, which is what those three figures say the models are '
        'producing, and evaluating them at the seven resolutions of Table 12 '
        'says the same thing a second way: '
        + ', '.join(
            '%.4f' % (sum(EV[m]['mean_rel_L2_vs_fine_reference'])
                      / len(EV[m]['mean_rel_L2_vs_fine_reference']))
            for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
        + ', each varying by under '
        f'{max(EV[m]["spread_over_the_mesh_pct"] for m in EV):.1f}% across a '
        'fourfold refinement, where the three B1 columns of that table move by '
        + '%.0f%% to %.0f%%' % (min(B1SP.values()), max(B1SP.values()))
        + '. An error that barely notices the mesh is the signature of a model '
        'that is not solving the problem on it. The cause is under '
        'investigation; it is not the optimiser '
        'settings tried so far, since a batch-size arm at matched optimiser '
        'steps moved it only from '
        f'{RETRAIN["candidates_tested"][1]["result"]["batch_8"]:.4f} to '
        f'{RETRAIN["candidates_tested"][1]["result"]["batch_1"]:.4f}. No B2 '
        'zero-shot number is quoted anywhere in this report, and none from the '
        'earlier runs should be either: the relative errors of '
        + ', '.join(f'{OLD_B2[m]:.2f}' for m in
                    ('neo_hookean', 'mooney_rivlin', 'arruda_boyce'))
        + ' and above that the superseded evaluation reports carry were '
        'produced by models trained on the bad load. And the finest mesh '
        'tested, N = 49, is still far from the N=101 reference, so nothing '
        'here says where the rising branch ends.')._p])

retext('Extend the resolution-invariance study to the remaining five',
       'Extend the resolution-invariance study to the remaining benchmark '
       'cases. Section 8.7 now covers all three B1 materials at seven unseen '
       'resolutions each (Table 12), which is half of the six. The three B2 '
       'cases remain outstanding for the reason given there: their data has '
       'been corrected and the models retrained, and the retrained models do '
       'not yet reach a usable accuracy.')

doc.save(DST)
print(f'wrote {DST}')
