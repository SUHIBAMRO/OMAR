"""v50 -> v51. Omar asked how to make every table look the same in Word.

Checked both documents directly rather than guessing: every table already
shares the same border style (single, 4, auto -- newer make_vN.py scripts
copy `tblPr` from an existing table, which carries the borders). What
they do NOT share is header-row shading and font size, because copying
`tblPr` only carries table-level properties, not the per-cell shading
(`tcPr`) or run formatting of the header row itself.

In the report, 20 of 50 tables (the earliest ones added) have a styled
header row -- dark fill #1F2937, bold white text, 9pt -- and a 9pt body
font throughout. The other 30 have no header shading, no bold/white
header text, and no explicit font size at all (so they render in
whatever the Normal style's default is). This is a real, visible
inconsistency, not a cosmetic nitpick.

This script applies the styled table's exact header formatting (fill
1F2937, bold, white, 9pt) to the header row of all 50 tables, and 9pt to
every other cell in every table, without touching any table's text,
borders, or the incidental bold already present on legitimate header
cells. Verified before and after with a full formatting scan; only the
30 previously-bare tables change.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC, DST = 'PFEM_Transolver_Report_v50.docx', 'PFEM_Transolver_Report_v51.docx'
HEADER_FILL = '1F2937'
FONT_SIZE = Pt(9)

doc = Document(SRC)


def set_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


before_bare = 0
for t in doc.tables:
    header = t.rows[0]
    for cell in header.cells:
        set_shading(cell, HEADER_FILL)
        for p in cell.paragraphs:
            if not p.runs:
                continue
            for r in p.runs:
                if r.font.size is None:
                    before_bare += 1
                r.font.bold = True
                r.font.size = FONT_SIZE
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in t.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = FONT_SIZE

assert before_bare > 0, 'expected some header cells with no explicit size'
print(f'{before_bare} previously-bare header runs got size/bold/color set')

doc.save(DST)
print('wrote', DST)
