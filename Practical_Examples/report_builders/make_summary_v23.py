"""v22 -> v23. Omar quoted a stale sentence straight out of this document:
"Still open: the same high-DOF study for B2, and the five remaining
resolution-invariance cases." Checking it found four leftover sentences
in the summary's opening "Summary of what was done" walkthrough and one
in its "Remaining work" closer, all describing a status this project
moved past long ago:

  - para 9 ("Resolution invariance...") still described the single-
    resolution, five-unseen-resolution protocol and said "The other five
    cases are still running."
  - para 10 -- the quoted sentence itself.
  - para 12 (the "1. Mesh refinement" item) said "B2's DOF-referenced
    study is not started yet" -- it was CANCELLED by Omar's own request
    on 2026-08-27, which is a scope decision, not a pending measurement.
  - para 69 (the "7. Resolution invariance" item) said "done for 1 of 6
    cases. The other 5 are running now" -- three paragraphs later (77)
    this same document already says "Every B1/B2 x material combination
    now has a valid zero-shot result."
  - para 168 ("Remaining work.") named "the out-of-distribution
    attribution extension named in the report's Section 10" as the one
    open item left -- that item closed too (Table 19/19a already do
    exactly that isolation; see make_v50.py, which fixes the same stale
    claim in the report itself).

All five rewritten to match what the rest of this same document (and the
report, as of v50) already says elsewhere.
"""
from docx import Document

SRC, DST = 'PFEM_Summary_Completed_Work.docx', 'PFEM_Summary_Completed_Work.docx'
BACKUP = 'PFEM_Summary_Completed_Work.pre_v23.docx'

doc = Document(SRC)
doc.save(BACKUP)


def retext_para(idx, old_start, new_text):
    p = doc.paragraphs[idx]
    assert p.text.strip().startswith(old_start), (
        f'paragraph {idx} does not start with {old_start!r}: {p.text!r}')
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = new_text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)


retext_para(9, 'Resolution invariance.', (
    'Resolution invariance. Following your correction, each model is now '
    'trained once — jointly on two resolutions (N = 21 and 33) — and '
    'evaluated with no retraining on seven resolutions it never saw, all '
    'compared against a common fine reference (N=101). All three B1 '
    'materials stay between 5.0% and 10.6% across every resolution. B2 '
    'shows the same real resolution invariance, weaker than B1’s. All '
    'six geometry–material combinations now have a valid result.'
))

retext_para(10, 'Still open:', (
    'Both items once open here are now closed: the B2 high-DOF study was '
    'cancelled at the sponsor’s own request, not merely postponed, and '
    'all six resolution-invariance cases have a valid result.'
))

retext_para(12, 'Confirmed for all 6 cases', (
    'Confirmed for all 6 cases (h-refinement sweep, N=6–51). B1’s '
    'Q4-vs-Q9 check against a ~10M/40M-DOF reference is also done. The '
    'equivalent B2 study was cancelled at the sponsor’s request and is '
    'not planned — a scope decision, not an unfinished measurement.'
))

retext_para(69, 'True zero-shot protocol', (
    'True zero-shot protocol (train once, evaluate on unseen resolutions '
    'vs. a common fine reference) done for all 6 cases.'
))

retext_para(168, 'Remaining work.', (
    'Remaining work. Zero-shot resolution invariance is reported for all '
    'six geometry–material cases — every B1/B2 × material combination '
    'has a valid result, above; nothing remains open there. The '
    'B2-geometry high-DOF, element-order study (Section 4.4’s deeper '
    'Q4-vs-Q9 check) was cancelled at the sponsor’s request and is not '
    'planned, not merely postponed. Point 2’s accuracy/cost Pareto is '
    'now measured for all six geometry × material combinations (Table '
    '18e closed the last of them). The out-of-distribution attribution '
    'extension once named as the report’s one remaining item is closed '
    'too — Table 19 isolates the material-stiffness and loading-'
    'magnitude shifts, and Table 19a tests normalization as a '
    'mitigation. No open item remains from the report’s stated next '
    'steps.'
))

doc.save(DST)
print('overwrote', DST, '(backup at', BACKUP, ')')
