"""v49 -> v50. The same leftover-status pattern fixed in v49's Table 1
row 7 turned up again while checking the rest of Section 10's "remaining
items" list for the same kind of staleness. Of its four bullets, three
(the B2 zero-shot cases, the B1/B2 accuracy gap, and resolution
invariance) already say plainly that they are closed. The fourth does
not:

  "Extend the out-of-distribution evaluation to isolate the individual
  contributions of the material-stiffness shift and the loading-
  magnitude shift, which were varied together in Section 8.6."

This is stale. Section 8.6 already contains exactly that isolation:
Table 19 varies material and loading separately (and combined) at seven
shift levels each, and Table 19a repeats it for a normalization
mitigation. PROJECT_STATUS.md's own master table lists this as done
("R6-1 progressive OOD: material vs loading... Table 19", "R6-1b
normalization tested as a mitigation... Table 19a"). The bullet was
simply never updated after Table 19/19a closed it -- the same failure
mode as Table 1's row 7.

Rewritten to match the style of the other three bullets in the same
list (state what was asked, then that it is closed and where). With all
four bullets now stating they are closed, the list's own heading, "The
remaining items are:", is also rewritten -- nothing in it remains.
"""
from docx import Document

SRC, DST = 'PFEM_Transolver_Report_v49.docx', 'PFEM_Transolver_Report_v50.docx'

doc = Document(SRC)
ORIGINAL = list(doc.paragraphs)


def find_para(prefix):
    hits = [p for p in ORIGINAL if p.text.strip().startswith(prefix)]
    assert len(hits) == 1, (
        f'{len(hits)} paragraphs start with {prefix!r}; expected exactly 1')
    return hits[0]


def retext(prefix, text):
    p = find_para(prefix)
    keep = p.runs[0] if p.runs else p.add_run()
    keep.text = text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)
    return p


old_prefix = 'Extend the out-of-distribution evaluation to isolate'
old = find_para(old_prefix)
assert old.text.strip() == (
    'Extend the out-of-distribution evaluation to isolate the individual '
    'contributions of the material-stiffness shift and the loading-'
    'magnitude shift, which were varied together in Section 8.6.'), (
    f'unexpected stale text: {old.text!r}')

retext(old_prefix, (
    'Extend the out-of-distribution evaluation to isolate the individual '
    'contributions of the material-stiffness shift and the loading-'
    'magnitude shift, which were varied together in the original single-'
    'shift measurement of Table 11 — done in Section 8.6: Table 19 shows '
    'the degradation comes entirely from the material-stiffness shift, '
    'not the loading shift, at seven shift levels; Table 19a then tests '
    'input normalization as a mitigation and finds it is not a clean fix. '
    'This item is closed.'
))

# With all four bullets below now closed, the heading itself is stale.
retext('The remaining items are:',
       'The four items originally listed here have all since closed:')

doc.save(DST)
print('wrote', DST)
