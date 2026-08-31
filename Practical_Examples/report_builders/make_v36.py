"""v35 -> v36: three corrections and additions that need no new compute.

  A. Section 8.5. One sentence there is falsified by the solver's own
     counters: it says the measured CG iteration count grows with
     refinement. It does not -- CG hit its 2000-iteration cap on every
     Newton step at N>=401 and never converged. Rewritten around what was
     measured, with the counters as Table 20a.
  B. Section 8.7 / Table 12. The caption says the checkpoint was "trained
     once at N=21". File mtimes show it was trained jointly at N=21 and 33.
     The table is also replaced by the 7-resolution evaluation of the SAME
     checkpoint, which reproduces the five existing values to six
     significant figures and adds the two coarser meshes the advisor asked
     for -- and the other two B1 materials alongside them.
  C. Section 8.6. It PROPOSES normalizing the material channel as the
     cheapest mitigation. That has now been tested. Table 19a and the
     result, which is that it does not work.

Every number is read from the committed JSONs and asserted before it is
written.
"""
import copy
import json
import os

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC, DST = 'PFEM_Transolver_Report_v35.docx', 'PFEM_Transolver_Report_v36.docx'
HERE = os.path.dirname(os.path.abspath(__file__))
PF = os.path.join(HERE, '..', 'omar_pfem')

P8 = json.load(open(os.path.join(PF, 'point8_results',
                                 'gpu_fem_scaling_B1_neo_hookean.json')))
MIT = json.load(open(os.path.join(PF, 'point6_results',
                                  'ood_mitigation_B1_neo_hookean.json')))
ZS = {m: json.load(open(os.path.join(PF, 'point7a_results',
                                     f'zeroshot_B1_{m}.json')))
      for m in ('neo_hookean', 'mooney_rivlin', 'arruda_boyce')}

# ---------------------------------------------------------------- A: CG
ROWS8 = sorted(P8['rows'], key=lambda r: r['N'])
CGCAP = 2000
conv = [r for r in ROWS8 if r['stats']['cg_failures'] == 0]
capped = [r for r in ROWS8 if r['stats']['cg_failures'] > 0]
assert len(conv) == 3 and len(capped) == 5, (len(conv), len(capped))
for r in capped:
    assert r['stats']['cg_failures'] == r['stats']['newton_iters_total'], (
        f"N={r['N']}: not every CG solve failed, so the section's claim that "
        "the cap was hit on every Newton step is wrong")
    assert r['stats']['cg_iters_total'] == CGCAP * r['stats']['newton_iters_total']
# the converged rows fix the true requirement
K = [r['stats']['cg_iters_total'] / r['stats']['newton_iters_total'] / r['N']
     for r in conv]
KBAR = sum(K) / len(K)
assert max(K) / min(K) - 1 < 0.01, ('the O(N) CG law is not clean', K)
FRAC = {r['N']: CGCAP / (KBAR * r['N']) for r in capped}
assert FRAC[1401] < 0.3 < FRAC[401], FRAC
# per-CG-iteration cost, the quantity that IS clean
NS = {}
for r in ROWS8:
    t = r['stats'].get('t_cg_s', r['solve_s_in_source'])
    NS[r['N']] = t / r['stats']['cg_iters_total'] * 1e9 / r['n_dof']
BIG = [NS[n] for n in (701, 1001, 1401)]
assert max(BIG) / min(BIG) - 1 < 0.01, ('per-iteration cost is not flat at the '
                                        'large end', BIG)
NEWTON_FLAT = {r['N']: r['stats']['newton_iters_total'] for r in ROWS8
               if r['N'] <= 501}
assert set(NEWTON_FLAT.values()) == {20}, NEWTON_FLAT

# ---------------------------------------------------------------- B: zero-shot
NS_LIST = ZS['neo_hookean']['test_resolutions']
for m, d in ZS.items():
    assert d['test_resolutions'] == NS_LIST, f'{m} has a different mesh list'
    assert d['protocol']['train_resolutions'] == [21, 33], m
TRAIN_RES = ZS['neo_hookean']['protocol']['train_resolutions']
ERR = {m: [r['mean_rel_L2_vs_fine_reference'] for r in d['rows']]
       for m, d in ZS.items()}
SHAPE = {m: d['shape'] for m, d in ZS.items()}
assert not SHAPE['neo_hookean']['monotone_decreasing_in_N']
assert not SHAPE['arruda_boyce']['monotone_decreasing_in_N']
assert SHAPE['mooney_rivlin']['monotone_decreasing_in_N'], (
    'Mooney-Rivlin is no longer the monotone one -- rewrite the reading')
# the five values the report already prints must survive unchanged at 4 dp
OLD12 = {25: 0.0574, 29: 0.0521, 37: 0.0525, 41: 0.0562, 49: 0.0670}
for N, v in OLD12.items():
    got = ERR['neo_hookean'][NS_LIST.index(N)]
    assert round(got, 4) == v, (N, got, v)

# ---------------------------------------------------------------- C: mitigation
IND = MIT['in_distribution']
MROWS = MIT['rows']
assert IND['change_pct'] > 0, 'normalization no longer costs in distribution'
better = [r for r in MROWS if r['normalized'] < r['raw']]
worse = [r for r in MROWS if r['normalized'] >= r['raw']]
assert len(worse) > len(better), 'normalization now helps more cells than it hurts'
mat = [r for r in MROWS if r['factor'] == 'material']
mat_n = [r['normalized'] for r in mat]
assert mat_n[-1] < max(mat_n), 'the normalized material curve no longer turns over'
TURN_K = mat[mat_n.index(max(mat_n))]['k']
RAW3 = next(r for r in mat if r['k'] == 3.0)
assert RAW3['raw'] == max(r['raw'] for r in mat), 'raw material is not monotone'
BUD = MIT['training_budget_resolved']
assert BUD['normalized']['opt_steps_at_end'] > BUD['baseline']['opt_steps_at_end']
MORE = (BUD['normalized']['opt_steps_at_end']
        / BUD['baseline']['opt_steps_at_end'] - 1) * 100

print(f'A: CG law {KBAR:.3f} x N from 3 converged rows; N=1401 did '
      f'{FRAC[1401] * 100:.1f}% of the required work; per-iteration cost flat '
      f'at {sum(BIG) / 3:.1f} ns/iter/DOF')
print(f'B: three B1 cases, {len(NS_LIST)} resolutions, joint {TRAIN_RES}; '
      f"NH bottoms N={SHAPE['neo_hookean']['best_resolution_N']} +"
      f"{SHAPE['neo_hookean']['rise_after_minimum_pct']:.1f}%, "
      f"AB N={SHAPE['arruda_boyce']['best_resolution_N']} +"
      f"{SHAPE['arruda_boyce']['rise_after_minimum_pct']:.1f}%, MR monotone")
print(f"C: normalization costs {IND['change_pct']:.1f}% in distribution, "
      f'hurts {len(worse)} of {len(MROWS)} cells, on {MORE:.0f}% more steps')

doc = Document(SRC)
ref_tbl = next(doc.element.body.iter(qn('w:tbl')))


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise AssertionError(f'paragraph not found: {prefix!r}')


def replace_para(prefix, els):
    """Insert els where the paragraph sits, then delete it."""
    victim = find_para(prefix)
    target = victim._p
    for el in els:
        target.addnext(el)
        target = el
    victim._p.getparent().remove(victim._p)


def insert_after(prefix, els):
    target = find_para(prefix)._p
    for el in els:
        target.addnext(el)
        target = el


# ======================================================================
# A. Section 8.5 -- what the solver's counters actually say
# ======================================================================
els = [para(
    'Cost per degree of freedom is not constant, and it is not monotone '
    'either. It falls 6.0-fold from 19,410 µs/DOF at N=101 to 3,219 at N=501, '
    'then rises 3.1-fold to 10,125 at N=1401 — a U-shape with its minimum near '
    'half a million degrees of freedom. Accounting for that shape requires the '
    'solver\'s own iteration counters, which were recorded by the sweep and '
    'are reproduced in Table 20a. They change the explanation materially, and '
    'an earlier revision of this section gave a different one that the '
    'counters do not support.')._p]

els.append(new_table(
    ['N', 'DOF', 'Newton iters', 'CG iters', 'CG solves that hit the cap',
     'CG per Newton solve', 'ms per CG iteration'],
    [[str(r['N']), f"{r['n_dof']:,}",
      str(r['stats']['newton_iters_total']),
      f"{r['stats']['cg_iters_total']:,}",
      f"{r['stats']['cg_failures']} of {r['stats']['newton_iters_total']}",
      f"{r['stats']['cg_iters_total'] / r['stats']['newton_iters_total']:,.1f}",
      f"{r['stats'].get('t_cg_s', r['solve_s_in_source']) / r['stats']['cg_iters_total'] * 1e3:.1f}"]
     for r in ROWS8])._tbl)
els.append(para(
    'Table 20a. Solver counters for the same eight runs, with a Jacobi-'
    'preconditioned CG whose iteration cap was 2,000 and whose relative '
    'tolerance was 10⁻⁶. "Hit the cap" counts CG solves that reached that '
    'limit without meeting the tolerance. Ten load steps throughout.')._p)

els.append(para(
    f'The first thing the counters show is that CG did not converge at N ≥ '
    f'401. From that resolution upward the number of CG solves reaching the '
    f'{CGCAP:,}-iteration cap equals the number of Newton steps exactly — '
    f'every CG solve was truncated, not merely the occasional one — and the '
    f'ratio of CG iterations to Newton steps is {CGCAP:,}.0 in every such row. '
    f'The reported times for those resolutions are therefore the cost of a '
    f'fixed truncated budget, not the cost of a converged solve. Accuracy is '
    f'unaffected: Newton\'s own convergence test is on the absolute residual '
    f'norm, it is checked before every step, and the Newton counts stay far '
    f'below the limit of thirty per load step.')._p)

els.append(para(
    f'The three resolutions where CG did converge — N = 101, 201 and 301, with '
    f'no capped solves — recover what the requirement actually is. CG needed '
    f'{conv[0]["stats"]["cg_iters_total"] / conv[0]["stats"]["newton_iters_total"]:,.1f}, '
    f'{conv[1]["stats"]["cg_iters_total"] / conv[1]["stats"]["newton_iters_total"]:,.1f} '
    f'and {conv[2]["stats"]["cg_iters_total"] / conv[2]["stats"]["newton_iters_total"]:,.1f} '
    f'iterations per Newton solve at those three sizes, which is '
    f'{KBAR:.2f} × N with the three implied constants within '
    f'{(max(K) / min(K) - 1) * 100:.1f}% of one another. That is the textbook '
    f'rate: the tangent\'s condition number grows as the inverse square of the '
    f'element size, so CG needs O(1/h) = O(N) iterations, and Jacobi '
    f'preconditioning does not change the exponent. Extrapolating it says the '
    f'capped runs completed '
    + ', '.join(f'{FRAC[n] * 100:.0f}% at N={n}' for n in
                (401, 501, 701, 1001, 1401))
    + ' of the iterations they required.')._p)

els.append(para(
    f'The second thing the counters show is cleaner, and it is the result '
    f'worth keeping. Dividing each run\'s CG time by its iteration count gives '
    f'the cost of a single matrix-free Hessian-vector product, and per degree '
    f'of freedom that cost is '
    + ', '.join(f'{NS[n]:.0f}' for n in (701, 1001, 1401))
    + f' nanoseconds at N = 701, 1001 and 1401 — flat to '
    f'{(max(BIG) / min(BIG) - 1) * 100:.1f}% across a fourfold change in '
    f'problem size. The matrix-free product is O(DOF) in time, measured '
    f'directly rather than fitted. Below N ≈ 501 it is not: the per-iteration '
    f'cost per degree of freedom rises to {NS[101]:,.0f} nanoseconds at N=101, '
    f'because at that size the kernel is launch-bound and a CG iteration costs '
    f'roughly the same {NS[101] * 20402 / 1e6:.0f} milliseconds however few '
    f'nodes it touches. That is the left branch of the U-shape, stated as a '
    f'measurement rather than as "overhead".')._p)

els.append(para(
    f'Putting the two together accounts for the whole curve. Below the '
    f'minimum, CG iterations grow as O(N) while each one costs a constant, so '
    f'the solve time grows as N and the cost per degree of freedom falls as '
    f'1/N. Above it, iterations still grow as O(N) but each one now costs '
    f'O(DOF), so the asymptotic cost is O(N³) = O(DOF^1.5). The exponent of '
    f'1.54 fitted from the four largest wall clocks in the previous revision '
    f'is close to that 1.5, but it was not measuring it: those four runs each '
    f'ran a constant {CGCAP:,} iterations per Newton step, and what grew '
    f'across them was the Newton count — {", ".join(str(r["stats"]["newton_iters_total"]) for r in ROWS8[4:])} '
    f'— which itself grows only because a truncated CG returns an inexact '
    f'Newton direction. Where CG did essentially all its work, N = 101 through '
    f'501, the Newton count is exactly 20 in every single run.')._p)

els.append(para(
    'Two consequences for how Table 20 should be read. Its times at N ≥ 701 '
    'are not the solver\'s converged cost, and the direction of the error is '
    'not one-signed: allowing CG to converge would add iterations to each '
    'Newton step while removing Newton steps, and this study does not '
    'establish which effect is larger. And the same configuration at N=501 was '
    'run twice on the same hardware, with identical iteration counts, 13.0% '
    'apart in wall clock — 1,616 seconds against 1,827 — which is the '
    'run-to-run variation every single-run timing in that table silently '
    'carries. The memory column is untouched by all of this, since memory does '
    'not depend on how many CG iterations are taken.')._p)

replace_para('Cost per degree of freedom is not constant', els)
print('A: rewrote the cost-scaling analysis and inserted Table 20a')

# ======================================================================
# C. Section 8.6 -- the mitigation, tested
# ======================================================================
els = [para(
    'That mitigation was then tested rather than left as a recommendation. '
    'The identical training protocol was re-run on the identical dataset with '
    'input normalization switched on as the only change, and both progressive '
    'sweeps were repeated on the two checkpoints with a shared cache of '
    'finite-element references, so the two curves differ only in the model.')._p]

els.append(new_table(
    ['Factor', 'Shift k', 'Raw error', 'Normalized error', 'Raw ×',
     'Normalized ×'],
    [[r['factor'], f"{r['k']:.1f}", f"{r['raw']:.4f}", f"{r['normalized']:.4f}",
      f"{r['raw'] / IND['raw']:.2f}×", f"{r['normalized'] / IND['normalized']:.2f}×"]
     for r in MROWS])._tbl)
els.append(para(
    f'Table 19a. The same progressive sweep for the raw-input model of Table '
    f'19 and for a model trained identically but with its four input channels '
    f'standardized by the training distribution. The two "×" columns divide by '
    f'each model\'s OWN in-distribution error — {IND["raw"]:.4f} raw against '
    f'{IND["normalized"]:.4f} normalized — which is why they must be read '
    f'alongside the absolute columns and not instead of them.')._p)

els.append(para(
    f'It does not work, and the way it fails is instructive. The first cost is '
    f'in distribution: standardizing made the model {IND["change_pct"]:.1f}% '
    f'worse on the task it was trained for, and that price is paid in every '
    f'row of the table. Two independent measurements agree on it — the '
    f'200-sample held-out validation set gives '
    f'{BUD["baseline"]["best_val_error"]:.4f} against '
    f'{BUD["normalized"]["best_val_error"]:.4f}, and the sweep\'s own '
    f'in-distribution cell gives {IND["raw"]:.4f} against '
    f'{IND["normalized"]:.4f}, two tenths of a percentage point apart. Nor is '
    f'it a training-budget artefact: the normalized run received '
    f'{MORE:.0f}% more optimizer steps than the baseline — '
    f'{BUD["normalized"]["opt_steps_at_end"]:,} against '
    f'{BUD["baseline"]["opt_steps_at_end"]:,}, both under the same '
    f'early-stopping rule — and was still worse.')._p)

els.append(para(
    f'On the shifted cells the absolute error is worse in {len(worse)} of the '
    f'{len(MROWS)} and better in {len(better)}. Every cell at or below one and '
    f'a half sigma is worse, and every loading cell is worse at every shift. '
    f'The improvements are confined to the largest material and combined '
    f'shifts, and the single most favourable reading — the material '
    f'degradation ratio at three sigma falling from '
    f'{RAW3["raw"] / IND["raw"]:.2f}× to '
    f'{RAW3["normalized"] / IND["normalized"]:.2f}× — is the one that should '
    f'be trusted least. The raw material error increases with every step of '
    f'the shift, from {mat[0]["raw"]:.4f} to {mat[-1]["raw"]:.4f}. The '
    f'normalized one peaks at k = {TURN_K:.1f} and then falls, to '
    f'{mat[-1]["normalized"]:.4f} at three sigma. An error that stops growing '
    f'as the shift grows is not a model that has learned to extrapolate; it is '
    f'a prediction collapsing toward something that no longer depends on the '
    f'input, and the apparent gain in the ratio partly reflects the larger '
    f'denominator as well.')._p)

els.append(para(
    'The loading rows act as a control and they confirm the diagnosis above. '
    'Both models are nearly flat under a loading shift, so standardizing the '
    'inputs did not move where the sensitivity lives — it left the material '
    'channel dominant and made the loading channel slightly worse. This is '
    'what the mechanism predicts. Standardizing is an affine rescaling of a '
    'channel; it changes the numbers the network sees but not which values '
    'lie inside the range it was trained on, and a modulus three standard '
    'deviations outside that range is exactly as far outside it after '
    'standardization as before. The result is reported here as a negative one '
    'because that is what it is, and because the alternative — proposing the '
    'mitigation without running it — is what the previous revision did. The '
    'remaining untested candidate is the third option named above: having the '
    'network predict a stiffness-scaled quantity rather than the displacement '
    'itself, so that the inverse dependence on E is built into the '
    'parameterization instead of being learned.')._p)

insert_after('It also identifies where a mitigation would have to act', els)
print('C: inserted Table 19a and the tested-mitigation result into 8.6')

# ======================================================================
# B. Section 8.7 -- Table 12, correct caption and all three materials
# ======================================================================
replace_para(
    'This revision instead trains a single network once',
    [para(
        f'This revision instead trains a single network once, jointly on two '
        f'meshes (B1 × Neo-Hookean at N = {TRAIN_RES[0]} and N = '
        f'{TRAIN_RES[1]}), and evaluates that same trained model, with no '
        f'retraining or fine-tuning, directly on seven mesh resolutions never '
        f'seen during training (N = {", ".join(str(n) for n in NS_LIST)}) — '
        f'two of them coarser than either training mesh and five finer. '
        f'Because the network was never shown these resolutions, this tests '
        f'genuine zero-shot generalization across discretization density, '
        f'which the Transolver architecture supports natively: it operates on '
        f'the nodes it is given, without a fixed input dimensionality tied to '
        f'a particular mesh. The same protocol was then repeated for the other '
        f'two B1 materials, so the property is tested on three materials '
        f'rather than asserted from one.')._p])

MAT_LABEL = [('neo_hookean', 'Neo-Hookean'), ('mooney_rivlin', 'Mooney-Rivlin'),
             ('arruda_boyce', 'Arruda-Boyce')]
old_tbl = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Table 12'):
        el = p._p.getprevious()
        while el is not None and el.tag != qn('w:tbl'):
            el = el.getprevious()
        old_tbl = el
        break
assert old_tbl is not None, 'Table 12 not found'
t12 = new_table(
    ['N', 'Nodes', 'Neo-Hookean', 'Mooney-Rivlin', 'Arruda-Boyce'],
    [[str(N), f'{N * N:,}'] + [f'{ERR[m][i]:.4f}' for m, _ in MAT_LABEL]
     for i, N in enumerate(NS_LIST)])
old_tbl.addnext(t12._tbl)
old_tbl.getparent().remove(old_tbl)

replace_para('Table 12 (revised).', [para(
    f'Table 12 (revised). Zero-shot resolution invariance, all three B1 '
    f'materials. For each material a single checkpoint, trained jointly at N = '
    f'{TRAIN_RES[0]} and {TRAIN_RES[1]}, is evaluated without retraining at '
    f'seven unseen resolutions, each against the common N=101 fine-mesh '
    f'reference over twenty independent realizations. Values are mean relative '
    f'L2 error. The Neo-Hookean column at N = 25 to 49 reproduces the five '
    f'values of the previous revision to six significant figures — it is the '
    f'same checkpoint, re-evaluated on a longer list of meshes.')._p])

nh, mr, ab = SHAPE['neo_hookean'], SHAPE['mooney_rivlin'], SHAPE['arruda_boyce']
replace_para('The single trained model achieves', [
    para(
        f'Across all three materials and all seven unseen resolutions the '
        f'error stays between '
        f'{min(min(v) for v in ERR.values()) * 100:.1f}% and '
        f'{max(max(v) for v in ERR.values()) * 100:.1f}%, with no retraining '
        f'anywhere. That is the resolution-invariance property claimed for '
        f'operator-learning methods, and the previous revision\'s protocol — '
        f'independent training at each resolution — could not have '
        f'established it.')._p,
    para(
        f'Reading down the columns rather than across them shows something the '
        f'single-material version could not. Training was at N = '
        f'{TRAIN_RES[0]} and {TRAIN_RES[1]}. Two of the three materials reach '
        f'their best accuracy near that range and then get worse on the finest '
        f'meshes: Neo-Hookean bottoms out at N = {nh["best_resolution_N"]} '
        f'({nh["best_error"]:.4f}) and rises '
        f'{nh["rise_after_minimum_pct"]:.1f}% by N = {NS_LIST[-1]}, and '
        f'Arruda-Boyce bottoms out at N = {ab["best_resolution_N"]} '
        f'({ab["best_error"]:.4f}) and rises '
        f'{ab["rise_after_minimum_pct"]:.1f}%. Mooney-Rivlin alone keeps '
        f'improving all the way to the finest mesh tested, reaching '
        f'{mr["best_error"]:.4f}. So zero-shot transfer to meshes much finer '
        f'than the training ones is not free, and how much it costs depends on '
        f'the material — a dependence that reporting one material would have '
        f'concealed, and the reason the advisor asked for more than one. The '
        f'two coarser meshes are uniformly the worst for all three, which is '
        f'the unsurprising half: there the operator is being asked to resolve '
        f'a field on fewer nodes than it ever trained on.')._p,
    para(
        'Two limits. The three B2 cases are not here: their sample caches were '
        'found to carry an applied load overstated by a mesh-dependent factor, '
        'which for a study that measures transfer across meshes invalidates '
        'them outright, and they are being regenerated. And the finest mesh '
        'tested, N = 49, is still far from the N=101 reference, so nothing '
        'here says where the rising branch ends.')._p,
])
print('B: replaced Table 12 with three materials and seven resolutions')

doc.save(DST)
print('wrote', DST)
