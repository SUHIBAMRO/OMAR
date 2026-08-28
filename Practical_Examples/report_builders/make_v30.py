"""v29 -> v30: the CPU and GPU break-even baselines side by side.

Timon, round 6: "the GPU-FEM break-even result of approximately 7,600-96,000
samples is very important. Please report both the CPU-FEM and GPU-FEM
baselines side by side. It also clarifies where the neural operator is
useful."

Both figures are already in the document -- 52-1,245 against CPU in Section
8.3, and Table 10c against GPU -- but they sit four pages apart in different
units of presentation, which is exactly why the range he quoted back was the
batch-size-128 column alone. One table with both, plus the correction, fixes
that.

Every number is recomputed here from the report's own tables rather than
copied: the CPU column from Table 7 (training wall-clock, native FEM cost,
inference latency) and the GPU columns from Table 10c. The CPU column is
asserted against the "52-1,245" range Section 8.3 already states, so a
mismatch fails the build instead of printing a second, disagreeing figure.
"""
import copy

from docx import Document
from docx.oxml.ns import qn

SRC, DST = 'PFEM_Transolver_Report_v29.docx', 'PFEM_Transolver_Report_v30.docx'
doc = Document(SRC)
body = doc.element.body
ref_tbl = next(body.iter(qn('w:tbl')))

CASES = ['B1 × Neo-Hookean', 'B1 × Mooney-Rivlin', 'B1 × Arruda-Boyce',
         'B2 × Neo-Hookean', 'B2 × Mooney-Rivlin', 'B2 × Arruda-Boyce']

# --- Table 7: training wall-clock (s), native CPU FEM cost (s), inference (ms)
TRAIN_S = [2873.8, 2785.3, 2855.6, 32244.0, 34164.1, 24846.8]
CPU_FEM_S = [25.432, 53.735, 52.542, 25.909, 61.712, 60.285]
INFER_MS = [4.586, 4.693, 4.745, 4.809, 4.908, 4.984]

# --- Table 10c: break-even against the GPU-native solver, per batch size
GPU_BE = [[1745, 6021, 7543, 8112],
          [1363, 5663, 7179, 7751],
          [1133, 5441, 6956, 7554],
          [19410, 67391, 84627, 90990],
          [17033, 69404, 87884, 95038],
          [9530, 46993, 60490, 65698]]

cpu_be = [t / (f - ms / 1000.0) for t, f, ms in zip(TRAIN_S, CPU_FEM_S, INFER_MS)]
lo, hi = min(cpu_be), max(cpu_be)
assert 51 <= lo <= 53 and 1240 <= hi <= 1250, \
    f'CPU break-even {lo:.0f}-{hi:.0f} does not reproduce Section 8.3\'s 52-1,245'
print(f'CPU break-even reproduces Section 8.3: {lo:.0f}-{hi:.0f}')

gpu_all = [v for row in GPU_BE for v in row]
gpu_bs1 = [row[0] for row in GPU_BE]
print(f'GPU break-even: {min(gpu_all):,}-{max(gpu_all):,} overall, '
      f'{min(gpu_bs1):,}-{max(gpu_bs1):,} at batch size 1')


def new_table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = doc.tables[0].style
    pr = ref_tbl.find(qn('w:tblPr'))
    if pr is not None:
        old = t._tbl.find(qn('w:tblPr'))
        if old is not None:
            t._tbl.remove(old)
        t._tbl.insert(0, copy.deepcopy(pr))
    for j, h in enumerate(header):
        c = t.cell(0, j)
        c.text = ''
        c.paragraphs[0].add_run(h).bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    return t


def para(text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


rows = [[c, f'{be:,.0f}'] + [f'{v:,}' for v in g]
        for c, be, g in zip(CASES, cpu_be, GPU_BE)]

els = []
h = doc.add_paragraph(style='Heading 3')
h.add_run('Break-even against both baselines, side by side')
els.append(h._p)

els.append(para(
    'The two break-even figures given above answer the same question against '
    'different baselines, and quoting either alone is misleading in opposite '
    'directions. Against the CPU reference solver — which is what generating a '
    'new solution actually costs today, and therefore the honest figure for a '
    'practitioner who has no GPU-native solver — the operator repays its '
    'training in tens to hundreds of problems. Against the GPU-native solver of '
    'this section, which is 71.7–171.5× faster, it repays in thousands to tens '
    'of thousands. Both are collected here so that neither can be quoted '
    'without the other.')._p)

els.append(new_table(
    ['Case', 'vs. CPU FEM', 'vs. GPU FEM, bs=1', 'bs=8', 'bs=32', 'bs=128'],
    rows)._tbl)
els.append(para(
    'Table 10d. Break-even in new problem instances, against both baselines. '
    'The CPU column divides each case\'s training wall-clock time (Table 7) by '
    'its saving over that case\'s own native FEM cost (Table 4a); the GPU '
    'columns are Table 10c. The CPU baseline is measured at batch size 1 only, '
    'since the reference solver is not batched.')._p)

els.append(para(
    'What this clarifies is where the operator is useful, and the answer is '
    f'narrower than either figure alone suggests. Against the CPU baseline '
    f'break-even is {lo:.0f}–{hi:.0f} problems, a threshold any parametric '
    'study crosses immediately. Against the GPU baseline it is '
    f'{min(gpu_all):,}–{max(gpu_all):,}, and the batch size assumed moves it '
    'further than the case does. The figure that matters for a deployment '
    'claim is the batch-size-1 column, where problems arrive one at a time and '
    f'break-even is {min(gpu_bs1):,}–{max(gpu_bs1):,}: a user who genuinely has '
    '128 problems in hand would batch the FEM solver too, which is what makes '
    'the right-hand column the least favourable one and the wrong one to quote '
    'as the headline. Any break-even figure taken from this work should name '
    'both the baseline and the batch size it assumes.')._p)

els.append(para(
    'Read across the table, the useful region is a strip rather than a claim: '
    'the operator is worth training when many problems from the same family '
    'must be solved, no GPU-native solver is available or the problems arrive '
    'singly, and a displacement error of roughly 7–12% (Section 8.8: rather '
    'more in the derived quantities) is acceptable. Outside that strip the '
    'finite-element solver is the better tool, and the three B2 rows show why '
    'the training recipe matters as much as the architecture: their corrected '
    'loss-normalized runs (Section 9.1) cost an order of magnitude more to '
    'train than the B1 ones, and break even an order of magnitude later for '
    'exactly that reason.')._p)

# insert after Table 10c's closing "Two limitations" paragraph, before §8.6
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('8.6 Out-of-distribution'):
        anchor = p
        break
assert anchor is not None, 'section 8.6 heading not found'
target = anchor._p.getprevious()
for el in els:
    target.addnext(el)
    target = el
print(f'inserted {len(els)} elements as a new subsection at the end of 8.5')

doc.save(DST)
print('wrote', DST)
